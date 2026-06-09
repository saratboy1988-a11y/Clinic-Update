# -*- coding: utf-8 -*-
import sys
import os
import ctypes
import json
import urllib.request
import urllib.error
import urllib.parse
import subprocess
import zipfile
import tempfile
import re
import hashlib
import platform
import ssl
from datetime import datetime
from PyQt5.QtCore import QTimer

# Set UTF-8 encoding for stdout to handle Khmer characters
if sys.platform == 'win32' and sys.stdout:
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Suppress noisy font warnings regarding OpenType support for Khmer script
os.environ["QT_LOGGING_RULES"] = "qt.qpa.fonts.warning=false"


def _subprocess_no_window_kwargs():
    kwargs = {}
    if os.name == "nt":
        kwargs["creationflags"] = getattr(subprocess, "CREATE_NO_WINDOW", 0)
        startupinfo = subprocess.STARTUPINFO()
        startupinfo.dwFlags |= getattr(subprocess, "STARTF_USESHOWWINDOW", 0)
        startupinfo.wShowWindow = 0
        kwargs["startupinfo"] = startupinfo
    return kwargs


def _github_update_ssl_context():
    try:
        import certifi
        return ssl.create_default_context(cafile=certifi.where())
    except Exception:
        return ssl.create_default_context()


def _urlopen_update_request(request, timeout=20):
    """Open update URLs with a bundled CA store and a GitHub-only fallback."""
    try:
        return urllib.request.urlopen(request, timeout=timeout, context=_github_update_ssl_context())
    except urllib.error.URLError as e:
        reason = getattr(e, "reason", None)
        is_cert_error = isinstance(reason, ssl.SSLError) or "CERTIFICATE_VERIFY_FAILED" in str(e)
        host = urllib.parse.urlparse(request.full_url).netloc.lower()
        is_github_update_host = host in {
            "github.com",
            "raw.githubusercontent.com",
            "objects.githubusercontent.com",
            "release-assets.githubusercontent.com",
        } or host.endswith(".githubusercontent.com")
        if is_cert_error and is_github_update_host:
            fallback_context = ssl._create_unverified_context()
            return urllib.request.urlopen(request, timeout=timeout, context=fallback_context)
        raise


import excel_handler
import report_handler
from version_manager import get_version
from ui_utils import (
    NumericTableWidgetItem, DateTableWidgetItem,
    copy_to_clipboard, open_url, parse_composite_field, get_composite_category, get_composite_value,
    show_error, show_warning, show_success, confirm_delete,
    validate_required_fields, validate_email, validate_not_empty,
    get_save_file_path, get_open_file_path, get_excel_save_path, get_excel_open_path,
    create_link_button, create_telegram_button, create_youtube_button,
    create_auto_field, create_button, create_primary_button, create_success_button, create_danger_button,
    parse_treatment_lines, get_treatment_lines, get_khmer_font, available_khmer_fonts
)
from widgets import CreatorHeader, CreatorFooter, BaseDialog, CompositeInputWidget, AutoInputWidget, FormSection
from constants import (
    CREATOR_NAME, TELEGRAM_URL, YOUTUBE_URL,
    COLOR_ACCENT_BLUE, COLOR_SUCCESS_GREEN, COLOR_ERROR_RED_ALT, COLOR_READ_ONLY_BG,
    MSG_TITLE_ERROR, MSG_TITLE_SUCCESS, MSG_TITLE_WARNING,
    STYLE_AUTO_FIELD, STYLE_BUTTON_PRIMARY, STYLE_BUTTON_SUCCESS, STYLE_BUTTON_DANGER,
    PLACEHOLDER_MACHINE_ID, PLACEHOLDER_EMAIL, PLACEHOLDER_LICENSE_KEY,
    MSG_VALIDATE_NAME, MSG_VALIDATE_DATE, MSG_VALIDATE_AGE,
    MSG_LICENSE_ACTIVATED, MSG_LICENSE_INVALID, MSG_LICENSE_MISSING_DATA,
    ICON_TELEGRAM, ICON_YOUTUBE, ICON_COPY, ICON_GENERATE, ICON_VALIDATE
)

# Get version from centralized version manager
APP_VERSION = get_version()
APP_NAME = "ClinicManager"
ONLINE_LICENSE_CONFIG_FILE = "license_server_config.json"
ONLINE_LICENSE_STORE_FILE = "online_license.json"
ONLINE_LICENSE_GRACE_DAYS = 3
import shutil
import sqlite3
from PyQt5.QtWidgets import (QApplication, QWidget, QLabel, QPushButton, QVBoxLayout,
                             QDialogButtonBox, QTextEdit, QTabWidget, QInputDialog,
                             QHBoxLayout, QGridLayout, QFrame, QLineEdit, QComboBox, QGroupBox, QDateEdit,
                             QMenu, QAction, QTableWidget, QTableWidgetItem, QDialog, QMessageBox,
                             QCalendarWidget, QShortcut, QFileDialog, QCompleter, QStatusBar, QSplitter,
                             QProgressDialog, QHeaderView, QTextBrowser)
from PyQt5.QtWidgets import QScrollArea
from PyQt5.QtGui import QFont, QColor, QFontDatabase, QKeySequence, QDesktopServices
from PyQt5.QtCore import Qt, QDate, QUrl, QLocale
import db
from datetime import datetime, timedelta
import configparser
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas # type: ignore
from matplotlib.figure import Figure
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import license management functions from dedicated module
from license_manager import generate_license_key, validate_license


MESSAGE_BOX_STYLESHEET = """
    QMessageBox {
        background-color: #1e272e;
    }
    QMessageBox QLabel {
        color: #ffffff;
        min-height: 24px;
        min-width: 260px;
        padding: 6px 10px;
    }
    QMessageBox QPushButton {
        background-color: #0fbcf9;
        color: #000000;
        border: none;
        border-radius: 5px;
        padding: 8px 18px;
        min-width: 90px;
        font-weight: bold;
    }
    QMessageBox QPushButton:hover {
        background-color: #00a8ff;
    }
"""


def build_message_box_stylesheet(
    button_bg="#0fbcf9",
    button_color="#000000",
    button_hover="#00a8ff",
    label_min_width=260,
    button_min_width=90,
):
    return f"""
        QMessageBox {{
            background-color: #1e272e;
        }}
        QMessageBox QLabel {{
            color: white;
            font-size: 13px;
            min-height: 24px;
            min-width: {label_min_width}px;
            padding: 6px 10px;
        }}
        QMessageBox QPushButton {{
            background-color: {button_bg};
            color: {button_color};
            font-weight: bold;
            padding: 8px 20px;
            border-radius: 5px;
            min-width: {button_min_width}px;
        }}
        QMessageBox QPushButton:hover {{
            background-color: {button_hover};
        }}
    """


def cloud_sync_period_options(include_entered_url=False):
    options = [
        "📅 ថ្ងៃនេះ (Today Only)",
        "📆 សប្តាហ៍នេះ (This Week)",
        "🗓️ ខែនេះ (This Month)",
        "📋 ខែមុន (Last Month)",
        "⚙️ កំណត់ដោយខ្លួនឯង (Custom Range)",
        "📦 ទាំងអស់ (Full Database)",
    ]
    if include_entered_url:
        options.append("🔗 ប្រើ URL ដែលបានបញ្ចូល (Use Entered URL)")
    return options


def setup_excel_report_workbook(sheet_title, title_text, title_color, header_color, column_widths):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook()
    ws = wb.active  # type: ignore
    ws.title = sheet_title  # type: ignore

    for column, width in column_widths.items():
        ws.column_dimensions[column].width = width  # type: ignore

    styles = {
        "header_fill": PatternFill(start_color=header_color, end_color=header_color, fill_type="solid"),
        "header_font": Font(bold=True, color="FFFFFFFF", size=12),
        "center_align": Alignment(horizontal="center", vertical="center"),
        "section_fill": PatternFill(start_color="FFD5F5E3", end_color="FFD5F5E3", fill_type="solid"),
        "section_font": Font(bold=True, size=12, color="FFC0392B"),
    }

    ws.merge_cells('A1:E1')  # type: ignore
    ws['A1'] = title_text  # type: ignore
    ws['A1'].font = Font(bold=True, size=16, color=title_color)  # type: ignore
    ws['A1'].alignment = styles["center_align"]  # type: ignore
    return wb, ws, styles


def add_excel_headers(ws, row, headers, styles):
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)  # type: ignore
        cell.fill = styles["header_fill"]  # type: ignore
        cell.font = styles["header_font"]  # type: ignore
        cell.alignment = styles["center_align"]  # type: ignore


def save_excel_report(parent, wb, filename, dialog_title, success_message):
    file_path, _ = QFileDialog.getSaveFileName(parent, dialog_title, filename, "Excel Files (*.xlsx)")
    if not file_path:
        return
    wb.save(file_path)
    parent.statusBar.showMessage(f"បាននាំចេញទៅកាន់ {os.path.basename(file_path)}", 5000)
    QMessageBox.information(parent, "ជោគជ័យ", f"{success_message}\n\n{os.path.basename(file_path)}")


def style_progress_dialog(progress):
    khmer_font = get_khmer_font()
    label = QLabel(progress.labelText())
    label.setFont(QFont(khmer_font, 11))
    label.setWordWrap(True)
    label.setMinimumWidth(320)
    label.setMinimumHeight(36)
    label.setStyleSheet(f"""
        QLabel {{
            color: #ffffff;
            background-color: transparent;
            font-family: "{khmer_font}";
            font-size: 13px;
            padding: 4px 6px;
        }}
    """)
    progress.setLabel(label)
    progress.setFont(QFont(khmer_font, 10))
    progress.setMinimumWidth(390)
    progress.setStyleSheet(f"""
        QProgressDialog {{
            background-color: #1e272e;
            color: white;
            font-family: "{khmer_font}";
            min-width: 360px;
        }}
        QProgressDialog QLabel {{
            color: white;
            background-color: transparent;
            font-size: 12px;
            min-width: 320px;
            min-height: 34px;
            padding: 4px;
        }}
        QProgressBar {{
            border: 1px solid #485460;
            border-radius: 4px;
            background-color: #dfe6e9;
            color: #2d3436;
            min-height: 18px;
            text-align: center;
        }}
        QProgressBar::chunk {{
            background-color: #05c46b;
            border-radius: 3px;
        }}
        QPushButton {{
            background-color: #0fbcf9;
            color: black;
            border: none;
            border-radius: 4px;
            padding: 7px 18px;
            font-weight: bold;
            min-width: 74px;
        }}
        QPushButton:hover {{
            background-color: #00cec9;
        }}
    """)
    return progress


def build_patient_share_database(target_db_path, patient_rows):
    """Create a share-safe database containing only the patient table/data."""
    if os.path.exists(target_db_path):
        os.remove(target_db_path)
    with sqlite3.connect(target_db_path) as temp_conn:
        temp_cur = temp_conn.cursor()

        with sqlite3.connect(db.DB_NAME) as source_conn:
            source_cur = source_conn.cursor()
            source_cur.execute(
                "SELECT sql FROM sqlite_master WHERE type='table' AND name='patient'"
            )
            patient_schema_row = source_cur.fetchone()
            if not patient_schema_row or not patient_schema_row[0]:
                raise Exception("Patient table schema was not found in the source database.")

            temp_cur.execute(patient_schema_row[0])
            patient_columns = [
                row[1] for row in source_cur.execute("PRAGMA table_info(patient)").fetchall()
            ]

            if patient_rows:
                placeholders = ",".join(["?"] * len(patient_columns))
                temp_cur.executemany(
                    f"INSERT INTO patient ({','.join(patient_columns)}) VALUES ({placeholders})",
                    patient_rows
                )

        temp_conn.commit()


def get_clinic_appdata_dir(fallback_dir=None):
    """Return a writable ClinicManager app data directory."""
    try:
        import ctypes.wintypes
        CSIDL_LOCAL_APPDATA = 28
        SHGFP_TYPE_CURRENT = 0
        buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
        ctypes.windll.shell32.SHGetFolderPathW(None, CSIDL_LOCAL_APPDATA, None, SHGFP_TYPE_CURRENT, buf)
        appdata_dir = os.path.join(buf.value, 'ClinicManager')
    except Exception:
        base_dir = fallback_dir or os.getenv('TEMP') or os.getenv('TMP') or os.getcwd()
        appdata_dir = os.path.join(base_dir, 'ClinicManager')

    os.makedirs(appdata_dir, exist_ok=True)
    return appdata_dir


def get_writable_settings_file(fallback_dir):
    return os.path.join(get_clinic_appdata_dir(fallback_dir), 'settings.ini')


def get_config_path(filename):
    return os.path.join(get_clinic_appdata_dir(os.path.dirname(os.path.abspath(__file__))), filename)


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS  # type: ignore[attr-defined]
    except Exception:
        base_path = os.path.dirname(os.path.abspath(__file__))

    path = os.path.join(base_path, relative_path)
    if not os.path.exists(path) and getattr(sys, 'frozen', False):
        path = os.path.join(os.path.dirname(sys.executable), relative_path)
    return path


def create_database_backup(backup_dir, prefix="clinic_backup_upload"):
    backup_dir = backup_dir or os.path.join(tempfile.gettempdir(), "ClinicManager", "backups")
    os.makedirs(backup_dir, exist_ok=True)
    backup_path = os.path.join(backup_dir, f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M')}.db")
    if os.path.exists(db.DB_NAME):
        shutil.copy2(db.DB_NAME, backup_path)
        db.logger.info(f"Backup created before upload: {backup_path}")
    return backup_path


def get_machine_id():
    """ទាញយកលេខសម្គាល់កុំព្យូទ័រ (UUID)"""
    """ទាញយកលេខសម្គាល់កុំព្យូទ័រ (UUID) ជាមួយការការពារប្រសិនបើបរាជ័យ (Improved)"""
    # ព្យាយាមវិធីទី១: ប្រើ Windows Registry (MachineGuid) - ជម្រើសល្អបំផុត
    if os.name == "nt":
        try:
            import winreg
            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Cryptography")
            guid, _ = winreg.QueryValueEx(key, "MachineGuid")
            winreg.CloseKey(key)
            if guid:
                return guid
        except Exception:
            pass

        # ព្យាយាមវិធីទី២: ប្រើ WMIC ប្រសិនបើមាន
        try:
            import shutil
            wmic_path = shutil.which("wmic")
            if wmic_path:
                output = subprocess.check_output(
                    [wmic_path, "csproduct", "get", "uuid"],
                    shell=False,
                    **_subprocess_no_window_kwargs()
                ).decode(errors="ignore")
                lines = [line.strip() for line in output.splitlines() if line.strip()]
                if len(lines) > 1:
                    uuid = lines[1]
                    if uuid and "FFFFFFFF" not in uuid and "Not Applicable" not in uuid:
                        return uuid
        except Exception:
            pass

    # ព្យាយាមវិធីទី៣: Fallback ប្រើ MAC Address និងឈ្មោះកុំព្យូទ័រ (ចុងក្រោយបង្អស់)
    try:
        import uuid as _uuid
        node = _uuid.getnode()
        machine_name = platform.node()
        fallback = hashlib.sha256(f"{node}-{machine_name}".encode()).hexdigest()
        return f"ID-{fallback[:20].upper()}"
    except Exception:
        return "UNKNOWN-MACHINE-ID"


def load_online_license_config():
    default_config = {
        "enabled": False,
        "api_base_url": "",
        "app_token": "",
        "strict_online": False,
        "timeout_seconds": 15,
    }
    config_paths = [
        get_config_path(ONLINE_LICENSE_CONFIG_FILE),
        resource_path(ONLINE_LICENSE_CONFIG_FILE),
    ]

    for path in config_paths:
        try:
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    merged = default_config.copy()
                    merged.update(data)
                    merged["api_base_url"] = str(merged.get("api_base_url", "")).rstrip("/")
                    return merged
        except Exception:
            continue

    return default_config


def online_license_is_enabled(config=None):
    config = config or load_online_license_config()
    return bool(config.get("enabled") and config.get("api_base_url"))


def online_license_is_strict(config=None):
    config = config or load_online_license_config()
    return bool(online_license_is_enabled(config) and config.get("strict_online"))


def is_online_license_key(license_key):
    key = str(license_key or "").strip()
    return bool(key and "-" in key)


def _license_api_post(path, payload, config=None):
    config = config or load_online_license_config()
    base_url = str(config.get("api_base_url", "")).rstrip("/")
    if not base_url:
        return False, {"message": "License server is not configured."}

    payload = dict(payload)
    is_google_apps_script = "script.google.com/macros/s/" in base_url
    if is_google_apps_script:
        url = base_url
        payload["_path"] = path
        app_token = str(config.get("app_token", "")).strip()
        if app_token:
            payload["_app_token"] = app_token
    else:
        url = f"{base_url}{path}"

    body = json.dumps(payload).encode("utf-8")
    headers = {
        "Content-Type": "application/json",
        "User-Agent": f"{APP_NAME}/{APP_VERSION}",
    }
    app_token = str(config.get("app_token", "")).strip()
    if app_token and not is_google_apps_script:
        headers["Authorization"] = f"Bearer {app_token}"

    request = urllib.request.Request(url, data=body, headers=headers, method="POST")
    timeout = int(config.get("timeout_seconds") or 15)

    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            response_body = response.read().decode("utf-8", errors="replace")
            if not response_body.strip():
                return False, {"message": "License server returned an empty response."}
            try:
                return True, json.loads(response_body)
            except json.JSONDecodeError:
                message = "License server returned a non-JSON response."
                if is_google_apps_script and (
                    "Script function not found: doPost" in response_body or "doPost" in response_body
                ):
                    message = (
                        "Google Apps Script does not have doPost(e). Paste google_apps_script_license_api.gs "
                        "into Apps Script, save, and deploy a new Web app version."
                    )
                elif is_google_apps_script and "Google Apps Script" in response_body:
                    message = "Google Apps Script returned an HTML error page. Check deployment and permissions."
                return False, {"message": message}
    except urllib.error.HTTPError as e:
        if is_google_apps_script and e.code in (401, 403):
            return False, {
                "message": (
                    "Google Apps Script denied access. Deploy the script as a Web app with "
                    "'Execute as: Me' and 'Who has access: Anyone', then copy the /exec URL again."
                )
            }
        try:
            error_body = e.read().decode("utf-8", errors="replace")
            if not error_body.strip():
                return False, {"message": str(e)}
            try:
                data = json.loads(error_body)
            except json.JSONDecodeError:
                return False, {"message": f"License server returned HTTP {e.code} with a non-JSON response."}
            if isinstance(data, dict) and isinstance(data.get("detail"), dict):
                return False, data["detail"]
            return False, data
        except Exception:
            return False, {"message": str(e)}
    except Exception as e:
        return False, {"message": str(e), "network_error": True}


def _online_license_config_fingerprint(config=None):
    config = config or load_online_license_config()
    raw = f"{config.get('api_base_url', '')}|{config.get('app_token', '')}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _online_token_matches_machine(token, machine_id):
    parts = str(token or "").split("|")
    return len(parts) >= 3 and parts[1] == machine_id


def _online_license_not_expired(saved):
    expires_at = saved.get("expires_at")
    if not expires_at:
        return True
    try:
        expires = datetime.fromisoformat(str(expires_at).replace("Z", "+00:00"))
        now = datetime.now(expires.tzinfo) if expires.tzinfo else datetime.now()
        return now <= expires
    except Exception:
        return False


def save_online_license(data, machine_id=None, config=None):
    payload = dict(data)
    payload["last_valid_at"] = datetime.now().isoformat()
    if machine_id:
        payload["activation_machine_id"] = machine_id
    payload["config_fingerprint"] = _online_license_config_fingerprint(config)
    with open(get_config_path(ONLINE_LICENSE_STORE_FILE), "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)


def load_saved_online_license():
    path = get_config_path(ONLINE_LICENSE_STORE_FILE)
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                return data
    except Exception:
        pass
    return {}


def activate_online_license(email, license_key, machine_id):
    config = load_online_license_config()
    if not online_license_is_enabled(config):
        return False, "Online license server is not configured.", {}

    ok, response = _license_api_post(
        "/api/v1/licenses/activate",
        {
            "email": email,
            "license_key": license_key,
            "machine_id": machine_id,
            "app_version": APP_VERSION,
        },
        config,
    )
    if ok and response.get("ok"):
        save_online_license(response, machine_id, config)
        return True, response.get("message", "Activated"), response

    return False, response.get("message", "Online activation failed."), response


def validate_saved_online_license(machine_id):
    config = load_online_license_config()
    if not online_license_is_enabled(config):
        return False, "Online license is disabled."

    saved = load_saved_online_license()
    token = str(saved.get("token", "")).strip()
    if not token:
        return False, "No online license token found."
    if saved.get("activation_machine_id") and saved.get("activation_machine_id") != machine_id:
        return False, "Saved online license belongs to another machine."
    if not _online_token_matches_machine(token, machine_id):
        return False, "Saved online license token does not match this machine."
    if saved.get("config_fingerprint") and saved.get("config_fingerprint") != _online_license_config_fingerprint(config):
        return False, "Saved online license belongs to another license server."
    if not _online_license_not_expired(saved):
        return False, "Saved online license expired."

    ok, response = _license_api_post(
        "/api/v1/licenses/check",
        {
            "token": token,
            "machine_id": machine_id,
            "app_version": APP_VERSION,
        },
        config,
    )
    if ok and response.get("ok"):
        merged = saved.copy()
        merged.update(response)
        save_online_license(merged, machine_id, config)
        return True, response.get("message", "Online license valid.")

    if response.get("network_error"):
        try:
            last_valid_at = saved.get("last_valid_at")
            if last_valid_at:
                last_dt = datetime.fromisoformat(last_valid_at)
                if last_dt > datetime.now() + timedelta(minutes=5):
                    return False, "Saved online license timestamp is invalid."
                days = (datetime.now() - last_dt).days
                if days <= ONLINE_LICENSE_GRACE_DAYS:
                    return True, f"Offline grace period active ({days}/{ONLINE_LICENSE_GRACE_DAYS} days)."
        except Exception:
            pass

    return False, response.get("message", "Online license invalid.")


class LicenseGeneratorDialog(BaseDialog):
    """ផ្ទាំងសម្រាប់ Admin បង្កើត License Key"""
    def __init__(self):
        super().__init__("License Key Generator (Admin)", size=(450, 450))
        
        # Input fields
        self.txt_mid = QLineEdit()
        self.txt_mid.setPlaceholderText(PLACEHOLDER_MACHINE_ID)
        self.txt_email = QLineEdit()
        self.txt_email.setPlaceholderText(PLACEHOLDER_EMAIL)
        
        self.duration_combo = QComboBox()
        self.duration_combo.addItems(["1 ខែ", "3 ខែ", "6 ខែ", "1 ឆ្នាំ", "Lifetime (មួយជីវិត)"])
        
        self.txt_result = QTextEdit()
        self.txt_result.setPlaceholderText(PLACEHOLDER_LICENSE_KEY)
        self.txt_result.setFixedHeight(100)
        self.txt_result.setReadOnly(True)
        
        # Buttons
        btn_gen = create_button(f"{ICON_GENERATE} Generate License", COLOR_ACCENT_BLUE, "black", self.generate)
        btn_copy = create_button(f"{ICON_COPY} Copy Key", "#3498db", "white", self.copy_key)
        
        # Add to layout
        self.add_widget(QLabel("Machine ID:"))
        self.add_widget(self.txt_mid)
        self.add_widget(QLabel("Email:"))
        self.add_widget(self.txt_email)
        self.add_widget(QLabel("រយះពេលផ្តល់ជូន:"))
        self.add_widget(self.duration_combo)
        self.add_widget(btn_gen)
        self.add_widget(self.txt_result)
        self.add_widget(btn_copy)
    
    def generate(self):
        key = generate_license_key(self.txt_email.text().strip(), self.txt_mid.text().strip(), self.duration_combo.currentText())
        self.txt_result.setPlainText(key)
        show_success(self, "License Key ត្រូវបានបង្កើតជោគជ័យ!")
    
    def copy_key(self):
        key = self.txt_result.toPlainText().strip()
        if not key:
            show_warning(self, "មិនមាន License Key ដើម្បីចម្លងទេ!")
            return
        copy_to_clipboard(key)
        show_success(self, "បានចម្លង License Key រួចរាល់!")


def create_license_contact_links():
    links = QHBoxLayout()
    links.setSpacing(12)
    links.addStretch()

    telegram = QLabel(f'<a href="{TELEGRAM_URL}" style="color: #19d3ff; text-decoration: none;">Telegram</a>')
    telegram.setOpenExternalLinks(True)
    telegram.setStyleSheet("background: transparent; font-size: 13px;")

    separator = QLabel("|")
    separator.setStyleSheet("color: #6c7a89; background: transparent;")

    youtube = QLabel(f'<a href="{YOUTUBE_URL}" style="color: #ff6b6b; text-decoration: none;">YouTube</a>')
    youtube.setOpenExternalLinks(True)
    youtube.setStyleSheet("background: transparent; font-size: 13px;")

    links.addWidget(telegram)
    links.addWidget(separator)
    links.addWidget(youtube)
    links.addStretch()
    return links


class LicenseDialog(BaseDialog):
    def __init__(self, machine_id):
        super().__init__("Register License", size=(620, 640),
                         show_creator_header=False, show_creator_footer=False)
        self.setMinimumSize(620, 640)
        self.setMaximumSize(16777215, 16777215)
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint)  # type: ignore
        self.online_config = load_online_license_config()
        self.setStyleSheet("""
            QDialog {
                background-color: #111820;
                color: #f5f6fa;
            }
            QLabel {
                color: #dfe6e9;
                font-size: 13px;
            }
            QLineEdit {
                background-color: #1e272e;
                color: #ffffff;
                border: 1px solid #485460;
                border-radius: 6px;
                padding: 10px 12px;
                min-height: 30px;
                selection-background-color: #0fbcf9;
                selection-color: #000000;
            }
            QLineEdit:focus {
                border: 1px solid #0fbcf9;
            }
        """)

        title = QLabel("Register License")
        title.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        title.setStyleSheet("font-size: 22px; font-weight: bold; color: #0fbcf9; padding-top: 4px;")
        subtitle = QLabel("Enter your email and license key to activate this computer.")
        subtitle.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        subtitle.setStyleSheet("color: #c8d6e5; font-size: 12px; padding-bottom: 8px;")

        self.txt_mid = QLineEdit(machine_id)
        self.txt_mid.setReadOnly(True)
        self.txt_mid.setStyleSheet("""
            background-color: #202a36;
            color: #ffffff;
            border: 1px solid #0fbcf9;
            border-radius: 6px;
            padding: 10px 12px;
            font-family: Consolas, monospace;
            font-weight: bold;
        """)

        btn_copy = create_button(f"{ICON_COPY} Copy Machine ID", COLOR_ACCENT_BLUE, "white", 
                                  lambda: copy_to_clipboard(machine_id))

        self.txt_email = QLineEdit()
        self.txt_email.setPlaceholderText("Email")

        self.txt_key = QLineEdit()
        self.txt_key.setPlaceholderText("License Key")

        btn_activate = create_button("Activate", COLOR_SUCCESS_GREEN, "white", self.activate)
        for btn in (btn_copy, btn_activate):
            btn.setMinimumHeight(48)
            btn.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]

        def make_card(title_text, body_widget):
            card = QFrame()
            card.setObjectName("licenseCard")
            card.setFrameShape(QFrame.NoFrame)
            card.setStyleSheet("""
                QFrame#licenseCard {
                    background-color: #18212b;
                    border: 1px solid #2f3b48;
                    border-radius: 8px;
                }
            """)
            layout = QVBoxLayout(card)
            layout.setContentsMargins(14, 10, 14, 14)
            layout.setSpacing(10)
            label = QLabel(title_text)
            label.setStyleSheet("background: transparent; border: none; color: #19d3ff; font-size: 14px; font-weight: bold;")
            layout.addWidget(label)
            layout.addWidget(body_widget)
            return card
        
        # Secret shortcut for Generator (Ctrl+Shift+G)
        self.shortcut_gen = QShortcut(QKeySequence("Ctrl+Shift+G"), self)
        self.shortcut_gen.activated.connect(self.open_generator)

        self.add_widget(title)
        self.add_widget(subtitle)
        self.add_widget(make_card("Machine ID", self.txt_mid))
        self.add_widget(btn_copy)
        self.add_widget(make_card("Email", self.txt_email))
        self.add_widget(make_card("License Key", self.txt_key))
        self.add_widget(btn_activate)
        self.add_layout(create_license_contact_links())
    
    def open_generator(self):
        gen = LicenseGeneratorDialog()
        gen.exec_()
    
    def activate(self):
        email = self.txt_email.text().strip()
        key = self.txt_key.text().strip()
        mid = self.txt_mid.text()

        if online_license_is_enabled(self.online_config):
            if not email:
                show_warning(self, "Please enter your email before online activation.")
                return

            is_valid, msg, _ = activate_online_license(email, key, mid)
            if is_valid:
                lic_info = db.get_license_info()
                install_date = lic_info[1] if lic_info else datetime.now().strftime("%Y-%m-%d")
                db.save_license_info(install_date, key, email, mid)
                show_success(self, f"{MSG_LICENSE_ACTIVATED}\nOnline activation successful!")
                self.accept()
                return

            if online_license_is_strict(self.online_config):
                show_warning(self, f"Online activation failed: {msg}")
                return

            if is_online_license_key(key):
                show_warning(self, f"Online activation failed: {msg}")
                return

            fallback_valid, fallback_msg = validate_license(email, mid, key)
            if not fallback_valid:
                show_warning(self, f"Online activation failed: {msg}")
                return
            is_valid, msg = fallback_valid, fallback_msg
        else:
            is_valid, msg = validate_license(email, mid, key)

        if is_valid:
            # Save to DB
            lic_info = db.get_license_info()
            install_date = lic_info[1] if lic_info else datetime.now().strftime("%Y-%m-%d")
            db.save_license_info(install_date, key, email, mid)
            show_success(self, f"{MSG_LICENSE_ACTIVATED} សុពលភាពដល់: {msg}")
            self.accept()
        else:
            show_warning(self, f"{MSG_LICENSE_INVALID}: {msg}")


class LicenseStatusDialog(BaseDialog):
    def __init__(self, parent=None):
        super().__init__("License & Register", size=(720, 700),
                         show_creator_header=False, show_creator_footer=False)
        self.setMinimumSize(720, 700)
        self.setMaximumSize(16777215, 16777215)
        self.parent_app = parent
        self.machine_id = get_machine_id()
        self.setStyleSheet("""
            QDialog {
                background-color: #111820;
                color: #f5f6fa;
            }
            QLabel {
                color: #dfe6e9;
                font-size: 13px;
            }
            QLineEdit {
                background-color: #1e272e;
                color: #ffffff;
                border: 1px solid #485460;
                border-radius: 6px;
                padding: 10px 12px;
                min-height: 30px;
                selection-background-color: #0fbcf9;
                selection-color: #000000;
            }
        """)

        title = QLabel("License & Register")
        title.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        title.setStyleSheet("font-size: 22px; font-weight: bold; color: #0fbcf9; padding-top: 4px;")
        subtitle = QLabel("Manage this computer's license status and registration.")
        subtitle.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        subtitle.setStyleSheet("color: #c8d6e5; font-size: 12px; padding-bottom: 8px;")

        self.txt_mid = QLineEdit(self.machine_id)
        self.txt_mid.setReadOnly(True)
        self.txt_mid.setStyleSheet("""
            background-color: #202a36;
            color: #ffffff;
            border: 1px solid #0fbcf9;
            border-radius: 6px;
            padding: 10px 12px;
            font-family: Consolas, monospace;
            font-weight: bold;
        """)

        self.status_label = QLabel()
        self.status_label.setWordWrap(True)
        self.status_label.setMinimumHeight(138)
        self.status_label.setTextInteractionFlags(Qt.TextSelectableByMouse)  # type: ignore[attr-defined]

        def make_card(title_text, body_widget):
            card = QFrame()
            card.setObjectName("licenseCard")
            card.setFrameShape(QFrame.NoFrame)
            card.setStyleSheet("""
                QFrame#licenseCard {
                    background-color: #18212b;
                    border: 1px solid #2f3b48;
                    border-radius: 8px;
                }
            """)
            layout = QVBoxLayout(card)
            layout.setContentsMargins(14, 10, 14, 14)
            layout.setSpacing(10)
            label = QLabel(title_text)
            label.setStyleSheet("background: transparent; border: none; color: #19d3ff; font-size: 14px; font-weight: bold;")
            layout.addWidget(label)
            layout.addWidget(body_widget)
            return card

        btn_register = create_button("Register / Activate", COLOR_SUCCESS_GREEN, "white", self.open_register)
        btn_check = create_button("Check License", COLOR_ACCENT_BLUE, "black", self.check_license)
        btn_copy = create_button(f"{ICON_COPY} Copy Machine ID", "#576574", "white",
                                 lambda: copy_to_clipboard(self.machine_id))
        btn_remove = create_button("Remove Saved License", COLOR_ERROR_RED_ALT, "white", self.remove_saved_license)
        for btn in (btn_register, btn_check, btn_copy, btn_remove):
            btn.setMinimumHeight(50)
            btn.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]

        row1 = QHBoxLayout()
        row1.setSpacing(10)
        row1.addWidget(btn_register)
        row1.addWidget(btn_check)

        row2 = QHBoxLayout()
        row2.setSpacing(10)
        row2.addWidget(btn_copy)
        row2.addWidget(btn_remove)

        self.add_widget(title)
        self.add_widget(subtitle)
        self.add_widget(make_card("Machine ID", self.txt_mid))
        self.add_widget(make_card("Status", self.status_label))
        self.add_layout(row1)
        self.add_layout(row2)
        self.add_layout(create_license_contact_links())

        self.refresh_status()

    def _get_license_status_text(self):
        online_valid, online_msg = validate_saved_online_license(self.machine_id)
        if online_valid:
            saved = load_saved_online_license()
            license_key = saved.get("license_key", "")
            expires_at = saved.get("expires_at") or "Lifetime"
            return True, f"License valid.\nKey: {license_key}\nExpires: {expires_at}"

        lic_info = db.get_license_info()
        if lic_info and lic_info[2]:
            if is_online_license_key(lic_info[2]):
                return False, f"Saved online license invalid: {online_msg}"
            valid, msg = validate_license(lic_info[3], self.machine_id, lic_info[2])
            if valid:
                return True, f"Offline license valid.\nExpires: {msg}"
            return False, f"Saved offline license invalid: {msg}"

        return False, "No registered license found. Please register to use this application."

    def refresh_status(self):
        valid, text = self._get_license_status_text()
        self.status_label.setText(text)
        color = "#123524" if valid else "#3d2f16"
        border = "#05c46b" if valid else "#ffa801"
        self.status_label.setStyleSheet(
            f"padding: 12px; background-color: {color}; border: 1px solid {border}; "
            "border-radius: 6px; color: white; font-weight: bold; font-size: 14px; line-height: 1.35;"
        )
        return valid, text

    def open_register(self):
        dlg = LicenseDialog(self.machine_id)
        if dlg.exec_() == QDialog.Accepted:
            self.refresh_status()

    def check_license(self):
        valid, text = self.refresh_status()
        if valid:
            QMessageBox.information(self, "License", text)
        else:
            QMessageBox.warning(self, "License", text)

    def remove_saved_license(self):
        reply = QMessageBox.question(
            self,
            "Remove License",
            "Remove saved license from this computer?\n\nលុប License ដែលបានរក្សាទុកចេញពីម៉ាស៊ីននេះ?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return

        removed = False
        for path in (get_config_path(ONLINE_LICENSE_STORE_FILE),):
            try:
                if os.path.exists(path):
                    os.remove(path)
                    removed = True
            except Exception as e:
                QMessageBox.warning(self, "License", f"Could not remove saved online license:\n{e}")
                return

        lic_info = db.get_license_info()
        if lic_info:
            db.save_license_info(lic_info[1], "", "", self.machine_id)
            removed = True

        self.refresh_status()
        if removed:
            QMessageBox.information(self, "License", "Saved license removed.")
        else:
            QMessageBox.information(self, "License", "No saved license was found.")


class ChangePasswordDialog(BaseDialog):
    """Dialog for changing user password"""
    def __init__(self, username):
        super().__init__("🔑 ប្តូរពាក្យសម្ងាត់ (Change Password)", size=(520, 430),
                         show_creator_header=False, show_creator_footer=False)
        self.username = username
        self.setMaximumSize(16777215, 16777215)
        self.setMinimumSize(440, 360)
        self.resize(520, 430)

        outer_layout = self.layout()
        if outer_layout:
            outer_layout.setContentsMargins(10, 10, 10, 10)  # type: ignore[attr-defined]
            outer_layout.setSpacing(8)  # type: ignore[attr-defined]
        self.content_layout.setContentsMargins(18, 14, 18, 14)
        self.content_layout.setSpacing(8)

        khmer_font_name = get_khmer_font()
        input_stylesheet = """
            QLineEdit {
                background-color: #ffffff;
                color: #1f2d3d;
                border: 1px solid #0fbcf9;
                border-radius: 5px;
                padding: 5px 10px;
                font-size: 11pt;
            }
            QLineEdit:focus {
                border: 2px solid #00a8ff;
            }
            QLineEdit::placeholder {
                color: #6c7a89;
            }
        """
        self.setStyleSheet("""
            QDialog {
                background-color: #1e272e;
                color: #ffffff;
            }
            QLabel {
                color: #ffffff;
                background: transparent;
                min-height: 22px;
            }
        """)

        # Input fields
        self.txt_old_pwd = QLineEdit()
        self.txt_old_pwd.setPlaceholderText("បញ្ចូលពាក្យសម្ងាត់ចាស់ (Old Password)")
        self.txt_old_pwd.setEchoMode(QLineEdit.Password)
        self.txt_old_pwd.setMinimumHeight(34)
        self.txt_old_pwd.setFont(QFont(khmer_font_name, 11))
        self.txt_old_pwd.setStyleSheet(input_stylesheet)

        self.txt_new_pwd = QLineEdit()
        self.txt_new_pwd.setPlaceholderText("បញ្ចូលពាក្យសម្ងាត់ថ្មី (New Password)")
        self.txt_new_pwd.setEchoMode(QLineEdit.Password)
        self.txt_new_pwd.setMinimumHeight(34)
        self.txt_new_pwd.setFont(QFont(khmer_font_name, 11))
        self.txt_new_pwd.setStyleSheet(input_stylesheet)

        self.txt_confirm_pwd = QLineEdit()
        self.txt_confirm_pwd.setPlaceholderText("បញ្ចូលពាក្យសម្ងាត់ថ្មីម្តងទៀត (Confirm)")
        self.txt_confirm_pwd.setEchoMode(QLineEdit.Password)
        self.txt_confirm_pwd.setMinimumHeight(34)
        self.txt_confirm_pwd.setFont(QFont(khmer_font_name, 11))
        self.txt_confirm_pwd.setStyleSheet(input_stylesheet)

        # Buttons
        self.btn_change = create_button("🔑 ប្តូរពាក្យសម្ងាត់", COLOR_SUCCESS_GREEN, "white", self.change_password)
        self.btn_change.setMinimumHeight(36)
        self.btn_change.setFont(QFont(khmer_font_name, 11, QFont.Bold))

        # Add to layout
        self.add_widget(QLabel(f"👤 អ្នកប្រើប្រាស់៖ <b>{username}</b>"))
        self.add_widget(QLabel("ពាក្យសម្ងាត់ចាស់ (Old Password):"))
        self.add_widget(self.txt_old_pwd)
        self.add_widget(QLabel("ពាក្យសម្ងាត់ថ្មី (New Password):"))
        self.add_widget(self.txt_new_pwd)
        self.add_widget(QLabel("បញ្ជាក់ពាក្យសម្ងាត់ថ្មី (Confirm New Password):"))
        self.add_widget(self.txt_confirm_pwd)
        self.add_widget(self.btn_change)

        # Set focus
        self.txt_old_pwd.setFocus()

    def change_password(self):
        old_pwd = self.txt_old_pwd.text().strip()
        new_pwd = self.txt_new_pwd.text().strip()
        confirm_pwd = self.txt_confirm_pwd.text().strip()

        if not old_pwd or not new_pwd or not confirm_pwd:
            show_warning(self, "សូមបញ្ចូលគ្រប់ចន្លោះ (Please fill all fields)")
            return

        if new_pwd != confirm_pwd:
            show_warning(self, "ពាក្យសម្ងាត់ថ្មីមិនត្រូវគ្នា (New passwords do not match)")
            self.txt_confirm_pwd.setFocus()
            return

        if len(new_pwd) < 4:
            show_warning(self, "ពាក្យសម្ងាត់ថ្មីត្រូវយ៉ាងតិច ៤ តួអក្សរ (Password must be at least 4 characters)")
            self.txt_new_pwd.setFocus()
            return

        success, message = db.change_password(self.username, old_pwd, new_pwd)
        if success:
            show_success(self, message)
            self.accept()
        else:
            show_warning(self, message)


class SignUpDialog(BaseDialog):
    def __init__(self):
        super().__init__("Sign Up", size=(550, 520))  # Increased size
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint)  # type: ignore
        
        self.user = QLineEdit()
        self.user.setPlaceholderText("Username")
        self.user.setFixedHeight(45)  # Larger input field
        self.user.setFont(QFont("Segoe UI", 14))  # Larger font
        
        self.pwd = QLineEdit()
        self.pwd.setPlaceholderText("Password")
        self.pwd.setEchoMode(QLineEdit.Password)
        self.pwd.setFixedHeight(45)  # Larger input field
        self.pwd.setFont(QFont("Segoe UI", 14))  # Larger font
        
        self.pwd_confirm = QLineEdit()
        self.pwd_confirm.setPlaceholderText("Confirm Password")
        self.pwd_confirm.setEchoMode(QLineEdit.Password)
        self.pwd_confirm.setFixedHeight(45)  # Larger input field
        self.pwd_confirm.setFont(QFont("Segoe UI", 14))  # Larger font

        self.branch_code = QLineEdit()
        self.branch_code.setPlaceholderText("Branch Code (e.g. PP01)")
        self.branch_code.setFixedHeight(45)
        self.branch_code.setFont(QFont("Segoe UI", 14))
        self.branch_code.setInputMethodHints(Qt.ImhLatinOnly | Qt.ImhPreferLatin)  # type: ignore[attr-defined]
        self.branch_code.textChanged.connect(self._force_branch_code_uppercase)
        
        self.btn_register = create_button("Register", COLOR_ACCENT_BLUE, "black", self.register)
        self.btn_register.setFixedHeight(45)
        self.btn_register.setFont(QFont("Segoe UI", 13, QFont.Bold))
        
        # Style for title label
        lbl_title = QLabel("Create Account")
        lbl_title.setFont(QFont("Segoe UI", 16, QFont.Bold))
        lbl_title.setStyleSheet("color: #0fbcf9; padding: 10px;")
        lbl_title.setAlignment(Qt.AlignCenter)  # type: ignore
        
        self.add_widget(lbl_title)
        self.add_widget(QLabel("Username:"))
        self.add_widget(self.user)
        self.add_widget(QLabel("Password:"))
        self.add_widget(self.pwd)
        self.add_widget(QLabel("Confirm Password:"))
        self.add_widget(self.pwd_confirm)
        self.add_widget(QLabel("Branch Code:"))
        self.add_widget(self.branch_code)
        self.add_widget(self.btn_register)

    def _force_branch_code_uppercase(self, text):
        upper_text = str(text or "").upper()
        if text == upper_text:
            return
        cursor_pos = self.branch_code.cursorPosition()
        self.branch_code.blockSignals(True)
        self.branch_code.setText(upper_text)
        self.branch_code.setCursorPosition(cursor_pos)
        self.branch_code.blockSignals(False)
    
    def register(self):
        user = self.user.text()
        pwd = self.pwd.text()
        confirm = self.pwd_confirm.text()
        branch_code = db.normalize_branch_code(self.branch_code.text())
        
        if not user or not pwd or not branch_code:
            show_warning(self, "Please fill all fields")
        elif pwd != confirm:
            show_warning(self, "Passwords do not match")
        elif db.check_username(user):
            show_warning(self, "Username already exists")
        else:
            db.add_user(user, pwd, branch_code)
            show_success(self, "Account created! You can now login.")
            self.accept()


class CustomDateRangeDialog(BaseDialog):
    """Dialog សម្រាប់ជ្រើសរើសកំណត់ថ្ងៃដោយខ្លួនឯង"""
    def __init__(self, parent=None, mode_label="Upload"):
        super().__init__("⚙️ កំណត់រយៈពេលដោយខ្លួនឯង", size=(500, 350),
                        show_creator_header=False, show_creator_footer=False)

        layout = self.content_layout
        layout.setContentsMargins(20, 15, 20, 15)
        layout.setSpacing(12)

        # Title
        title_lbl = QLabel("📅 កំណត់រយៈពេលដែលចង់ Upload")
        title_lbl.setFont(QFont("Segoe UI", 16, QFont.Bold))
        title_lbl.setText(f"📅 កំណត់រយៈពេលដែលចង់ {mode_label}")
        title_lbl.setStyleSheet("""
            color: #00cec9;
            padding: 10px;
            background-color: rgba(0, 206, 201, 0.1);
            border-radius: 8px;
            border-left: 4px solid #00cec9;
        """)
        title_lbl.setWordWrap(True)
        layout.addWidget(title_lbl)

        # Start Date
        layout.addWidget(QLabel("📆 ថ្ងៃចាប់ផ្តើម:"))
        self.start_date = QDateEdit()
        self.start_date.setCalendarPopup(True)
        self.start_date.setDisplayFormat("dd/MM/yyyy")
        self.start_date.setDate(QDate.currentDate().addMonths(-1))
        self.start_date.setFixedHeight(40)
        self.start_date.setFont(QFont("Segoe UI", 12))
        layout.addWidget(self.start_date)

        # End Date
        layout.addWidget(QLabel("📆 ថ្ងៃបញ្ចប់:"))
        self.end_date = QDateEdit()
        self.end_date.setCalendarPopup(True)
        self.end_date.setDisplayFormat("dd/MM/yyyy")
        self.end_date.setDate(QDate.currentDate())
        self.end_date.setFixedHeight(40)
        self.end_date.setFont(QFont("Segoe UI", 12))
        layout.addWidget(self.end_date)

        # Quick buttons
        quick_layout = QHBoxLayout()
        quick_layout.setSpacing(5)

        btn_today = QPushButton("ថ្ងៃនេះ")
        btn_today.clicked.connect(lambda: self.set_date_range(0))
        btn_week = QPushButton("សប្តាហ៍នេះ")
        btn_week.clicked.connect(lambda: self.set_date_range(7))
        btn_month = QPushButton("ខែនេះ")
        btn_month.clicked.connect(lambda: self.set_date_range(30))

        for btn in [btn_today, btn_week, btn_month]:
            btn.setFixedHeight(32)
            btn.setStyleSheet("""
                QPushButton {
                    background-color: #485460;
                    color: white;
                    border-radius: 5px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #57606f;
                }
            """)
            quick_layout.addWidget(btn)

        layout.addLayout(quick_layout)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()

        btn_ok = QPushButton("✅ OK")
        btn_ok.setStyleSheet("""
            QPushButton {
                background-color: #05c46b;
                color: white;
                font-weight: bold;
                padding: 10px 30px;
                border-radius: 6px;
            }
        """)
        btn_ok.clicked.connect(self.accept)
        btn_layout.addWidget(btn_ok)

        btn_cancel = QPushButton("❌ Cancel")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #ff3f34;
                color: white;
                font-weight: bold;
                padding: 10px 30px;
                border-radius: 6px;
            }
        """)
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_cancel)
        btn_layout.addStretch()

        layout.addLayout(btn_layout)

    def set_date_range(self, days):
        """កំណត់រយៈពេលរហ័ស"""
        now = QDate.currentDate()
        self.start_date.setDate(now.addDays(-days))
        self.end_date.setDate(now)


class TelegramBotSetupDialog(BaseDialog):
    """Dialog សម្រាប់កំណត់ Telegram Bot Token និង Chat ID"""
    def __init__(self, parent=None):
        super().__init__("⚙️ កំណត់ Telegram Bot", size=(600, 550),
                        show_creator_header=False, show_creator_footer=False)

        self.khmer_font_name = get_khmer_font()

        layout = self.content_layout
        layout.setContentsMargins(20, 15, 20, 15)
        layout.setSpacing(12)

        # Title
        title_lbl = QLabel("🤖 កំណត់ Telegram Bot សម្រាប់ Upload")
        title_lbl.setFont(QFont(self.khmer_font_name, 16, QFont.Bold))
        title_lbl.setStyleSheet("""
            color: #05c46b;
            padding: 10px;
            background-color: rgba(5, 196, 107, 0.1);
            border-radius: 8px;
            border-left: 4px solid #05c46b;
        """)
        title_lbl.setWordWrap(True)
        layout.addWidget(title_lbl)

        # Instructions
        steps_lbl = QLabel(
            "📋 ជំហានបង្កើត Telegram Bot៖\n\n"
            "១. បើក Telegram → ស្វែងរក @BotFather\n"
            "២. ចុច /start → វាយ /newbot\n"
            "៣. វាយឈ្មោះ Bot (ឧ. MyClinicBot)\n"
            "៤. វាយ Username Bot (ត្រូវបញ្ចប់ដោយ 'bot' ឧ. my_clinic_data_bot)\n"
            "៥. Copy Bot Token ដែលបាន (ឧ. 123456789:ABCdef...)\n\n"
            "📋 រក Chat ID៖\n"
            "៦. បើក Bot របស់អ្នក → ចុច /start\n"
            "៧. បន្ថែម @userinfobot → វាយ /start ក្នុង Bot របស់អ្នក\n"
            "៨. Copy Chat ID (ឧ. 123456789)"
        )
        steps_lbl.setFont(QFont(self.khmer_font_name, 12))
        steps_lbl.setStyleSheet("color: #d2dae2; background: transparent;")
        steps_lbl.setWordWrap(True)
        layout.addWidget(steps_lbl)

        # Bot Token Input
        layout.addWidget(QLabel("🔑 Bot Token:"))
        self.txt_bot = QLineEdit()
        self.txt_bot.setPlaceholderText("123456789:ABCdefGHIjklMNOpqrsTUVwxyz")
        self.txt_bot.setFixedHeight(40)
        self.txt_bot.setFont(QFont("Consolas", 11))
        self.txt_bot.setStyleSheet("""
            QLineEdit {
                background-color: #1e272e;
                color: #00cec9;
                border: 1px solid #485460;
                border-radius: 6px;
                padding: 0 10px;
            }
            QLineEdit:focus {
                border: 1px solid #05c46b;
            }
        """)
        layout.addWidget(self.txt_bot)

        # Chat ID Input
        layout.addWidget(QLabel("💬 Chat ID:"))
        self.txt_chat = QLineEdit()
        self.txt_chat.setPlaceholderText("123456789")
        self.txt_chat.setFixedHeight(40)
        self.txt_chat.setFont(QFont("Consolas", 11))
        self.txt_chat.setStyleSheet("""
            QLineEdit {
                background-color: #1e272e;
                color: #00cec9;
                border: 1px solid #485460;
                border-radius: 6px;
                padding: 0 10px;
            }
            QLineEdit:focus {
                border: 1px solid #05c46b;
            }
        """)
        layout.addWidget(self.txt_chat)

        # Test Button
        btn_test = QPushButton("🧪 ពិនិត្យតភ្ជាប់")
        btn_test.setStyleSheet("""
            QPushButton {
                background-color: #f39c12;
                color: white;
                font-weight: bold;
                padding: 10px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #e67e22;
            }
        """)
        btn_test.clicked.connect(self.test_connection)
        layout.addWidget(btn_test)

        # Save Button
        btn_save = QPushButton("✅ រក្សាទុក")
        btn_save.setStyleSheet("""
            QPushButton {
                background-color: #05c46b;
                color: white;
                font-weight: bold;
                padding: 10px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #06bc5c;
            }
        """)
        btn_save.clicked.connect(self.accept)
        layout.addWidget(btn_save)

    def test_connection(self):
        """ពិនិត្យថា Bot Token និង Chat ID ត្រឹមត្រូវទេ"""
        bot_token = self.txt_bot.text().strip()
        chat_id = self.txt_chat.text().strip()

        if not bot_token or not chat_id:
            QMessageBox.warning(self, "⚠️ មិនពេញលេញ", "សូមបញ្ចូលទាំង Bot Token និង Chat ID!")
            return

        try:
            # ពិនិត្យ Bot Token
            url = f"https://api.telegram.org/bot{bot_token}/getMe"
            response = urllib.request.urlopen(url, timeout=10)
            result = json.loads(response.read().decode('utf-8'))

            if result.get('ok'):
                bot_name = result['result'].get('first_name', 'Unknown')
                QMessageBox.information(
                    self,
                    "✅ តភ្ជាប់ជោគជ័យ!",
                    f"Bot Token ត្រឹមត្រូវ!\n\n"
                    f"🤖 Bot Name: {bot_name}\n"
                    f"💬 Chat ID: {chat_id}\n\n"
                    f"ឥឡូវអ្នកអាចចុច '✅ រក្សាទុក' បាន។"
                )
            else:
                QMessageBox.critical(self, "❌ ខុស", f"Bot Token មិនត្រឹមត្រូវទេ: {result.get('description')}")
        except Exception as e:
            QMessageBox.critical(self, "❌ កំហុស", f"មិនអាចតភ្ជាប់បានទេ: {str(e)}")


class CloudUploadReviewDialog(BaseDialog):
    def __init__(self, patient_rows, parent=None, period_label=""):
        super().__init__("ពិនិត្យទិន្នន័យមុន Upload", size=(1180, 720),
                         show_creator_header=False, show_creator_footer=False,
                         parent=parent)
        self.setMinimumSize(980, 620)
        self.setMaximumSize(16777215, 16777215)
        self.patient_rows = [list(row) for row in patient_rows]
        self.headers = [
            "ID", "កាលបរិច្ឆេទ", "លេខរៀង", "លេខប័ណ្ណ", "ឈ្មោះ", "អាណាព្យាបាល",
            "អាយុ", "ភេទ", "តំបន់", "ផ្ទៃពោះ", "អាសយដ្ឋាន", "ទូរស័ព្ទ",
            "បញ្ជូនមកពី", "ករណីជំងឺ", "រោគសញ្ញា", "អមវេជ្ជសាស្រ្ត",
            "រោគវិនិច្ឆ័យ", "ព្យាបាល", "IMCI", "អាហារូបត្ថម្ភ", "បញ្ជូនទៅ",
            "សេវា", "កំណត់សម្គាល់", "ប្រភេទអ្នកជំងឺ", "Branch"
        ]

        layout = self.content_layout
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(10)

        title = QLabel(f"ពិនិត្យមុនផ្ញើទៅ Cloud - {period_label}")
        title.setFont(QFont(get_khmer_font(), 15, QFont.Bold))
        title.setStyleSheet("color: #0fbcf9;")
        title.setWordWrap(True)
        layout.addWidget(title)

        hint = QLabel("អ្នកអាចកែ cell ឬលុប row មុន Upload។ ការកែនេះសម្រាប់ file ដែលផ្ញើទៅ Cloud ប៉ុណ្ណោះ មិនប៉ះ database ដើមទេ។")
        hint.setWordWrap(True)
        hint.setStyleSheet("color: #d2dae2; background-color: #1e272e; padding: 8px; border-radius: 6px;")
        layout.addWidget(hint)

        self.count_label = QLabel()
        self.count_label.setStyleSheet("color: #05c46b; font-weight: bold;")
        layout.addWidget(self.count_label)

        self.review_table = QTableWidget()
        self.review_table.setColumnCount(len(self.headers))
        self.review_table.setHorizontalHeaderLabels(self.headers)
        self.review_table.setAlternatingRowColors(True)
        self.review_table.setSelectionBehavior(QTableWidget.SelectRows)  # type: ignore[attr-defined]
        self.review_table.setStyleSheet("""
            QTableWidget {
                background-color: #111820;
                alternate-background-color: #1e272e;
                color: #ffffff;
                gridline-color: #485460;
            }
            QTableWidget::item:selected {
                background-color: #0fbcf9;
                color: #000000;
            }
            QHeaderView::section {
                background-color: #2f3640;
                color: #ffffff;
                padding: 6px;
                border: none;
                font-weight: bold;
            }
        """)
        layout.addWidget(self.review_table)

        buttons = QHBoxLayout()
        btn_delete = create_button("លុប Row ដែលបានជ្រើស", COLOR_ERROR_RED_ALT, "white", self.delete_selected_rows)
        btn_cancel = create_button("បោះបង់ Upload", "#7f8c8d", "white", self.reject)
        btn_continue = create_button("បន្ត Upload", COLOR_SUCCESS_GREEN, "white", self.accept)
        for btn in (btn_delete, btn_cancel, btn_continue):
            btn.setMinimumHeight(42)
        buttons.addWidget(btn_delete)
        buttons.addStretch()
        buttons.addWidget(btn_cancel)
        buttons.addWidget(btn_continue)
        layout.addLayout(buttons)

        self.populate_table()

    def populate_table(self):
        self.review_table.setRowCount(len(self.patient_rows))
        for row_idx, row in enumerate(self.patient_rows):
            for col_idx in range(len(self.headers)):
                value = row[col_idx] if col_idx < len(row) else ""
                item = QTableWidgetItem("" if value is None else str(value))
                if col_idx == 0:
                    item.setFlags(item.flags() & ~Qt.ItemIsEditable)  # type: ignore[attr-defined]
                self.review_table.setItem(row_idx, col_idx, item)
        self.review_table.resizeColumnsToContents()
        self.review_table.horizontalHeader().setStretchLastSection(True)
        self.update_count_label()

    def update_count_label(self):
        self.count_label.setText(f"ចំនួនត្រូវ Upload: {self.review_table.rowCount()} នាក់")

    def delete_selected_rows(self):
        rows = sorted({index.row() for index in self.review_table.selectedIndexes()}, reverse=True)
        if not rows:
            QMessageBox.information(self, "Info", "សូមជ្រើសរើស row ដែលចង់លុបចេញពី Upload។")
            return
        reply = QMessageBox.question(
            self,
            "លុប Row",
            f"លុប {len(rows)} row ចេញពី package Upload មែនទេ?\n\nDatabase ដើមមិនត្រូវបានប៉ះពាល់ទេ។",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return
        for row_idx in rows:
            self.review_table.removeRow(row_idx)
        self.update_count_label()

    def get_patient_rows(self):
        rows = []
        for row_idx in range(self.review_table.rowCount()):
            row = []
            for col_idx in range(len(self.headers)):
                item = self.review_table.item(row_idx, col_idx)
                row.append(item.text() if item else "")
            rows.append(row)
        return rows


class CloudSyncHelpDialog(BaseDialog):
    """Dialog ដែលបង្ហាញការណែនាំអំពីរបៀបបង្កើត URL សម្រាប់ Cloud Sync"""
    def __init__(self, parent=None):
        super().__init__("☁️ ការណែនាំ Cloud Sync - របៀបបង្កើត URL", size=(750, 650),
                        show_creator_header=False, show_creator_footer=False,
                        parent=parent)
        self.setWindowModality(Qt.WindowModal)  # type: ignore[attr-defined]

        # Get Khmer font
        self.khmer_font = get_khmer_font()

        layout = self.content_layout
        layout.setContentsMargins(20, 15, 20, 15)
        layout.setSpacing(12)

        # Title
        title_lbl = QLabel("📚 របៀបបង្កើត URL សម្រាប់ Cloud Sync")
        title_lbl.setFont(QFont(self.khmer_font, 16, QFont.Bold))
        title_lbl.setStyleSheet("""
            color: #00cec9;
            padding: 10px;
            background-color: rgba(0, 206, 201, 0.1);
            border-radius: 8px;
            border-left: 4px solid #00cec9;
        """)
        title_lbl.setWordWrap(True)
        layout.addWidget(title_lbl)

        # Description
        desc_lbl = QLabel(
            "Cloud Sync អនុញ្ញាតឱ្យអ្នកចែករំលែកទិន្នន័យអ្នកជំងឺរវាងកុំព្យូទ័រច្រើន។\n"
            "ខាងក្រោមនេះជាជម្រើសក្នុងការបង្កើត URL ដើម្បីទាញយកទិន្នន័យ៖"
        )
        desc_lbl.setFont(QFont(self.khmer_font, 11))
        desc_lbl.setStyleSheet("color: #ecf0f1; padding: 8px; background: transparent;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Create scrollable content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("""
            QScrollArea {
                background-color: transparent;
                border: 1px solid #485460;
                border-radius: 8px;
            }
            QScrollBar:vertical {
                background-color: #2f3640;
                width: 15px;
            }
            QScrollBar::handle:vertical {
                background-color: #485460;
                border-radius: 7px;
            }
        """)

        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        content_layout.setSpacing(15)
        content_layout.setContentsMargins(15, 15, 15, 15)

        # ===== Option 1: Google Drive =====
        content_layout.addWidget(self._create_option_card(
            icon="📁",
            title="ជម្រើសទី ១: Google Drive (ងាយស្រួលបំផុត)",
            steps=[
                "១. បើក Google Drive (drive.google.com)",
                "២. Upload file clinic.db ចូលទៅក្នុង Google Drive",
                "៣. ចុចស្តាំលើ file → ចែករំលែក (Share) → កំណត់ជា 'Anyone with the link'",
                "៤. Copy Link ដែលបាន",
                "៥. ប្តូរ URL ពី:",
                "   https://drive.google.com/file/d/FILE_ID/view?usp=sharing",
                "   ទៅជា:",
                "   https://drive.google.com/uc?export=download&id=FILE_ID",
            ],
            example="ឧទាហរណ៍ URL:\nhttps://drive.google.com/uc?export=download&id=1ABC123xyz",
            color="#4285f4"
        ))

        # ===== Option 2: GitHub =====
        content_layout.addWidget(self._create_option_card(
            icon="🐙",
            title="ជម្រើសទី ២: GitHub (ឥតគិតថ្លៃ)",
            steps=[
                "១. បង្កើត GitHub Account (github.com)",
                "២. បង្កើត Repository ថ្មី (ឧ. clinic-data)",
                "៣. Upload file clinic.db ចូលទៅក្នុង Repository",
                "៤. ចុចលើ file → ចុច 'Raw' button",
                "៥. Copy URL ពី browser address bar",
            ],
            example="ឧទាហរណ៍ URL:\nhttps://raw.githubusercontent.com/USERNAME/REPO/main/clinic.db",
            color="#6e40c9"
        ))

        # ===== Option 3: Dropbox =====
        content_layout.addWidget(self._create_option_card(
            icon="📦",
            title="ជម្រើសទី ៣: Dropbox (ងាយស្រួល)",
            steps=[
                "១. បើក Dropbox (dropbox.com)",
                "២. Upload file clinic.db",
                "៣. ចុច 'Share' → 'Create link'",
                "៤. Copy Link",
                "៥. ប្តូរ URL ពី:",
                "   https://www.dropbox.com/s/FILE_ID/clinic.db?dl=0",
                "   ទៅជា:",
                "   https://dl.dropboxusercontent.com/s/FILE_ID/clinic.db",
            ],
            example="ឧទាហរណ៍ URL:\nhttps://dl.dropboxusercontent.com/s/abc123/clinic.db",
            color="#0061ff"
        ))

        # ===== Option 4: Local Network =====
        content_layout.addWidget(self._create_option_card(
            icon="🏠",
            title="ជម្រើសទី ៤: Local Network (LAN)",
            steps=[
                "១. ដាក់ file clinic.db ក្នុង Shared Folder",
                "២. ចុចស្តាំលើ folder → Properties → Sharing → Share",
                "៣. Copy Network Path",
                "៤. ប្រើ URL ជាទម្រង់: file://COMPUTER_NAME/shared/clinic.db",
            ],
            example="ឧទាហរណ៍ URL:\nfile://DESKTOP-ABC123/Users/Public/clinic.db",
            color="#20bf6b"
        ))

        # ===== Important Notes =====
        notes_card = self._create_option_card(
            icon="⚠️",
            title="ចំណាំសំខាន់ៗ",
            steps=[
                "✓ URL ត្រូវចាប់ផ្តើមដោយ http:// ឬ https:// (ឬ file:// សម្រាប់ LAN)",
                "✓ File ត្រូវតែជា SQLite database (.db) ត្រឹមត្រូវ",
                "✓ ទិន្នន័យនឹងត្រូវបាន Backup មុនពេល Sync គ្រប់ពេល",
                "✓ អ្នកអាចជ្រើសរើស 'Replace' (ជំនួស) ឬ 'Merge' (បញ្ចូលគ្នា)",
                "✗ កុំប្រើ URL ពីប្រភពដែលមិនទុកចិត្ត (អាចមាន malware)",
            ],
            example="",
            color="#ffa801"
        )
        content_layout.addWidget(notes_card)

        scroll.setWidget(content_widget)
        layout.addWidget(scroll)

        # Buttons
        btn_layout = QHBoxLayout()

        btn_copy_example = QPushButton("📋 Copy Example URL")
        btn_copy_example.setStyleSheet("""
            QPushButton {
                background-color: #00cec9;
                color: black;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #00b5b0;
            }
        """)
        btn_copy_example.clicked.connect(lambda: copy_to_clipboard(
            "https://raw.githubusercontent.com/USERNAME/REPO/main/clinic.db"
        ))

        btn_close = QPushButton("✅ យល់ហើយ")
        btn_close.setStyleSheet("""
            QPushButton {
                background-color: #05c46b;
                color: white;
                font-weight: bold;
                padding: 10px 20px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #06bc5c;
            }
        """)
        btn_close.clicked.connect(self.accept)

        btn_layout.addWidget(btn_copy_example)
        btn_layout.addStretch()
        btn_layout.addWidget(btn_close)
        layout.addLayout(btn_layout)

    def _create_option_card(self, icon, title, steps, example, color):
        """Create a styled card for each option"""
        card = QFrame()
        card.setStyleSheet(f"""
            QFrame {{
                background-color: rgba(30, 39, 46, 0.8);
                border: 1px solid {color};
                border-radius: 10px;
                border-left: 4px solid {color};
                padding: 12px;
            }}
        """)

        layout = QVBoxLayout(card)
        layout.setSpacing(8)

        # Title
        title_lbl = QLabel(f"{icon} {title}")
        title_lbl.setFont(QFont(self.khmer_font, 13, QFont.Bold))
        title_lbl.setStyleSheet(f"color: {color}; background: transparent;")
        layout.addWidget(title_lbl)

        # Steps
        steps_text = "\n".join([f"  • {step}" for step in steps])
        steps_lbl = QLabel(steps_text)
        steps_lbl.setFont(QFont(self.khmer_font, 12))
        steps_lbl.setStyleSheet("color: #d2dae2; background: transparent;")
        steps_lbl.setWordWrap(True)
        layout.addWidget(steps_lbl)

        # Example
        if example:
            example_lbl = QLabel(example)
            example_lbl.setFont(QFont("Consolas", 9))
            example_lbl.setStyleSheet(f"""
                color: {color};
                background-color: rgba(0, 0, 0, 0.3);
                padding: 8px;
                border-radius: 5px;
                font-weight: bold;
            """)
            example_lbl.setWordWrap(True)
            layout.addWidget(example_lbl)

        return card


class LoginDialog(BaseDialog):
    def __init__(self):
        # Disable creator header and footer, we'll add our own
        super().__init__("System Login", size=(550, 580),
                        show_creator_header=False, show_creator_footer=False)
        self._configure_login_window_for_screen()
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint)  # type: ignore

        # កំណត់ផ្លូវឯកសារ
        if getattr(sys, 'frozen', False):
            self.base_dir = os.path.dirname(sys.executable)
        else:
            self.base_dir = os.path.dirname(os.path.abspath(__file__))

        # កំណត់ Backup directory (use AppData to avoid permission issues)
        self.backup_dir = os.path.join(get_clinic_appdata_dir(os.getenv('TEMP') or os.getenv('TMP') or os.getcwd()), 'backups')
        os.makedirs(self.backup_dir, exist_ok=True)

        # កំណត់ current_user (នឹងត្រូវបានកំណត់ពេល Login ជោគជ័យ)
        self.current_user = ""
        self.user_context = None
        self.branch_code = "MAIN"
        self.active_branch_code = None

        # Load settings
        self.config = configparser.ConfigParser()
        
        # Use AppData for writable settings file
        writable_settings_file = get_writable_settings_file(self.base_dir)

        # Try to read from AppData first, then fall back to program dir
        settings_to_read = writable_settings_file if os.path.exists(writable_settings_file) else os.path.join(self.base_dir, 'settings.ini')
        self.settings_file = writable_settings_file  # Point to writable location for future saves
        
        if os.path.exists(settings_to_read):
            # បើក file ជាមួយ utf-8 encoding (សម្រាប់ Khmer characters)
            with open(settings_to_read, 'r', encoding='utf-8-sig') as f:
                self.config.read_file(f)
        
        # Set gradient background for main dialog
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #1e272e, stop:1 #34495e);
            }
        """ + MESSAGE_BOX_STYLESHEET)

        layout = self.content_layout
        compact_login = getattr(self, "compact_login_layout", False)
        outer_layout = self.layout()
        if outer_layout:
            outer_layout.setContentsMargins(6 if compact_login else 10, 6 if compact_login else 10, 6 if compact_login else 10, 6 if compact_login else 10)  # type: ignore[attr-defined]
            outer_layout.setSpacing(6 if compact_login else 10)  # type: ignore[attr-defined]
        layout.setContentsMargins(14 if compact_login else 30, 8 if compact_login else 20, 14 if compact_login else 30, 8 if compact_login else 20)
        layout.setSpacing(6 if compact_login else 12)

        # Use Khmer-compatible font
        khmer_font_name = get_khmer_font()

        # Common stylesheet for input fields
        input_stylesheet = """
            QLineEdit {
                padding: 5px 10px;
                border: 2px solid #bdc3c7;
                border-radius: 6px;
                background-color: white;
                color: #1f2d3d;
                font-size: 11pt;
            }
            QLineEdit:focus {
                border: 2px solid #3498db;
            }
            QLineEdit::placeholder {
                color: #6c7a89;
            }
        """

        # Common button cursor type
        hand_cursor = Qt.PointingHandCursor  # type: ignore[attr-defined]

        # Load English keyboard layout (US)
        try:
            self.english_keyboard = ctypes.windll.user32.LoadKeyboardLayoutA("00000409", 1)  # type: ignore[union-attr]  # US English
        except:
            self.english_keyboard = None

        # ===== TITLE SECTION =====
        title_label = QLabel("🏥 កម្មវិធីគ្រប់គ្រងព័ត៌មានអ្នកជំងឺ OPD")
        title_label.setFont(QFont(khmer_font_name, 16 if compact_login else 18, QFont.Bold))
        title_label.setStyleSheet(f"""
            color: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #00d2ff, stop:1 #3a7bd5);
            padding: {6 if compact_login else 10}px;
            background-color: rgba(0, 0, 0, 0.3);
            border-radius: {8 if compact_login else 10}px;
            border: 1px solid rgba(0, 210, 255, 0.3);
        """)
        title_label.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        title_label.setWordWrap(True)
        layout.addWidget(title_label)

        # ===== LOGIN FORM CONTAINER =====
        form_container = QWidget()
        form_container.setStyleSheet("""
            QWidget {
                background-color: rgba(255, 255, 255, 0.95);
                border-radius: 12px;
                border: 1px solid rgba(0, 210, 255, 0.5);
            }
        """)
        form_layout = QVBoxLayout(form_container)
        form_layout.setContentsMargins(14 if compact_login else 25, 10 if compact_login else 20, 14 if compact_login else 25, 10 if compact_login else 20)
        form_layout.setSpacing(6 if compact_login else 12)

        # Welcome message
        welcome_lbl = QLabel("សូមស្វាគមន៍!")
        welcome_lbl.setFont(QFont(khmer_font_name, 14, QFont.Bold))
        welcome_lbl.setStyleSheet("color: #2c3e50;")
        welcome_lbl.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        form_layout.addWidget(welcome_lbl)

        # Username section
        lbl_user = QLabel("👤 Username")
        lbl_user.setFont(QFont("Segoe UI", 10, QFont.Bold))
        lbl_user.setStyleSheet("color: #3498db;")

        self.user = QLineEdit()
        self.user.setPlaceholderText("បញ្ចូលឈ្មោះអ្នកប្រើប្រាស់...")
        self.user.setMinimumHeight(34 if compact_login else 36)
        self.user.setFont(QFont(khmer_font_name, 11))
        self.user.setStyleSheet(input_stylesheet)
        # Set English keyboard layout for Username
        self.user.setInputMethodHints(Qt.ImhLatinOnly | Qt.ImhPreferLatin)  # type: ignore[attr-defined]

        # Password section
        lbl_pwd = QLabel("🔒 Password")
        lbl_pwd.setFont(QFont("Segoe UI", 10, QFont.Bold))
        lbl_pwd.setStyleSheet("color: #e74c3c;")

        self.pwd = QLineEdit()
        self.pwd.setPlaceholderText("បញ្ចូលពាក្យសម្ងាត់...")
        self.pwd.setEchoMode(QLineEdit.Password)
        self.pwd.setMinimumHeight(34 if compact_login else 36)
        self.pwd.setFont(QFont(khmer_font_name, 11))
        self.pwd.setStyleSheet(input_stylesheet)
        # Set English keyboard layout for Password
        self.pwd.setInputMethodHints(Qt.ImhLatinOnly | Qt.ImhPreferLatin)  # type: ignore[attr-defined]

        form_layout.addWidget(lbl_user)
        form_layout.addWidget(self.user)
        form_layout.addWidget(lbl_pwd)
        form_layout.addWidget(self.pwd)
        layout.addWidget(form_container)

        # ===== BUTTONS SECTION =====
        button_container = QWidget()
        button_layout = QVBoxLayout(button_container)
        button_layout.setContentsMargins(0, 0, 0, 0)
        button_layout.setSpacing(8)

        # Login button
        self.btn_login = QPushButton("🔓 Login / ចូលប្រើប្រាស់")
        self.btn_login.setFixedHeight(38 if compact_login else 42)
        self.btn_login.setFont(QFont(khmer_font_name, 12, QFont.Bold))
        self.btn_login.setCursor(hand_cursor)
        self.btn_login.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #11998e, stop:1 #38ef7d);
                color: white;
                font-weight: bold;
                font-size: 12px;
                border-radius: 8px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #0f8a7f, stop:1 #2ed86f);
            }
        """)
        self.btn_login.clicked.connect(self.check_login)

        # Sign Up button
        self.btn_signup = QPushButton("📝 Sign Up / ចុះឈ្មោះថ្មី")
        self.btn_signup.setFixedHeight(34 if compact_login else 38)
        self.btn_signup.setFont(QFont(khmer_font_name, 12, QFont.Bold))
        self.btn_signup.setCursor(hand_cursor)
        self.btn_signup.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #667eea, stop:1 #764ba2);
                color: white;
                font-weight: bold;
                font-size: 12px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #5a6fd6, stop:1 #6a4190);
            }
        """)
        self.btn_signup.clicked.connect(self.show_signup)

        def style_cloud_button(button, start_color, end_color, hover_start, hover_end):
            button.setMinimumHeight(42 if compact_login else 46)
            button.setMinimumWidth(130)
            button.setFont(QFont(get_khmer_font(), 11, QFont.Bold))
            button.setCursor(hand_cursor)
            button.setStyleSheet(f"""
            QPushButton {{
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 {start_color}, stop:1 {end_color});
                color: white;
                font-weight: bold;
                font-size: 12px;
                border-radius: 7px;
                padding: 7px 10px;
                border: 1px solid rgba(255,255,255,0.14);
            }}
            QPushButton:hover {{
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 {hover_start}, stop:1 {hover_end});
            }}
            QPushButton:pressed {{
                padding-top: 8px;
                padding-left: 11px;
            }}
        """)

        # Cloud action buttons
        self.btn_sync = QPushButton("☁️ ទាញពី Cloud")
        style_cloud_button(self.btn_sync, "#3d5368", "#25313d", "#4d657d", "#314252")
        self.btn_sync.clicked.connect(self.sync_data_initial)

        self.btn_upload = QPushButton("⬆️ ផ្ញើទៅ Cloud")
        style_cloud_button(self.btn_upload, "#00b894", "#00a36c", "#11cfa6", "#00b77a")
        self.btn_upload.clicked.connect(self.upload_to_cloud)

        self.btn_edit_cloud = QPushButton("✏️ កែ Cloud")
        style_cloud_button(self.btn_edit_cloud, "#0984e3", "#006bbf", "#1e90ff", "#0875cf")
        self.btn_edit_cloud.clicked.connect(self.edit_uploaded_cloud_data)

        self.btn_delete_cloud = QPushButton("🗑️ លុប Cloud")
        style_cloud_button(self.btn_delete_cloud, "#e74c3c", "#c0392b", "#ff5a4d", "#d64535")
        self.btn_delete_cloud.clicked.connect(self.delete_uploaded_cloud_data)

        self.btn_cloud_token = QPushButton("🔑 Token")
        style_cloud_button(self.btn_cloud_token, "#8e44ad", "#6c3483", "#9b59b6", "#7d3c98")
        self.btn_cloud_token.clicked.connect(self.show_github_token_setup)

        self.btn_sync_help = QPushButton("❓ ជំនួយ")
        style_cloud_button(self.btn_sync_help, "#f39c12", "#d87900", "#ffad22", "#ec8d00")
        self.btn_sync_help.clicked.connect(self.show_cloud_sync_help)

        button_layout.addWidget(self.btn_login)
        button_layout.addWidget(self.btn_signup)

        # Cloud actions are available from the main app Cloud Sync menu after login.
        sync_layout = QGridLayout()
        sync_layout.setSpacing(8)
        sync_layout.addWidget(self.btn_sync, 0, 0)
        sync_layout.addWidget(self.btn_upload, 0, 1)
        sync_layout.addWidget(self.btn_edit_cloud, 1, 0)
        sync_layout.addWidget(self.btn_delete_cloud, 1, 1)
        sync_layout.addWidget(self.btn_cloud_token, 2, 0)
        sync_layout.addWidget(self.btn_sync_help, 2, 1)
        sync_layout.setColumnStretch(0, 1)
        sync_layout.setColumnStretch(1, 1)

        layout.addWidget(button_container)

        # ===== CREATOR INFO SECTION =====
        creator_frame = QWidget()
        creator_frame.setStyleSheet("""
            QWidget {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #232526, stop:1 #414345);
                border-radius: 10px;
                border: 1px solid rgba(52, 152, 219, 0.4);
            }
        """)
        creator_frame.setMinimumHeight(54 if compact_login else 70)
        creator_frame_layout = QVBoxLayout(creator_frame)
        creator_frame_layout.setContentsMargins(10 if compact_login else 15, 6 if compact_login else 10, 10 if compact_login else 15, 6 if compact_login else 10)
        creator_frame_layout.setSpacing(2 if compact_login else 4)

        # Line 1: Created by
        creator_line1 = QLabel('💻 <span style="color: #f5f6fa;">បង្កើតដោយ៖</span> <b style="color: #5dade2;">នូរ សារ៉ាត់ (NOU SARAT)</b>')
        creator_line1.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        creator_line1.setFont(QFont(khmer_font_name, 9 if compact_login else 10, QFont.Bold))
        creator_line1.setStyleSheet("background: transparent; color: #f5f6fa;")
        creator_line1.setTextFormat(Qt.RichText)  # type: ignore[attr-defined]

        # Line 2: Social links
        social_line = QLabel('''
            <a href="https://t.me/nousarat" style="color: #0088cc; text-decoration: none; font-weight: bold;">✈️ Telegram</a>
            <span style="color: #d2dae2;">&nbsp;|&nbsp;</span>
            <a href="https://www.youtube.com/@TechFree2026" style="color: #ff0000; text-decoration: none; font-weight: bold;">📺 YouTube</a>
        ''')
        social_line.setOpenExternalLinks(True)
        social_line.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        social_line.setFont(QFont(khmer_font_name, 8 if compact_login else 9))
        social_line.setStyleSheet("background: transparent; color: #d2dae2;")
        social_line.setTextFormat(Qt.RichText)  # type: ignore[attr-defined]

        creator_frame_layout.addWidget(creator_line1)
        creator_frame_layout.addWidget(social_line)
        layout.addWidget(creator_frame)

    def _configure_login_window_for_screen(self):
        """Allow the login dialog to grow on high-DPI and smaller laptop screens."""
        self.setMaximumSize(16777215, 16777215)

        screen = QApplication.primaryScreen()
        if not screen:
            self.setMinimumSize(500, 620)
            self.resize(560, 680)
            return

        available = screen.availableGeometry()
        self.compact_login_layout = available.height() < 760 or available.width() < 900
        width = min(560, max(500, int(available.width() * 0.9)), available.width())
        height = min(720, max(620, int(available.height() * 0.9)), available.height())

        self.setMinimumSize(min(500, available.width()), min(560, available.height()))
        self.resize(width, height)

        frame = self.frameGeometry()
        frame.moveCenter(available.center())
        self.move(frame.topLeft())

    def switch_to_english(self):
        """Switch keyboard layout to English (US)"""
        if self.english_keyboard:
            try:
                hwnd = int(self.winId())
                ctypes.windll.user32.ActivateKeyboardLayout(self.english_keyboard, 0)
            except:
                pass

    def show_cloud_sync_help(self):
        """បង្ហាញការណែនាំអំពីរបៀបបង្កើត URL សម្រាប់ Cloud Sync"""
        dialog = CloudSyncHelpDialog(self)
        dialog.raise_()
        dialog.activateWindow()
        dialog.exec_()

    def upload_to_cloud(self):
        """ផ្ញើទិន្នន័យទៅ Cloud (GitHub)"""
        # ពិនិត្យថា Git ដំឡើងហើយឬនៅ
        if not self._check_git_installed():
            QMessageBox.critical(
                self,
                "❌ មិនមាន Git",
                "Git មិនទាន់បានដំឡើងលើកុំព្យូទ័រនេះទេ!\n\n"
                "សូមដំឡើង Git មុនពេលប្រើមុខងារនេះ៖\n"
                "https://git-scm.com/downloads"
            )
            return

        # ពិនិត្យថាកំណត់ GitHub Repository សម្រាប់ Cloud Sync ហើយឬនៅ
        repo_url = self.config.get('CATEGORIES', 'cloud_sync_repo_url', fallback="").strip()
        if not repo_url:
            reply = QMessageBox.question(
                self,
                "⚙️ មិនមានការកំណត់ GitHub",
                "អ្នកត្រូវកំណត់ GitHub Repository URL សម្រាប់ Cloud Sync មុនពេលប្រើមុខងារនេះ។\n\n"
                "តើអ្នកចង់កំណត់ឥឡូវនេះទេ?",
                QMessageBox.Yes | QMessageBox.No
            )

            if reply == QMessageBox.Yes:
                self.show_github_setup()
            return

        if "clinic-update" in repo_url.lower():
            QMessageBox.critical(
                self,
                "❌ Repository មិនត្រឹមត្រូវ",
                "មុខងារ Cloud Sync មិនអាចប្រើ Clinic-Update repository បានទេ ព្រោះវាជា repository សម្រាប់ Update កម្មវិធី។\n\n"
                "សូមកំណត់ repository ផ្សេងសម្រាប់ទិន្នន័យ Cloud Sync។"
            )
            self.show_github_setup()
            return

        # Upload ទិន្នន័យ
        self._upload_to_github(repo_url)

    def edit_uploaded_cloud_data(self):
        """Download an uploaded Cloud database, edit rows, then upload it back."""
        if not self._check_git_installed():
            QMessageBox.critical(
                self,
                "❌ មិនមាន Git",
                "Git មិនទាន់បានដំឡើងលើកុំព្យូទ័រនេះទេ!\n\n"
                "សូមដំឡើង Git មុនពេលកែទិន្នន័យ Cloud:\n"
                "https://git-scm.com/downloads"
            )
            return

        repo_url = self.config.get('CATEGORIES', 'cloud_sync_repo_url', fallback="").strip()
        if not repo_url:
            reply = QMessageBox.question(
                self,
                "⚙️ មិនមាន GitHub Repository",
                "ត្រូវកំណត់ GitHub Repository URL មុនពេលកែទិន្នន័យ Cloud។\n\n"
                "តើចង់កំណត់ឥឡូវនេះទេ?",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                self.show_github_setup()
            return

        if "clinic-update" in repo_url.lower():
            QMessageBox.critical(
                self,
                "❌ Repository មិនត្រឹមត្រូវ",
                "Repository នេះសម្រាប់ Update កម្មវិធី មិនមែនសម្រាប់ Cloud Sync ទេ។\n\n"
                "សូមកំណត់ repository ផ្សេងសម្រាប់ទិន្នន័យ Cloud Sync។"
            )
            self.show_github_setup()
            return

        if not self._ensure_cloud_token_for_write():
            return

        download_file_name, period_label, ok = self._get_cloud_download_period_choice()
        if not ok:
            return

        if download_file_name:
            cloud_url = self._github_repo_file_raw_url(repo_url, download_file_name)
            upload_file_name = download_file_name
        else:
            saved_url = self.config.get('CATEGORIES', 'cloud_sync_url', fallback="").strip()
            cloud_url, url_ok = self._get_cloud_sync_url_input(saved_url)
            if not url_ok or not cloud_url.strip():
                return
            cloud_url = cloud_url.strip()
            parsed_path = urllib.parse.urlparse(cloud_url).path
            upload_file_name = os.path.basename(parsed_path) or "clinic_full.db"
            if not upload_file_name.lower().endswith(".db"):
                upload_file_name = "clinic_full.db"

        if not cloud_url:
            QMessageBox.warning(
                self,
                "⚠️ URL មិនត្រឹមត្រូវ",
                "មិនអាចបង្កើត Cloud URL ពី GitHub Repository បានទេ។\n\n"
                "សូមពិនិត្យ GitHub Repository URL ម្ដងទៀត។"
            )
            return

        ok_url, url_error = self._preflight_cloud_sync_url(cloud_url)
        if not ok_url:
            saved_url = self.config.get('CATEGORIES', 'cloud_sync_url', fallback="").strip()
            if download_file_name and saved_url and saved_url != cloud_url:
                reply = QMessageBox.question(
                    self,
                    "⚠️ រក file មិនឃើញ",
                    f"រកមិនឃើញ file នេះនៅ Cloud:\n{download_file_name}\n\n"
                    "វាអាចមិនទាន់បាន Upload ឬ file period ចាស់ត្រូវបានជំនួសដោយ Upload ថ្មី។\n\n"
                    "តើចង់បើក file ចុងក្រោយដែលបាន Upload វិញទេ?",
                    QMessageBox.Yes | QMessageBox.No
                )
                if reply == QMessageBox.Yes:
                    cloud_url = saved_url
                    parsed_path = urllib.parse.urlparse(cloud_url).path
                    upload_file_name = os.path.basename(parsed_path) or upload_file_name
                    period_label = "Cloud file ចុងក្រោយ"
                    ok_url, url_error = self._preflight_cloud_sync_url(cloud_url)
                else:
                    return
            if not ok_url:
                QMessageBox.warning(self, "⚠️ មិនអាចទាញ Cloud File", url_error)
                return

        progress = None
        download_temp_dir = None
        temp_git_dir = None
        cloud_backup_path = ""
        try:
            progress = QProgressDialog("កំពុងទាញទិន្នន័យពី Cloud...", "បោះបង់", 0, 100, self)
            style_progress_dialog(progress)
            progress.setWindowTitle("✏️ កែទិន្នន័យ Cloud")
            progress.setWindowModality(Qt.WindowModal)  # type: ignore[attr-defined]
            progress.setMinimumDuration(0)
            progress.setAutoClose(False)
            progress.show()
            QApplication.processEvents()

            temp_root = os.path.join(tempfile.gettempdir(), "ClinicManager")
            os.makedirs(temp_root, exist_ok=True)
            download_temp_dir = tempfile.mkdtemp(prefix="cloud_edit_download_", dir=temp_root)
            downloaded_db = os.path.join(download_temp_dir, upload_file_name)

            progress.setValue(15)
            urllib.request.urlretrieve(cloud_url, downloaded_db)

            progress.setValue(30)
            progress.setLabelText("កំពុងពិនិត្យ database...")
            QApplication.processEvents()

            if not self._validate_sqlite_file(downloaded_db):
                progress.close()
                QMessageBox.critical(
                    self,
                    "❌ File មិនត្រឹមត្រូវ",
                    "File ដែលទាញពី Cloud មិនមែនជា SQLite database ត្រឹមត្រូវទេ។"
                )
                return

            os.makedirs(self.backup_dir, exist_ok=True)
            backup_stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_backup_name = os.path.splitext(upload_file_name)[0]
            cloud_backup_path = os.path.join(self.backup_dir, f"cloud_before_edit_{safe_backup_name}_{backup_stamp}.db")
            shutil.copy2(downloaded_db, cloud_backup_path)

            with sqlite3.connect(downloaded_db) as conn:
                patient_rows = conn.execute("SELECT * FROM patient ORDER BY id ASC").fetchall()

            progress.setValue(45)
            progress.close()

            if not patient_rows:
                QMessageBox.information(
                    self,
                    "📭 គ្មានទិន្នន័យ",
                    "Cloud file នេះមិនមានទិន្នន័យអ្នកជំងឺសម្រាប់កែទេ។"
                )
                return

            review_dialog = CloudUploadReviewDialog(patient_rows, self, f"កែ Cloud - {period_label}")
            if review_dialog.exec_() != QDialog.Accepted:
                return

            edited_rows = review_dialog.get_patient_rows()
            if not edited_rows:
                reply = QMessageBox.question(
                    self,
                    "⚠️ លុបទិន្នន័យទាំងអស់?",
                    "អ្នកបានលុប row ទាំងអស់ចេញ។\n\n"
                    "តើចង់ Upload database ទទេត្រឡប់ទៅ Cloud មែនទេ?",
                    QMessageBox.Yes | QMessageBox.No
                )
                if reply != QMessageBox.Yes:
                    return

            progress = QProgressDialog("កំពុង Upload ទិន្នន័យដែលបានកែ...", "បោះបង់", 0, 100, self)
            style_progress_dialog(progress)
            progress.setWindowTitle("✏️ Upload Cloud ដែលបានកែ")
            progress.setWindowModality(Qt.WindowModal)  # type: ignore[attr-defined]
            progress.setMinimumDuration(0)
            progress.setAutoClose(False)
            progress.show()
            QApplication.processEvents()

            progress.setValue(10)
            progress.setLabelText("កំពុងភ្ជាប់ទៅ Cloud repository...")
            QApplication.processEvents()
            temp_git_dir = self._create_cloud_git_workspace(repo_url, temp_root, "temp_github_cloud_edit_")
            edited_db = os.path.join(temp_git_dir, upload_file_name)

            progress.setValue(20)
            progress.setLabelText("កំពុងបង្កើត database ថ្មី...")
            QApplication.processEvents()
            build_patient_share_database(edited_db, edited_rows)

            metadata = {
                "version": APP_VERSION,
                "upload_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "upload_by": self.current_user or "login-screen",
                "period": period_label,
                "patient_count": len(edited_rows),
                "edited_cloud_file": upload_file_name,
                "cloud_backup": cloud_backup_path,
                "note": "Cloud data edited from uploaded database"
            }
            with open(os.path.join(temp_git_dir, "report_metadata.json"), "w", encoding="utf-8") as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)

            readme_content = (
                f"# Cloud Data Edited: {period_label}\n\n"
                f"**Edited Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                f"**Edited By:** {self.current_user or 'login-screen'}\n"
                f"**File:** {upload_file_name}\n"
                f"**Patient Count:** {len(edited_rows)}\n"
            )
            with open(os.path.join(temp_git_dir, "README.md"), "w", encoding="utf-8") as f:
                f.write(readme_content)

            published, failed_step, git_error = self._publish_cloud_git_workspace(
                temp_git_dir,
                f"Edit cloud data {period_label} - {len(edited_rows)} patients",
                progress,
                35,
                60,
            )
            if not published:
                progress.close()
                QMessageBox.critical(
                    self,
                    "❌ Upload បរាជ័យ",
                    f"កំហុសពេល {failed_step.lower()}\n\n"
                    f"{git_error}\n\n"
                    f"💾 Backup Cloud ដើម:\n{cloud_backup_path}"
                )
                return

            uploaded_db_url = self._github_repo_file_raw_url(repo_url, upload_file_name)
            if uploaded_db_url:
                self.config.set('CATEGORIES', 'cloud_sync_url', uploaded_db_url)
                with open(self.settings_file, 'w', encoding='utf-8') as f:
                    self.config.write(f)

            file_size_mb = f"{os.path.getsize(edited_db) / (1024 * 1024):.2f} MB"
            progress.setValue(100)
            progress.close()

            QMessageBox.information(
                self,
                "✅ កែ Cloud ជោគជ័យ",
                f"ទិន្នន័យ Cloud ត្រូវបានកែ និង Upload ត្រឡប់រួចរាល់។\n\n"
                f"រយៈពេល: {period_label}\n"
                f"ចំនួន row: {len(edited_rows)}\n"
                f"ទំហំ: {file_size_mb}\n\n"
                f"Backup Cloud ដើម:\n{cloud_backup_path}"
            )

        except subprocess.TimeoutExpired:
            if progress:
                progress.close()
            QMessageBox.critical(
                self,
                "❌ Timeout",
                "Upload ចំណាយពេលយូរពេក។ សូមពិនិត្យ Internet Connection។"
            )
        except sqlite3.Error as e:
            if progress:
                progress.close()
            QMessageBox.critical(
                self,
                "❌ Database Error",
                f"មិនអាចអានតារាង patient ពី Cloud database បានទេ:\n{str(e)}"
            )
        except Exception as e:
            if progress:
                progress.close()
            QMessageBox.critical(
                self,
                "❌ កំហុស",
                f"មិនអាចកែទិន្នន័យ Cloud បានទេ:\n{str(e)}"
            )
        finally:
            for folder in (download_temp_dir, temp_git_dir):
                if folder:
                    try:
                        shutil.rmtree(folder, ignore_errors=True)
                    except Exception:
                        pass

    def delete_uploaded_cloud_data(self):
        """Delete an uploaded Cloud database file from the configured GitHub repo."""
        if not self._check_git_installed():
            QMessageBox.critical(
                self,
                "❌ មិនមាន Git",
                "Git មិនទាន់បានដំឡើងលើកុំព្យូទ័រនេះទេ!\n\n"
                "សូមដំឡើង Git មុនពេលលុបទិន្នន័យ Cloud:\n"
                "https://git-scm.com/downloads"
            )
            return

        repo_url = self.config.get('CATEGORIES', 'cloud_sync_repo_url', fallback="").strip()
        if not repo_url:
            reply = QMessageBox.question(
                self,
                "⚙️ មិនមាន GitHub Repository",
                "ត្រូវកំណត់ GitHub Repository URL មុនពេលលុបទិន្នន័យ Cloud។\n\n"
                "តើចង់កំណត់ឥឡូវនេះទេ?",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                self.show_github_setup()
            return

        if "clinic-update" in repo_url.lower():
            QMessageBox.critical(
                self,
                "❌ Repository មិនត្រឹមត្រូវ",
                "Repository នេះសម្រាប់ Update កម្មវិធី មិនមែនសម្រាប់ Cloud Sync ទេ។\n\n"
                "សូមកំណត់ repository ផ្សេងសម្រាប់ទិន្នន័យ Cloud Sync។"
            )
            self.show_github_setup()
            return

        if not self._ensure_cloud_token_for_write():
            return

        delete_file_name, period_label, ok = self._get_cloud_download_period_choice()
        if not ok:
            return

        if not delete_file_name:
            saved_url = self.config.get('CATEGORIES', 'cloud_sync_url', fallback="").strip()
            cloud_url, url_ok = self._get_cloud_sync_url_input(saved_url)
            if not url_ok or not cloud_url.strip():
                return
            parsed_path = urllib.parse.urlparse(cloud_url.strip()).path
            delete_file_name = os.path.basename(parsed_path) or ""
            period_label = "URL ដែលបានបញ្ចូល"

        delete_file_name = str(delete_file_name or "").strip().replace("\\", "/")
        if not delete_file_name or "/" in delete_file_name or not delete_file_name.lower().endswith(".db"):
            QMessageBox.warning(
                self,
                "⚠️ File មិនត្រឹមត្រូវ",
                "មិនអាចកំណត់ file Cloud ដែលត្រូវលុបបានទេ។\n\n"
                "សូមជ្រើសរើសរយៈពេល ឬ URL ដែលជា file .db ត្រឹមត្រូវ។"
            )
            return

        cloud_url = self._github_repo_file_raw_url(repo_url, delete_file_name)
        ok_url, url_error = self._preflight_cloud_sync_url(cloud_url) if cloud_url else (False, "មិនអាចបង្កើត Cloud URL បានទេ។")
        if not ok_url:
            QMessageBox.warning(
                self,
                "⚠️ រក file មិនឃើញ",
                f"រកមិនឃើញ file នេះនៅ Cloud:\n{delete_file_name}\n\n{url_error}"
            )
            return

        first_confirm = QMessageBox.question(
            self,
            "🗑️ បញ្ជាក់ការលុប Cloud",
            f"តើអ្នកចង់លុប file នេះពី Cloud មែនទេ?\n\n"
            f"File: {delete_file_name}\n"
            f"រយៈពេល: {period_label}\n\n"
            "ការលុបនេះនឹងលុបចេញពី GitHub repository។",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        if first_confirm != QMessageBox.Yes:
            return

        second_confirm = QMessageBox.question(
            self,
            "⚠️ បញ្ជាក់ម្ដងទៀត",
            "សូមបញ្ជាក់ម្ដងទៀត៖ តើពិតជាចង់លុប file Cloud នេះមែនទេ?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        if second_confirm != QMessageBox.Yes:
            return

        progress = None
        temp_git_dir = None
        try:
            progress = QProgressDialog("កំពុងលុបទិន្នន័យពី Cloud...", "បោះបង់", 0, 100, self)
            style_progress_dialog(progress)
            progress.setWindowTitle("🗑️ លុបទិន្នន័យ Cloud")
            progress.setWindowModality(Qt.WindowModal)  # type: ignore[attr-defined]
            progress.setMinimumDuration(0)
            progress.setAutoClose(False)
            progress.show()
            QApplication.processEvents()

            temp_root = os.path.join(tempfile.gettempdir(), "ClinicManager")
            os.makedirs(temp_root, exist_ok=True)

            progress.setValue(10)
            progress.setLabelText("កំពុងភ្ជាប់ទៅ Cloud repository...")
            QApplication.processEvents()
            temp_git_dir = self._create_cloud_git_workspace(repo_url, temp_root, "temp_github_cloud_delete_")

            target_file = os.path.abspath(os.path.join(temp_git_dir, delete_file_name))
            repo_root = os.path.abspath(temp_git_dir)
            if os.path.commonpath([repo_root, target_file]) != repo_root:
                raise RuntimeError("Cloud file path is outside the repository.")
            if not os.path.exists(target_file):
                raise FileNotFoundError(delete_file_name)

            progress.setValue(35)
            progress.setLabelText("កំពុងលុប file...")
            QApplication.processEvents()
            os.remove(target_file)

            metadata = {
                "version": APP_VERSION,
                "deleted_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "deleted_by": self.current_user or "login-screen",
                "deleted_file": delete_file_name,
                "period": period_label,
                "note": "Cloud data file deleted from application"
            }
            with open(os.path.join(temp_git_dir, "report_metadata.json"), "w", encoding="utf-8") as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)

            published, failed_step, git_error = self._publish_cloud_git_workspace(
                temp_git_dir,
                f"Delete cloud data {period_label} - {delete_file_name}",
                progress,
                45,
                50,
            )
            if not published:
                progress.close()
                QMessageBox.critical(
                    self,
                    "❌ លុប Cloud បរាជ័យ",
                    f"កំហុសពេល {failed_step.lower()}\n\n{git_error}"
                )
                return

            saved_url = self.config.get('CATEGORIES', 'cloud_sync_url', fallback="").strip()
            if saved_url and saved_url == cloud_url:
                self.config.remove_option('CATEGORIES', 'cloud_sync_url')
                self._save_settings()

            progress.setValue(100)
            progress.close()
            QMessageBox.information(
                self,
                "✅ លុប Cloud ជោគជ័យ",
                f"បានលុប file ពី Cloud រួចរាល់។\n\n"
                f"File: {delete_file_name}\n"
                f"រយៈពេល: {period_label}"
            )

        except subprocess.TimeoutExpired:
            if progress:
                progress.close()
            QMessageBox.critical(
                self,
                "❌ Timeout",
                "ការលុប Cloud ចំណាយពេលយូរពេក។ សូមពិនិត្យ Internet Connection។"
            )
        except Exception as e:
            if progress:
                progress.close()
            QMessageBox.critical(
                self,
                "❌ កំហុស",
                f"មិនអាចលុបទិន្នន័យ Cloud បានទេ:\n{self._redact_cloud_token(str(e))}"
            )
        finally:
            if temp_git_dir:
                try:
                    shutil.rmtree(temp_git_dir, ignore_errors=True)
                except Exception:
                    pass

    def show_telegram_bot_setup(self):
        """បង្ហាញ Dialog សម្រាប់កំណត់ Telegram Bot"""
        dialog = TelegramBotSetupDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            bot_token = dialog.txt_bot.text().strip()
            chat_id = dialog.txt_chat.text().strip()

            if bot_token and chat_id:
                self.config.set('CATEGORIES', 'telegram_bot_token', bot_token)
                self.config.set('CATEGORIES', 'telegram_chat_id', chat_id)
                with open(self.settings_file, 'w', encoding='utf-8') as f:
                    self.config.write(f)
                db.logger.info(f"Telegram Bot configured: {bot_token[:10]}... / {chat_id}")
                QMessageBox.information(
                    self,
                    "✅ ជោគជ័យ",
                    "ការកំណត់ Telegram Bot បានជោគជ័យ!\n\n"
                    "អ្នកអាចចុច '⬆️ ផ្ញើទៅ Cloud' ឥឡូវនេះ។"
                )
            else:
                QMessageBox.warning(self, "⚠️ មិនពេញលេញ", "សូមបញ្ចូលទាំង Bot Token និង Chat ID!")

    def _ensure_settings_categories(self):
        if 'CATEGORIES' not in self.config:
            self.config['CATEGORIES'] = {}

    def _save_settings(self):
        self._ensure_settings_categories()
        with open(self.settings_file, 'w', encoding='utf-8') as f:
            self.config.write(f)

    def _cloud_token_username(self):
        username = (self.current_user or "").strip() or "login_screen"
        return re.sub(r"[^A-Za-z0-9_-]", "_", username)

    def _cloud_token_key(self):
        return f"cloud_github_token_{self._cloud_token_username()}"

    def _get_cloud_github_token(self):
        self._ensure_settings_categories()
        token = self.config.get('CATEGORIES', self._cloud_token_key(), fallback="").strip()
        if token:
            return token

        legacy_token = self.config.get('CATEGORIES', 'cloud_github_token', fallback="").strip()
        if legacy_token and (self.current_user or "").strip():
            self.config.set('CATEGORIES', self._cloud_token_key(), legacy_token)
            self.config.remove_option('CATEGORIES', 'cloud_github_token')
            self._save_settings()
            return legacy_token
        return ""

    def _set_cloud_github_token(self, token):
        self._ensure_settings_categories()
        self.config.set('CATEGORIES', self._cloud_token_key(), token)
        if self.config.has_option('CATEGORIES', 'cloud_github_token'):
            self.config.remove_option('CATEGORIES', 'cloud_github_token')
        self._save_settings()

    def _ensure_cloud_token_for_write(self):
        token = self._get_cloud_github_token()
        if token:
            return True

        reply = QMessageBox.question(
            self,
            "🔑 ត្រូវការ GitHub Token",
            f"User នេះ ({self.current_user or 'login-screen'}) មិនទាន់មាន GitHub Token ទេ។\n\n"
            "ការកែ ឬលុបទិន្នន័យ Cloud ត្រូវការ GitHub Token ដើម្បីសរសេរទៅ repository។\n\n"
            "តើអ្នកចង់កំណត់ Token ឥឡូវនេះទេ?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )
        if reply != QMessageBox.Yes:
            return False

        self.show_github_token_setup()
        return bool(self._get_cloud_github_token())

    def show_github_token_setup(self):
        """Save a GitHub token for Cloud Sync edit/delete on this computer."""
        self._ensure_settings_categories()
        current_token = self._get_cloud_github_token()
        username = self.current_user or "login-screen"
        dialog = QInputDialog(self)
        dialog.setInputMode(QInputDialog.TextInput)
        dialog.setTextEchoMode(QLineEdit.Password)
        dialog.setWindowTitle("🔑 កំណត់ GitHub Token")
        dialog.setLabelText(
            f"បញ្ចូល GitHub Personal Access Token សម្រាប់ User: {username}\n\n"
            "• Token ត្រូវមានសិទ្ធិ write ទៅ repository Cloud Sync\n"
            "• User ផ្សេងត្រូវបញ្ចូល token ដោយឡែក\n"
            "• ទុកឲ្យទទេ រួចចុច OK ដើម្បីលុប token ចាស់\n"
            "• Token នឹងរក្សាទុកតែលើ PC នេះប៉ុណ្ណោះ"
        )
        dialog.setTextValue(current_token)
        dialog.setOkButtonText("រក្សាទុក")
        dialog.setCancelButtonText("បោះបង់")
        dialog.resize(620, 220)
        self._style_url_input_dialog(dialog)

        if dialog.exec_() != QDialog.Accepted:
            return

        token = dialog.textValue().strip()
        self._set_cloud_github_token(token)
        if token:
            QMessageBox.information(
                self,
                "✅ បានរក្សាទុក",
                f"GitHub Token ត្រូវបានរក្សាទុកសម្រាប់ User: {username} លើ PC នេះ។"
            )
        else:
            QMessageBox.information(self, "✅ បានលុប", f"GitHub Token របស់ User: {username} ត្រូវបានលុបចេញរួចហើយ។")

    def _cloud_git_env(self):
        env = os.environ.copy()
        env["GIT_TERMINAL_PROMPT"] = "0"
        return env

    def _cloud_repo_url_for_git(self, repo_url):
        token = self._get_cloud_github_token()
        repo_url = str(repo_url or "").strip()
        if not token:
            return repo_url

        quoted_token = urllib.parse.quote(token, safe="")
        if repo_url.startswith("https://github.com/"):
            parsed = urllib.parse.urlparse(repo_url)
            return urllib.parse.urlunparse((
                parsed.scheme,
                f"x-access-token:{quoted_token}@{parsed.netloc}",
                parsed.path,
                parsed.params,
                parsed.query,
                parsed.fragment,
            ))

        if repo_url.startswith("git@github.com:"):
            repo_path = repo_url[len("git@github.com:"):].strip("/")
            return f"https://x-access-token:{quoted_token}@github.com/{repo_path}"

        return repo_url

    def _redact_cloud_token(self, text):
        text = str(text or "")
        token = self._get_cloud_github_token()
        if token:
            text = text.replace(token, "***")
            text = text.replace(urllib.parse.quote(token, safe=""), "***")
        return re.sub(r"https://x-access-token:[^@\s]+@github\.com/", "https://x-access-token:***@github.com/", text)

    def _check_git_installed(self):
        """ពិនិត្យថា Git ដំឡើងហើយឬនៅ"""
        try:
            result = subprocess.run(
                ['git', '--version'],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                timeout=5,
                **_subprocess_no_window_kwargs()
            )
            return result.returncode == 0
        except Exception:
            return False

    def _create_cloud_git_workspace(self, repo_url, temp_root, prefix):
        """Clone the Cloud repo when possible so older uploaded files are preserved."""
        os.makedirs(temp_root, exist_ok=True)
        temp_git_dir = tempfile.mkdtemp(prefix=prefix, dir=temp_root)
        git_repo_url = self._cloud_repo_url_for_git(repo_url)
        clone_result = subprocess.run(
            ['git', 'clone', '--depth', '1', git_repo_url, temp_git_dir],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',
            timeout=120,
            env=self._cloud_git_env(),
            **_subprocess_no_window_kwargs()
        )

        if clone_result.returncode == 0:
            return temp_git_dir

        shutil.rmtree(temp_git_dir, ignore_errors=True)
        temp_git_dir = tempfile.mkdtemp(prefix=prefix, dir=temp_root)
        init_commands = [
            ['git', 'init'],
            ['git', 'branch', '-M', 'main'],
            ['git', 'remote', 'add', 'origin', git_repo_url],
        ]
        for cmd in init_commands:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                cwd=temp_git_dir,
                timeout=120,
                env=self._cloud_git_env(),
                **_subprocess_no_window_kwargs()
            )
            if result.returncode != 0:
                output = self._redact_cloud_token(result.stderr.strip() or result.stdout.strip())
                raise RuntimeError(output or "Git workspace setup failed.")
        return temp_git_dir

    def _publish_cloud_git_workspace(self, temp_git_dir, commit_message, progress, start_value=60, span=40):
        commands = [
            ("Setting git user name...", ['git', 'config', 'user.name', 'NOU SARAT']),
            ("Setting git user email...", ['git', 'config', 'user.email', 'saratboy1988-a11y@users.noreply.github.com']),
            ("Adding files...", ['git', 'add', '.']),
            ("Committing...", ['git', 'commit', '-m', commit_message]),
            ("Setting branch...", ['git', 'branch', '-M', 'main']),
            ("Pushing to GitHub...", ['git', 'push', '-u', 'origin', 'main']),
        ]

        for i, (msg, cmd) in enumerate(commands):
            if progress:
                progress.setValue(start_value + int((i / len(commands)) * span))
                progress.setLabelText(msg)
                QApplication.processEvents()

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                cwd=temp_git_dir,
                timeout=120,
                env=self._cloud_git_env(),
                **_subprocess_no_window_kwargs()
            )
            if result.returncode != 0:
                output = self._redact_cloud_token(f"{result.stderr}\n{result.stdout}".strip())
                if "nothing to commit" in output.lower():
                    continue
                return False, msg, output
        return True, "", ""

    def _upload_to_github(self, repo_url):
        """ផ្ញើរបាយការណ៍ទៅ GitHub (ជម្រើសច្រើន: ថ្ងៃ/សប្តាហ៍/ខែ/កំណត់ដោយខ្លួនឯង)"""
        try:
            # កំណត់ខែបច្ចុប្បន្ន
            now = datetime.now()
            today_str = now.strftime("%d/%m/%Y")

            # សួរអ្នកប្រើប្រាស់ថាតើចង់ upload បែបណា
            period_choice, ok = self._get_upload_period_choice(today_str)

            if not ok or not period_choice:
                return

            # កំណត់ថ្ងៃចាប់ផ្តើម + បញ្ចប់
            start_date = None
            end_date = None
            period_label = ""

            if "ថ្ងៃនេះ" in period_choice or "Today" in period_choice:
                start_date = now.strftime("%d/%m/%Y")
                end_date = start_date
                period_label = f"ថ្ងៃនេះ ({today_str})"

            elif "សប្តាហ៍នេះ" in period_choice or "This Week" in period_choice:
                # ថ្ងៃច័ន្ទ ដល់ ថ្ងៃអាទិត្យ
                monday = now - timedelta(days=now.weekday())
                sunday = monday + timedelta(days=6)
                start_date = monday.strftime("%d/%m/%Y")
                end_date = sunday.strftime("%d/%m/%Y")
                period_label = f"សប្តាហ៍នេះ ({start_date} - {end_date})"

            elif "ខែនេះ" in period_choice or "This Month" in period_choice:
                first_day = now.replace(day=1)
                last_day = now.replace(day=1, month=now.month + 1) - timedelta(days=1) if now.month < 12 else now.replace(year=now.year + 1, month=1, day=1) - timedelta(days=1)
                start_date = first_day.strftime("%d/%m/%Y")
                end_date = last_day.strftime("%d/%m/%Y")
                period_label = f"ខែនេះ ({start_date} - {end_date})"

            elif "ខែមុន" in period_choice or "Last Month" in period_choice:
                first_day_last_month = now.replace(day=1, month=now.month - 1) if now.month > 1 else now.replace(year=now.year - 1, month=12, day=1)
                last_day_last_month = now.replace(day=1, month=now.month) - timedelta(days=1)
                start_date = first_day_last_month.strftime("%d/%m/%Y")
                end_date = last_day_last_month.strftime("%d/%m/%Y")
                period_label = f"ខែមុន ({start_date} - {end_date})"

            elif "កំណត់ដោយខ្លួនឯង" in period_choice or "Custom" in period_choice:
                # បង្ហាញ Date Range Dialog
                dialog = CustomDateRangeDialog(self)
                if dialog.exec_() == QDialog.Accepted:
                    start_date = dialog.start_date.date().toString("dd/MM/yyyy")
                    end_date = dialog.end_date.date().toString("dd/MM/yyyy")
                    period_label = f"កំណត់ដោយខ្លួនឯង ({start_date} - {end_date})"
                else:
                    return

            elif "ទាំងអស់" in period_choice or "Full" in period_choice:
                start_date = None
                end_date = None
                period_label = "ទាំងអស់ (Full Database)"

            # បង្កើត Backup មុនពេល Upload
            backup_path = create_database_backup(self.backup_dir)

            # ទាញយកអ្នកជំងឺតាមរយៈពេល
            if start_date and end_date:
                # បំប្លែង date format សម្រាប់ SQL
                start_sql = start_date  # dd/mm/yyyy
                end_sql = end_date
                branch_clause = ""
                branch_params = ()
                if self.active_branch_code:
                    branch_clause = " AND (branch_code = ? OR branch_code IS NULL OR branch_code = '')"
                    branch_params = (self.active_branch_code,)

                patients_selected = db.execute_read(
                    f"""
                    SELECT * FROM patient
                    WHERE (SUBSTR(date, 7, 4) || SUBSTR(date, 4, 2) || SUBSTR(date, 1, 2))
                          BETWEEN (SUBSTR(?, 7, 4) || SUBSTR(?, 4, 2) || SUBSTR(?, 1, 2))
                              AND (SUBSTR(?, 7, 4) || SUBSTR(?, 4, 2) || SUBSTR(?, 1, 2))
                    {branch_clause}
                    """,
                    (start_sql, start_sql, start_sql, end_sql, end_sql, end_sql, *branch_params)
                ) or []
            else:
                # Full database
                patients_selected = db.view(self.active_branch_code) or []

            if not patients_selected:
                QMessageBox.information(
                    self,
                    "📭 គ្មានទិន្នន័យ",
                    f"មិនមានទិន្នន័យសម្រាប់ {period_label} ទេ។"
                )
                return

            review_dialog = CloudUploadReviewDialog(patients_selected, self, period_label)
            if review_dialog.exec_() != QDialog.Accepted:
                return
            patients_selected = review_dialog.get_patient_rows()
            if not patients_selected:
                QMessageBox.information(
                    self,
                    "📭 គ្មានទិន្នន័យ",
                    "អ្នកបានលុប row ទាំងអស់ចេញពី package Upload។"
                )
                return

            # សួរថាតើចង់បន្ថែមការកត់សម្គាល់ទេ
            note, note_ok = self._get_upload_note_input(period_label)

            # បង្កើត Progress Dialog
            progress = QProgressDialog(
                f"កំពុងរៀបចំរបាយការណ៍...",
                "បោះបង់", 0, 100, self
            )
            style_progress_dialog(progress)
            progress.setWindowTitle(f"⬆️ Upload: {period_label}")
            progress.setWindowModality(Qt.WindowModal)  # type: ignore[attr-defined]
            progress.setMinimumDuration(0)
            progress.setAutoClose(False)
            progress.show()

            temp_root = os.path.join(tempfile.gettempdir(), "ClinicManager")

            try:
                os.makedirs(temp_root, exist_ok=True)
                temp_git_dir = self._create_cloud_git_workspace(repo_url, temp_root, "temp_github_upload_")
            except Exception as e:
                progress.close()
                QMessageBox.critical(
                    self,
                    "❌ កំហុស",
                    f"មិនអាចបង្កើត folder បានទេ: {str(e)}\n\n"
                    f"សូមបិទកម្មវិធី ហើយបើកម្តងទៀត។"
                )
                return

            progress.setValue(20)
            progress.setLabelText("កំពុងបង្កើត database...")
            QApplication.processEvents()

            # កំណត់ filename
            if start_date and end_date:
                # ជៀសវាង characters ដែលបញ្ហាក្នុង filename
                safe_start = start_date.replace('/', '-')
                safe_end = end_date.replace('/', '-')
                filename_db = f"clinic_{safe_start}_to_{safe_end}.db"
            else:
                filename_db = "clinic_full.db"

            temp_db = os.path.join(temp_git_dir, filename_db)

            progress.setValue(25)
            progress.setLabelText("កំពុងបង្កើត database...")
            QApplication.processEvents()

            # បង្កើត database ថ្មី
            try:
                build_patient_share_database(temp_db, patients_selected)

                # ពិនិត្យថា file ត្រូវបានបង្កើត
                if not os.path.exists(temp_db):
                    raise Exception(f"Database file មិនត្រូវបានបង្កើត: {temp_db}")

            except Exception as e:
                progress.close()
                QMessageBox.critical(
                    self,
                    "❌ កំហុស Database",
                    f"មិនអាចបង្កើត database បានទេ: {str(e)}\n\n"
                    f"សូមពិនិត្យ:\n"
                    f"1. មានទិន្នន័យអ្នកជំងឺក្នុងរយៈពេលនេះ\n"
                    f"2. មានសិទ្ធិសរសេរក្នុង folder\n\n"
                    f"💾 Backup ក្នុង:\n{backup_path}"
                )
                return

            progress.setValue(40)
            progress.setLabelText("កំពុងបង្កើតឯកសារ...")
            QApplication.processEvents()

            # បង្កើត metadata file សម្រាប់ Cloud Sync ដោយមិនប៉ះ updater manifest
            version_data = {
                "version": APP_VERSION,
                "upload_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "upload_by": self.current_user,
                "period": period_label,
                "start_date": start_date,
                "end_date": end_date,
                "patient_count": len(patients_selected),
                "note": note if note_ok and note else ""
            }
            with open(os.path.join(temp_git_dir, "report_metadata.json"), 'w', encoding='utf-8') as f:
                json.dump(version_data, f, indent=2, ensure_ascii=False)

            progress.setValue(50)
            QApplication.processEvents()

            # បង្កើត README.md
            note_section = f"\n**ការកត់សម្គាល់:** {note}\n" if note else ""

            readme_content = (
                f"# របាយការណ៍: {period_label}\n\n"
                f"**Upload Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                f"**Uploaded By:** {self.current_user}\n"
                f"**Version:** {APP_VERSION}\n"
                f"**ចំនួនអ្នកជំងឺ:** {len(patients_selected)} នាក់\n"
                f"**រយៈពេល:** {start_date or 'ទាំងអស់'} - {end_date or ''}\n"
                f"{note_section}\n"
                f"## សង្ខេប\n\n"
                f"- ចំនួនអ្នកជំងឺសរុប: {len(patients_selected)} នាក់\n"
            )
            with open(os.path.join(temp_git_dir, "README.md"), 'w', encoding='utf-8') as f:
                f.write(readme_content)

            progress.setValue(60)
            QApplication.processEvents()

            published, failed_step, git_error = self._publish_cloud_git_workspace(
                temp_git_dir,
                f"Report {period_label} - {len(patients_selected)} patients",
                progress,
                60,
                40,
            )
            if not published:
                progress.close()
                QMessageBox.critical(
                    self,
                    "❌ Upload បរាជ័យ",
                    f"កំហុសពេល {failed_step.lower()}\n\n"
                    f"{git_error}\n\n"
                    f"💾 Backup ក្នុង:\n{backup_path}"
                )
                return

            uploaded_db_url = self._github_repo_file_raw_url(repo_url, filename_db)
            if uploaded_db_url:
                try:
                    self.config.set('CATEGORIES', 'cloud_sync_url', uploaded_db_url)
                    with open(self.settings_file, 'w', encoding='utf-8') as f:
                        self.config.write(f)
                    db.logger.info(f"Cloud Sync download URL saved: {uploaded_db_url}")
                except Exception as e:
                    db.logger.error(f"Failed to save Cloud Sync download URL: {e}")

            # គណនា file size មុនពេលសម្អាត temp folder
            file_size_mb = "N/A"
            if os.path.exists(temp_db):
                file_size_mb = f"{os.path.getsize(temp_db) / (1024*1024):.2f} MB"
            else:
                # បើ file មិនមាន បង្ហាញ error
                progress.close()
                QMessageBox.critical(
                    self,
                    "❌ កំហុស",
                    f"មិនអាចបង្កើត database file បានទេ!\n\n"
                    f"សូមពិនិត្យ:\n"
                    f"1. មានទិន្នន័យអ្នកជំងឺក្នុងរយៈពេលនេះ\n"
                    f"2. មានសិទ្ធិសរសេរក្នុង folder\n\n"
                    f"💾 Backup ក្នុង:\n{backup_path}"
                )
                return

            # សម្អាត temp folder (ignore errors ព្រោះ Windows អាច lock files)
            import shutil
            try:
                shutil.rmtree(temp_git_dir, ignore_errors=True)
            except Exception:
                pass

            progress.setValue(100)
            progress.setLabelText("✅ ផ្ញើជោគជ័យ!")
            progress.close()

            # រៀបចំ note section (ជៀសវាង backslash ក្នុង f-string)
            note_html = ""
            if note:
                note_html = f"<span style='color: #00cec9;'>📝 ការកត់សម្គាល់:</span> <span style='color: white;'>{note}</span><br><br>"

            # បង្កើត custom message box ជាមួយពណ៌អក្សរស (white text on dark background)
            success_msg = QMessageBox(self)
            success_msg.setIcon(QMessageBox.Information)
            success_msg.setWindowTitle("✅ Upload ជោគជ័យ!")
            success_msg.setText(
                f"<span style='color: white;'>របាយការណ៍ត្រូវបានផ្ញើទៅ GitHub រួចរាល់!</span><br><br>"
                f"<span style='color: #00cec9;'>📅 រយៈពេល:</span> <span style='color: white;'>{period_label}</span><br>"
                f"<span style='color: #00cec9;'>👥 ចំនួនអ្នកជំងឺ:</span> <span style='color: white;'>{len(patients_selected)} នាក់</span><br>"
                f"<span style='color: #00cec9;'>📦 ទំហំ:</span> <span style='color: white;'>{file_size_mb}</span><br><br>"
                f"{note_html}"
                f"<span style='color: #00cec9;'>🔗 Cloud URL:</span><br>"
                f"<span style='color: #d2dae2;'>{uploaded_db_url or 'N/A'}</span><br><br>"
                f"<span style='color: #ffa801;'>💾 Backup ក្នុង:</span><br>"
                f"<span style='color: #d2dae2;'>{backup_path}</span><br><br>"
                f"<span style='color: #05c46b;'>អ្នកអាច Merge ពី GitHub លើ PC ផ្សេង!</span>"
            )
            success_msg.setStyleSheet(build_message_box_stylesheet(
                button_bg="#05c46b",
                button_color="white",
                button_hover="#06bc5c",
                label_min_width=300,
                button_min_width=80,
            ))
            success_msg.addButton("យល់ហើយ 👍", QMessageBox.AcceptRole)
            success_msg.exec_()

        except subprocess.TimeoutExpired:
            QMessageBox.critical(
                self,
                "❌ Timeout",
                "Upload ចំណាយពេលយូរពេក (២ នាទី)!\n\n"
                "សូមពិនិត្យ Internet Connection។"
            )
        except Exception as e:
            QMessageBox.critical(
                self,
                "❌ កំហុស",
                f"មិនអាច Upload ទិន្នន័យបានទេ: {str(e)}"
            )

    def _github_repo_file_raw_url(self, repo_url, file_name):
        """Build raw.githubusercontent.com URL for a file in a GitHub repo."""
        repo_url = str(repo_url or "").strip()
        file_name = str(file_name or "").strip().replace("\\", "/")
        if not repo_url or not file_name:
            return ""

        if repo_url.startswith("https://github.com/"):
            repo_path = repo_url[len("https://github.com/"):]
            if repo_path.endswith(".git"):
                repo_path = repo_path[:-4]
            repo_path = repo_path.strip("/")
            return f"https://raw.githubusercontent.com/{repo_path}/main/{file_name}"

        if repo_url.startswith("git@github.com:"):
            repo_path = repo_url[len("git@github.com:"):]
            if repo_path.endswith(".git"):
                repo_path = repo_path[:-4]
            repo_path = repo_path.strip("/")
            return f"https://raw.githubusercontent.com/{repo_path}/main/{file_name}"

        return ""

    def show_github_setup(self):
        """បង្ហាញ Dialog សម្រាប់កំណត់ GitHub Repository របស់ Cloud Sync"""
        repo_url, ok = QInputDialog.getText(
            self,
            "⚙️ កំណត់ Cloud Sync Repository",
            "សូមបញ្ចូល GitHub Repository URL សម្រាប់ Cloud Sync:\n\n"
            "ឧទាហរណ៍:\n"
            "https://github.com/saratboy1988-a11y/Clinic-Cloud-Sync.git\n\n"
            "ចំណាំ៖\n"
            "• បង្កើត Repository មុន (Public ឬ Private)\n"
            "• កុំប្រើ Clinic-Update repository ព្រោះវាសម្រាប់ Update កម្មវិធី\n"
            "• កំណត់ GitHub Token ដាច់ដោយឡែកសម្រាប់ Upload/Edit",
            text=self.config.get('CATEGORIES', 'cloud_sync_repo_url', fallback="https://github.com/saratboy1988-a11y/Clinic-Cloud-Sync.git")
        )

        if ok and repo_url and repo_url.strip():
            repo_url = repo_url.strip()

            # ពិនិត្យ URL បឋម
            if not repo_url.startswith(('https://github.com/', 'git@github.com:')):
                QMessageBox.warning(
                    self,
                    "⚠️ URL មិនត្រឹមត្រូវ",
                    "Repository URL ត្រូវចាប់ផ្តើមដោយ:\n"
                    "https://github.com/USERNAME/REPO.git"
                )
                return
            if "clinic-update" in repo_url.lower():
                QMessageBox.warning(
                    self,
                    "⚠️ Repository មិនត្រឹមត្រូវ",
                    "Clinic-Update repository ត្រូវរក្សាទុកសម្រាប់ Update កម្មវិធី។\n\n"
                    "សូមបញ្ចូល repository ផ្សេងសម្រាប់ Cloud Sync។"
                )
                return

            # រក្សាទុក
            self._ensure_settings_categories()
            self.config.set('CATEGORIES', 'cloud_sync_repo_url', repo_url)
            self._save_settings()
            db.logger.info(f"Cloud Sync GitHub Repo URL saved: {repo_url}")

            QMessageBox.information(
                self,
                "✅ ជោគជ័យ",
                f"GitHub Repository ត្រូវបានកំណត់!\n\n"
                f"📦 URL: {repo_url}\n\n"
                f"ឥឡូវអ្នកអាចចុច '⬆️ ផ្ញើទៅ Cloud' បាន។"
            )
        else:
            QMessageBox.warning(self, "⚠️ បោះបង់", "សូមបញ្ចូល GitHub Repository URL!")

    def _send_to_telegram_bot(self, bot_token, chat_id):
        """ផ្ញើ database ទៅ Telegram Bot"""
        try:
            import json

            # បង្កើត Backup មុនពេលផ្ញើ
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            backup_path = create_database_backup(self.backup_dir)

            # បង្កើត Progress Dialog
            progress = QProgressDialog("កំពុងផ្ញើទិន្នន័យទៅ Telegram...", "បោះបង់", 0, 100, self)
            style_progress_dialog(progress)
            progress.setWindowTitle("⬆️ Upload to Cloud")
            progress.setWindowModality(Qt.WindowModal)  # type: ignore[attr-defined]
            progress.setMinimumDuration(0)
            progress.setAutoClose(False)
            progress.show()

            # ប្រើ urllib ជំនួស requests (មិនចាំបាច់ install បន្ថែម)
            boundary = "----WebKitFormBoundary7MA4YWxkTrZu0gW"
            filename = f"ClinicDB_{timestamp}.db"

            # អាន file
            with open(db.DB_NAME, 'rb') as f:
                file_data = f.read()

            file_size = len(file_data)

            # បង្កើត multipart/form-data body
            body = (
                f"--{boundary}\r\n"
                f"Content-Disposition: form-data; name=\"chat_id\"\r\n\r\n"
                f"{chat_id}\r\n"
                f"--{boundary}\r\n"
                f"Content-Disposition: form-data; name=\"document\"; filename=\"{filename}\"\r\n"
                f"Content-Type: application/octet-stream\r\n\r\n"
            ).encode('utf-8')

            body += file_data
            body += f"\r\n--{boundary}--\r\n".encode('utf-8')

            # បង្កើត request
            url = f"https://api.telegram.org/bot{bot_token}/sendDocument"
            req = urllib.request.Request(url, data=body, method='POST')
            req.add_header('Content-Type', f'multipart/form-data; boundary={boundary}')

            # ផ្ញើ request ជាមួយ progress
            import io
            import socket

            old_timeout = socket.getdefaulttimeout()
            socket.setdefaulttimeout(60)  # ៦០ វិនាទី timeout

            try:
                # ប្រើ chunked upload សម្រាប់ progress
                response = urllib.request.urlopen(req)
                response_data = response.read().decode('utf-8')

                progress.setValue(100)
                progress.setLabelText("✅ ផ្ញើជោគជ័យ!")

                # ពិនិត្យ response
                result = json.loads(response_data)
                if result.get('ok'):
                    QMessageBox.information(
                        self,
                        "✅ ផ្ញើជោគជ័យ!",
                        f"ទិន្នន័យត្រូវបានផ្ញើទៅ Telegram រួចរាល់!\n\n"
                        f"📦 ឈ្មោះ: {filename}\n"
                        f"📊 ទំហំ: {file_size / (1024*1024):.2f} MB\n\n"
                        f"💾 Backup ត្រូវបានរក្សាទុកក្នុង:\n{backup_path}\n\n"
                        f"អ្នកអាចទាញយក file នេះពី Telegram លើ PC ផ្សេង!"
                    )
                else:
                    error_msg = result.get('description', 'Unknown error')
                    QMessageBox.critical(
                        self,
                        "❌ ផ្ញើបរាជ័យ",
                        f"Telegram មិនទទួលយក file ទេ: {error_msg}\n\n"
                        f"សូមពិនិត្យ Bot Token និង Chat ID។"
                    )
            except urllib.error.HTTPError as e:
                error_body = e.read().decode('utf-8') if e.fp else str(e)
                QMessageBox.critical(
                    self,
                    "❌ HTTP Error",
                    f"កំហុស HTTP {e.code}: {error_body}\n\n"
                    f"សូមពិនិត្យ Bot Token និង Chat ID។"
                )
            except urllib.error.URLError as e:
                QMessageBox.critical(
                    self,
                    "❌ Connection Error",
                    f"មិនអាចតភ្ជាប់ទៅ Telegram បានទេ: {str(e.reason)}\n\n"
                    f"សូមពិនិត្យ Internet Connection។"
                )
            finally:
                socket.setdefaulttimeout(old_timeout)
                progress.close()

        except Exception as e:
            QMessageBox.critical(
                self,
                "❌ កំហុស",
                f"មិនអាចផ្ញើទិន្នន័យបានទេ: {str(e)}"
            )

    def sync_data_initial(self):
        """
        ទាញយកទិន្នន័យពី Cloud (Cloud Sync)
        កែលម្អ៖ មាន Backup, Validation, Progress Bar, Timeout, Remember URL
        """
        # ១. ទាញយក URL ដែលបានរក្សាទុក (បើមាន)
        saved_url = self.config.get('CATEGORIES', 'cloud_sync_url', fallback="")

        # Default URL ពី GitHub
        default_url = "https://raw.githubusercontent.com/saratboy1988-a11y/Clinic-Cloud-Sync/main/clinic_01-04-2026_to_30-04-2026.db"

        # បើគ្មាន saved URL ទេ ប្រើ default
        initial_url = saved_url if saved_url else default_url

        # ២. ជ្រើសរើសរយៈពេលទាញយកមុន ដើម្បីបង្កើត URL ដោយស្វ័យប្រវត្តិ
        download_file_name, download_period_label, ok = self._get_cloud_download_period_choice()
        if not ok:
            return

        url = ""
        if download_file_name:
            repo_url = self.config.get('CATEGORIES', 'cloud_sync_repo_url', fallback="").strip()
            url = self._github_repo_file_raw_url(repo_url, download_file_name)
            if not url:
                url = self._replace_cloud_sync_url_filename(initial_url, download_file_name)
        else:
            # ប្រើ URL ដែលបានបញ្ចូលដោយដៃ
            url, ok = self._get_cloud_sync_url_input(initial_url)
            if not ok or not url or not url.strip():
                return
            url = url.strip()

        if not url:
            QMessageBox.warning(
                self,
                "⚠️ មិនអាចបង្កើត URL",
                "មិនអាចបង្កើត URL សម្រាប់ទាញយកបានទេ។\n\n"
                "សូមកំណត់ Cloud Sync Repository ឬជ្រើស 'ប្រើ URL ដែលបានបញ្ចូល'។"
            )
            return

        # ២. ផ្ទៀងផ្ទាត់ URL បឋម
        if not url.startswith(('http://', 'https://')):
            QMessageBox.warning(
                self,
                "⚠️ URL មិនត្រឹមត្រូវ",
                "URL ត្រូវចាប់ផ្តើមដោយ http:// ឬ https://"
            )
            return

        # ៣.១ ពិនិត្យ URL មុនពេលទាញយក
        is_valid_url, validation_message = self._preflight_cloud_sync_url(url)
        if not is_valid_url:
            fallback_url = ""
            if download_file_name and saved_url and saved_url.strip() and saved_url.strip() != url:
                reply = QMessageBox.question(
                    self,
                    "⚠️ រក file មិនឃើញ",
                    f"រកមិនឃើញ file នេះនៅ Cloud:\n{download_file_name}\n\n"
                    "តើចង់ប្រើ URL ចុងក្រោយដែលបានរក្សាទុកវិញទេ?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )
                if reply == QMessageBox.Yes:
                    fallback_url = saved_url.strip()
                    download_period_label = "Cloud file ចុងក្រោយ"

            if not fallback_url:
                reply = QMessageBox.question(
                    self,
                    "⚠️ URL មិនអាចប្រើបាន",
                    f"{validation_message}\n\n"
                    "តើចង់បញ្ចូល URL ដោយផ្ទាល់វិញទេ?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )
                if reply == QMessageBox.Yes:
                    fallback_url, ok = self._get_cloud_sync_url_input(saved_url or initial_url)
                    if not ok or not fallback_url or not fallback_url.strip():
                        return
                    fallback_url = fallback_url.strip()
                    download_period_label = "URL ដែលបានបញ្ចូល"

            if not fallback_url:
                return

            fallback_valid, fallback_message = self._preflight_cloud_sync_url(fallback_url)
            if not fallback_valid:
                QMessageBox.warning(self, "❌ URL មិនអាចប្រើបាន", fallback_message)
                return
            url = fallback_url

        # ៣.២ រក្សាទុក URL ក្រោយពេលពិនិត្យថាប្រើបាន
        try:
            self.config.set('CATEGORIES', 'cloud_sync_url', url)
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                self.config.write(f)
            db.logger.info(f"Cloud Sync URL saved: {url}")
        except Exception as e:
            db.logger.error(f"Failed to save Cloud Sync URL: {e}")

        # ៤. បង្កើត Backup មុនពេលធ្វើអ្វីផ្សេង
        try:
            import shutil
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_dir = self.backup_dir or os.path.join(tempfile.gettempdir(), "ClinicManager", "backups")
            os.makedirs(backup_dir, exist_ok=True)
            backup_path = os.path.join(backup_dir, f"clinic_backup_before_sync_{timestamp}.db")

            if os.path.exists(db.DB_NAME):
                shutil.copy2(db.DB_NAME, backup_path)
                db.logger.info(f"Backup created before sync: {backup_path}")
        except Exception as e:
            QMessageBox.critical(
                self,
                "❌ កំហុស Backup",
                f"មិនអាចបង្កើត Backup បានទេ: {str(e)}\n\n"
                f"សូមពិនិត្យមើលថា folder អាចសរសេរបាន។"
            )
            return

        # ៥. ទាញយក file ជាមួយ Progress Dialog
        progress = None
        try:
            # បង្កើត Progress Dialog
            progress = QProgressDialog("កំពុងទាញយកទិន្នន័យពី Cloud...", "បោះបង់", 0, 100, self)
            style_progress_dialog(progress)
            progress.setWindowTitle("☁️ Cloud Sync")
            progress.setWindowModality(Qt.WindowModal)  # type: ignore[attr-defined]
            progress.setMinimumDuration(0)
            progress.setAutoClose(False)
            progress.show()

            # ប្រើ Temp file ដើម្បីសុវត្ថិភាព
            import tempfile
            fd, temp_db = tempfile.mkstemp(suffix=".db", prefix="cloud_sync_")
            os.close(fd)

            # Download file ជាមួយ timeout (៣០ វិនាទី)
            def download_progress(block_num, block_size, total_size):
                """Update progress bar"""
                if total_size > 0:
                    downloaded = block_num * block_size
                    percent = min(int((downloaded / total_size) * 100), 100)
                    progress.setValue(percent)
                    progress.setLabelText(
                        f"កំពុងទាញយក... {percent}%\n"
                        f"ទំហំ: {total_size / (1024*1024):.2f} MB"
                    )
                QApplication.processEvents()

                # ពិនិត្យថាអ្នកប្រើប្រាស់ចុច Cancel
                if progress.wasCanceled():
                    raise Exception("ការទាញយកត្រូវបានបោះបង់")

            # កំណត់ timeout socket
            import socket
            old_timeout = socket.getdefaulttimeout()
            socket.setdefaulttimeout(30)  # ៣០ វិន្ទី

            try:
                urllib.request.urlretrieve(url, temp_db, download_progress)
            except urllib.error.HTTPError as e:
                if e.code == 404:
                    raise Exception(
                        "រកមិនឃើញ file នៅលើ server (HTTP 404)\n"
                        "សូមពិនិត្យថា URL ត្រឹមត្រូវ ហើយ file clinic.db មានពិតនៅ GitHub/Cloud។"
                    )
                if e.code == 403:
                    raise Exception(
                        "server មិនអនុញ្ញាតឱ្យទាញយក file នេះទេ (HTTP 403)\n"
                        "សូមពិនិត្យថា repository/file ជា Public ឬមានសិទ្ធិចូលប្រើ។"
                    )
                raise Exception(f"Server error HTTP {e.code}: {e.reason}")
            except urllib.error.URLError as e:
                raise Exception(f"មិនអាចតភ្ជាប់ទៅ URL បានទេ: {e.reason}")
            finally:
                socket.setdefaulttimeout(old_timeout)

            progress.setValue(100)
            progress.setLabelText("✅ ទាញយកជោគជ័យ! កំពុងផ្ទៀងផ្ទាត់...")
            QApplication.processEvents()

            # ៦. ផ្ទៀងផ្ទាត់ថា file ជា SQLite database ត្រឹមត្រូវ
            if not self._validate_sqlite_file(temp_db):
                os.remove(temp_db)
                progress.close()
                QMessageBox.critical(
                    self,
                    "❌ File មិនត្រឹមត្រូវ",
                    "File ដែលទាញយកមិនមែនជា SQLite database ត្រឹមត្រូវទេ!\n\n"
                    "សូមពិនិត្យ URL ឡើងវិញ។\n\n"
                    f"Backup ត្រូវបានរក្សាទុកក្នុង:\n{backup_path}"
                )
                return

            # ៧. Auto Merge ទិន្ននយ (ជានិច្ច មិនជំនួសទេ)
            progress.setLabelText("✅ ផ្ទៀងផ្ទាត់រួចរាល់! កំពុងបញ្ចូលទិន្នន័យ...")
            QApplication.processEvents()
            merged_count, skipped_count = db.merge_database_file(temp_db, self.branch_code)
            msg = (
                "✅ បញ្ចូលទិន្នន័យជោគជ័យ! (Auto Merge)\n\n"
                f"📅 រយៈពេល: {download_period_label}\n"
                f"📊 ចំនួនអ្នកជំងឺថ្មី: {merged_count} នាក់\n"
                f"⏭️ រំលង (Duplicate): {skipped_count} នាក់\n"
                f"💾 Backup ត្រូវបានរក្សាទុកក្នុង:\n{backup_path}"
            )

            # សម្អាត temp file
            os.remove(temp_db)
            progress.close()

            success_dialog = QMessageBox(self)
            success_dialog.setIcon(QMessageBox.Information)
            success_dialog.setWindowTitle("☁️ Cloud Sync ជោគជ័យ")
            success_dialog.setText(
                "<span style='color: #05c46b; font-weight: bold;'>✅ បញ្ចូលទិន្នន័យជោគជ័យ! (Auto Merge)</span><br><br>"
                f"<span style='color: #00cec9;'>📅 រយៈពេល:</span> "
                f"<span style='color: white;'>{download_period_label}</span><br>"
                f"<span style='color: #00cec9;'>📊 ចំនួនអ្នកជំងឺថ្មី:</span> "
                f"<span style='color: white;'>{merged_count} នាក់</span><br>"
                f"<span style='color: #00cec9;'>⏭️ រំលង (Duplicate):</span> "
                f"<span style='color: white;'>{skipped_count} នាក់</span><br><br>"
                f"<span style='color: #ffa801;'>💾 Backup ត្រូវបានរក្សាទុកក្នុង:</span><br>"
                f"<span style='color: #d2dae2;'>{backup_path}</span>"
            )
            success_dialog.setStyleSheet(build_message_box_stylesheet(
                button_bg="#0fbcf9",
                button_color="black",
                button_hover="#00cec9",
                label_min_width=300,
                button_min_width=80,
            ))
            success_dialog.addButton("OK", QMessageBox.AcceptRole)
            success_dialog.exec_()

        except Exception as e:
            if progress is not None:
                progress.close()
            error_msg = str(e)
            if "cancelled" in error_msg.lower() or "បោះបង់" in error_msg:
                QMessageBox.information(self, "បោះបង់", "ការទាញយកត្រូវបានបោះបង់។")
            else:
                # បង្ហាញ URL ពេញលេញ + ពណ៌អក្សរស អានងាយ
                error_dialog = QMessageBox(self)
                error_dialog.setIcon(QMessageBox.Critical)
                error_dialog.setWindowTitle("❌ ទាញយកបរាជ័យ")

                message = (
                    f"<span style='color: #ff3f34; font-weight: bold;'>មិនអាចទាញយកទិន្ននយបានទេ!</span><br><br>"
                    f"<span style='color: #00cec9; font-weight: bold;'>🔗 URL ដែលបានប្រើ:</span><br>"
                    f"<span style='color: white; font-size: 11px;'>{url}</span><br><br>"
                    f"<span style='color: #00cec9; font-weight: bold;'>📋 កំហុស:</span> "
                    f"<span style='color: white;'>{error_msg}</span><br><br>"
                    f"<span style='color: #ffa801; font-weight: bold;'>សូមពិនិត្យ:</span><br>"
                    f"<span style='color: white;'>1. URL ត្រឹមត្រូវ (file មាននៅលើ GitHub)<br>"
                    f"2. ការតភ្ជាប់ Internet<br>"
                    f"3. Repository ជា Public (ឬមាន Token)</span><br><br>"
                    f"<span style='color: #05c46b; font-weight: bold;'>💡 ដំណោះស្រាយ:</span><br>"
                    f"<span style='color: white;'>• ចុច '⬆️ ផ្ញើទៅ Cloud' ដើម្បី Upload database មុន<br>"
                    f"• ឬប្តូរទៅ Repository ផ្សេងដែលមាន file</span><br><br>"
                    f"<span style='color: #d2dae2;'>💾 Backup ត្រូវបានរក្សាទុកក្នុង:<br>{backup_path}</span>"
                )

                error_dialog.setText(message)
                error_dialog.setStyleSheet('''
                    QMessageBox { background-color: #1e272e; }
                    QMessageBox QLabel {
                        color: white;
                        font-size: 12px;
                        min-height: 24px;
                        min-width: 300px;
                        padding: 6px 10px;
                    }
                    QMessageBox QPushButton {
                        background-color: #ff3f34; color: white; font-weight: bold;
                        padding: 8px 20px; border-radius: 5px; min-width: 80px;
                    }
                ''')
                error_dialog.addButton("យល់ហើយ 👍", QMessageBox.AcceptRole)
                error_dialog.exec_()

    def _clear_saved_cloud_sync_url(self):
        """Clear the remembered Cloud Sync URL from settings."""
        try:
            if self.config.has_option('CATEGORIES', 'cloud_sync_url'):
                self.config.remove_option('CATEGORIES', 'cloud_sync_url')
                with open(self.settings_file, 'w', encoding='utf-8') as f:
                    self.config.write(f)
            QMessageBox.information(self, "✅ បានលុប URL", "URL ចាស់សម្រាប់ Cloud Sync ត្រូវបានលុបរួចហើយ។")
            self.statusBar.showMessage("បានលុប URL ចាស់សម្រាប់ Cloud Sync។", 5000)
        except Exception as e:
            QMessageBox.warning(self, "⚠️ មិនអាចលុប URL", f"មិនអាចលុប URL ដែលបានចងចាំបានទេ: {str(e)}")

    def _style_url_input_dialog(self, dialog):
        dialog.setStyleSheet("""
            QInputDialog {
                background-color: #1e272e;
            }
            QLabel {
                color: #f5f6fa;
                font-size: 12px;
                background: transparent;
            }
            QLineEdit {
                background-color: #485460;
                color: #ffffff;
                border: 1px solid #00cec9;
                border-radius: 6px;
                padding: 8px;
                selection-background-color: #00cec9;
                selection-color: #000000;
            }
            QPushButton {
                background-color: #00cec9;
                color: #000000;
                font-weight: bold;
                border-radius: 6px;
                padding: 8px 16px;
                min-width: 90px;
            }
            QPushButton:hover {
                background-color: #00b5b0;
            }
        """)

    def _style_light_input_dialog(self, dialog, editor_selector="QComboBox"):
        editor_style = {
            "QTextEdit": """
            QTextEdit {
                background-color: #f5f6fa;
                color: #000000;
                border: 1px solid #3498db;
                border-radius: 4px;
                padding: 8px;
                selection-background-color: #74b9ff;
                selection-color: #000000;
            }
            """,
            "QComboBox": """
            QComboBox {
                background-color: #f5f6fa;
                color: #000000;
                border: 1px solid #95a5a6;
                border-radius: 4px;
                padding: 6px 8px;
                min-height: 30px;
                selection-background-color: #dfe6e9;
                selection-color: #000000;
            }
            QComboBox QAbstractItemView {
                background-color: #ffffff;
                color: #000000;
                selection-background-color: #74b9ff;
                selection-color: #000000;
                border: 1px solid #95a5a6;
            }
            """,
        }.get(editor_selector, "")

        dialog.setStyleSheet(f"""
            QInputDialog {{
                background-color: #2c3e50;
            }}
            QLabel {{
                color: #f5f6fa;
                font-size: 12px;
                background: transparent;
            }}
            {editor_style}
            QPushButton {{
                background-color: #f5f6fa;
                color: #000000;
                border: 1px solid #74b9ff;
                padding: 6px 18px;
                min-width: 90px;
            }}
            QPushButton:hover {{
                background-color: #dfe6e9;
            }}
        """)

    def _configure_input_dialog(
        self,
        title,
        label,
        text_value="",
        ok_text="OK",
        cancel_text="Cancel",
        size=(430, 180),
        style="light",
        combo_items=None,
        combo_index=None,
        multiline=False,
    ):
        dialog = QInputDialog(self)
        if multiline:
            dialog.setInputMode(QInputDialog.TextInput)
            dialog.setOption(QInputDialog.UsePlainTextEditForTextInput, True)
        dialog.setWindowTitle(title)
        dialog.setLabelText(label)
        if combo_items is not None:
            dialog.setComboBoxItems(combo_items)
            dialog.setComboBoxEditable(False)
        dialog.setTextValue(text_value)
        dialog.setOkButtonText(ok_text)
        dialog.setCancelButtonText(cancel_text)
        dialog.resize(*size)
        if style == "url":
            self._style_url_input_dialog(dialog)
        else:
            self._style_light_input_dialog(dialog, "QTextEdit" if multiline else "QComboBox")
        if combo_index is not None:
            combo = dialog.findChild(QComboBox)
            if combo:
                combo.setCurrentIndex(combo_index)
        return dialog

    def _exec_configured_input_dialog(self, *args, **kwargs):
        dialog = self._configure_input_dialog(*args, **kwargs)
        ok = dialog.exec_() == QDialog.Accepted
        return dialog.textValue(), ok

    def _get_cloud_sync_url_input(self, initial_url):
        """Show a styled input dialog for Cloud Sync URL."""
        return self._exec_configured_input_dialog("Cloud Sync - ទាញយកទិន្នន័យពី Cloud", "សូមបញ្ចូល Link សម្រាប់ទាញយក Database (URL):\n\nURL នឹងត្រូវបានចងចាំស្វ័យប្រវត្តិ។", text_value=initial_url, ok_text="ទាញយក", cancel_text="បោះបង់", size=(640, 180), style="url")

    def _exec_period_choice_dialog(self, title, label, period_options, default_index, ok_text, cancel_text, size):
        dialog = self._configure_input_dialog(
            title,
            label,
            combo_items=period_options,
            text_value=period_options[default_index],
            ok_text=ok_text,
            cancel_text=cancel_text,
            size=size,
            combo_index=default_index,
        )
        if dialog.exec_() != QDialog.Accepted:
            return "", False
        return dialog.textValue(), True

    def _get_cloud_download_period_choice(self):
        """Show period choices and return the matching Cloud Sync database filename."""
        now = datetime.now()
        today_str = now.strftime("%d/%m/%Y")
        period_options = cloud_sync_period_options(include_entered_url=True)
        choice, ok = self._exec_period_choice_dialog(
            "📅 ជ្រើសរើសកាលបរិច្ឆេទទាញយក",
            f"តើអ្នកចង់ទាញយកទិន្នន័យពី Cloud រយៈពេលណា?\n\n"
            f"ថ្ងៃនេះ: {today_str}",
            period_options,
            2,
            "ទាញយក",
            "បោះបង់",
            (460, 190),
        )
        if not ok:
            return None, "", False

        start_date = None
        end_date = None
        period_label = choice

        if "Use Entered URL" in choice:
            return None, "URL ដែលបានបញ្ចូល", True
        if "Today" in choice:
            start_date = now.strftime("%d/%m/%Y")
            end_date = start_date
            period_label = f"ថ្ងៃនេះ ({start_date})"
        elif "This Week" in choice:
            monday = now - timedelta(days=now.weekday())
            sunday = monday + timedelta(days=6)
            start_date = monday.strftime("%d/%m/%Y")
            end_date = sunday.strftime("%d/%m/%Y")
            period_label = f"សប្ដាហ៍នេះ ({start_date} - {end_date})"
        elif "This Month" in choice:
            first_day = now.replace(day=1)
            last_day = now.replace(day=1, month=now.month + 1) - timedelta(days=1) if now.month < 12 else now.replace(year=now.year + 1, month=1, day=1) - timedelta(days=1)
            start_date = first_day.strftime("%d/%m/%Y")
            end_date = last_day.strftime("%d/%m/%Y")
            period_label = f"ខែនេះ ({start_date} - {end_date})"
        elif "Last Month" in choice:
            first_day_last_month = now.replace(day=1, month=now.month - 1) if now.month > 1 else now.replace(year=now.year - 1, month=12, day=1)
            last_day_last_month = now.replace(day=1, month=now.month) - timedelta(days=1)
            start_date = first_day_last_month.strftime("%d/%m/%Y")
            end_date = last_day_last_month.strftime("%d/%m/%Y")
            period_label = f"ខែមុន ({start_date} - {end_date})"
        elif "Custom Range" in choice:
            date_dialog = CustomDateRangeDialog(self, mode_label="ទាញយក")
            if date_dialog.exec_() != QDialog.Accepted:
                return None, "", False
            start_date = date_dialog.start_date.date().toString("dd/MM/yyyy")
            end_date = date_dialog.end_date.date().toString("dd/MM/yyyy")
            period_label = f"កំណត់ដោយខ្លួនឯង ({start_date} - {end_date})"
        elif "Full Database" in choice:
            return "clinic_full.db", "ទាំងអស់ (Full Database)", True

        if not start_date or not end_date:
            return None, "URL ដែលបានបញ្ចូល", True

        safe_start = start_date.replace("/", "-")
        safe_end = end_date.replace("/", "-")
        return f"clinic_{safe_start}_to_{safe_end}.db", period_label, True

    def _replace_cloud_sync_url_filename(self, url, file_name):
        """Replace the last URL path segment with a generated Cloud Sync filename."""
        url = str(url or "").strip()
        file_name = str(file_name or "").strip()
        if not url or not file_name or "/" not in url:
            return ""
        base_url = url.split("?", 1)[0].rsplit("/", 1)[0]
        return f"{base_url}/{file_name}"

    def _get_upload_period_choice(self, today_str):
        """Show a styled selection dialog for upload period."""
        period_options = cloud_sync_period_options()
        return self._exec_period_choice_dialog(
            "📅 ជ្រើសរើសរយៈពេល Upload",
            f"តើអ្នកចង់ Upload របាយការណ៍បែបណា?\n\n"
            f"📊 ថ្ងៃនេះ: {today_str}",
            period_options,
            0,
            "OK",
            "Cancel",
            (430, 180),
        )

    def _get_upload_note_input(self, period_label):
        """Show a styled multiline input dialog for optional upload notes."""
        return self._exec_configured_input_dialog("ការកត់សម្គាល់ (ស្រេចចិត្ត)", f"តើអ្នកចង់បន្ថែមការកត់សម្គាល់សម្រាប់ {period_label}?\n\nឧ. ការងារច្រើន, មានអ្នកជំងឺច្រើន, ខ្វះបុគ្គលិក", text_value="", ok_text="OK", cancel_text="Cancel", size=(520, 360), multiline=True)

    def _preflight_cloud_sync_url(self, url):
        """Lightweight validation before downloading the database file."""
        try:
            request = urllib.request.Request(
                url,
                headers={"User-Agent": f"ClinicManager/{APP_VERSION}"},
                method="HEAD"
            )
            try:
                response = urllib.request.urlopen(request, timeout=10)
            except urllib.error.HTTPError as e:
                if e.code == 405:
                    response = urllib.request.urlopen(
                        urllib.request.Request(url, headers={"User-Agent": f"ClinicManager/{APP_VERSION}"}),
                        timeout=10
                    )
                elif e.code == 404:
                    return False, "រកមិនឃើញ file នៅ URL នេះទេ (HTTP 404)\n\nសូមពិនិត្យថា file មានពិតនៅ GitHub/Cloud។"
                elif e.code == 403:
                    return False, "URL នេះមិនអនុញ្ញាតឱ្យចូលប្រើទេ (HTTP 403)\n\nសូមពិនិត្យថា repository/file ជា Public។"
                else:
                    return False, f"URL នេះមានបញ្ហា HTTP {e.code}: {e.reason}"

            final_url = response.geturl()
            content_type = response.headers.get("Content-Type", "").lower()
            response.close()

            if "text/html" in content_type and "githubusercontent.com" in final_url:
                return False, (
                    "URL នេះត្រឡប់ជា HTML មិនមែន database file ទេ។\n\n"
                    "សូមចុចលើ file ហើយ copy តំណ `Raw` មកប្រើវិញ។"
                )

            return True, ""
        except urllib.error.URLError as e:
            return False, f"មិនអាចតភ្ជាប់ទៅ URL នេះបានទេ:\n{e.reason}"
        except Exception as e:
            return False, f"មិនអាចផ្ទៀងផ្ទាត់ URL បានទេ:\n{str(e)}"
    def _validate_sqlite_file(self, file_path):
        """
        ផ្ទៀងផ្ទាត់ថា file ជា SQLite database ត្រឹមត្រូវ
        """
        try:
            import sqlite3

            # ពិនិត្យ file size (មិនតូចពេក មិនធំពេក)
            file_size = os.path.getsize(file_path)
            if file_size < 512:  # SQLite តូចបំផុត = ៥១២ bytes
                return False
            if file_size > 100 * 1024 * 1024:  # ១០០ MB
                return False

            # ពិនិត្យ SQLite header (bytes ដំបូង = "SQLite format 3")
            with open(file_path, 'rb') as f:
                header = f.read(16)
                if not header.startswith(b'SQLite format 3\x00'):
                    return False

            # ពិនិត្យថាអាច connect និង query បាន
            conn = sqlite3.connect(file_path)
            cur = conn.cursor()
            cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = cur.fetchall()
            conn.close()

            # ត្រូវមានយ៉ាងហោចណាស់ ១ table
            return len(tables) > 0

        except Exception:
            return False

    def check_login(self):
        username = self.user.text().strip()
        if db.check_user(username, self.pwd.text().strip()):
            self.user_context = db.get_user_context(username)
            db.log_login(username) # Log the successful login
            self.accept()
        else:
            QMessageBox.warning(self, "Error", "Invalid Username or Password")

    def show_signup(self):
        signup = SignUpDialog()
        signup.exec_()


class TelegramShareDialog(BaseDialog):
    """Dialog for selecting what type of data to share to Telegram"""
    def __init__(self, parent=None):
        super().__init__("Share Database to Telegram", size=(700, 700),
                        show_creator_header=False, show_creator_footer=False)
        self.parent_app = parent
        self.share_type = None
        self.telegram_contact = ""
        
        # Get Khmer-compatible font
        self.khmer_font = get_khmer_font()
        
        layout = self.content_layout
        layout.setContentsMargins(25, 15, 25, 15)
        layout.setSpacing(12)

        # Title
        title_lbl = QLabel("📤 ចែករំលែកទិន្នន័យទៅ Telegram")
        title_lbl.setFont(QFont(self.khmer_font, 18, QFont.Bold))
        title_lbl.setStyleSheet("""
            color: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #00d2ff, stop:1 #3a7bd5);
            padding: 12px;
            background-color: rgba(0, 0, 0, 0.3);
            border-radius: 8px;
            border: 1px solid rgba(0, 210, 255, 0.3);
        """)
        title_lbl.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        title_lbl.setWordWrap(True)
        layout.addWidget(title_lbl)

        # Description
        desc_lbl = QLabel("សូមជ្រើសរើសប្រភេទទិន្នន័យ និងអ្នកដែលចង់ផ្ញើ:")
        desc_lbl.setFont(QFont(self.khmer_font, 12))
        desc_lbl.setStyleSheet("color: #ecf0f1; padding: 8px; background: transparent;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Telegram Contact Input
        contact_group = QGroupBox("👤 អាសយដ្ឋាន Telegram")
        contact_group.setFont(QFont(self.khmer_font, 11, QFont.Bold))
        contact_group.setStyleSheet("""
            QGroupBox {
                color: #f39c12;
                border: 2px solid #f39c12;
                border-radius: 8px;
                margin-top: 8px;
                padding-top: 8px;
                background-color: rgba(243, 156, 18, 0.05);
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 10px;
            }
        """)
        contact_layout = QVBoxLayout(contact_group)
        contact_layout.setContentsMargins(12, 12, 12, 12)
        contact_layout.setSpacing(8)
        
        self.telegram_input = QLineEdit()
        self.telegram_input.setPlaceholderText("ឧ. @Phunsinouen ឬ +85512345678")
        self.telegram_input.setFixedHeight(40)
        self.telegram_input.setFont(QFont("Segoe UI", 13))
        self.telegram_input.setStyleSheet("""
            QLineEdit {
                padding: 10px;
                border: 2px solid #3498db;
                border-radius: 6px;
                background-color: white;
                color: #2c3e50;
                font-size: 13px;
            }
            QLineEdit:focus {
                border: 2px solid #0fbcf9;
                background-color: #fff;
            }
            QLineEdit:hover {
                border: 2px solid #95a5a6;
            }
        """)
        if self.parent_app and hasattr(self.parent_app, "_get_default_telegram_contact"):
            self.telegram_input.setText(self.parent_app._get_default_telegram_contact())
        contact_layout.addWidget(self.telegram_input)
        
        # Quick contacts
        quick_lbl = QLabel("ជម្រើសរហ័ស:")
        quick_lbl.setFont(QFont(self.khmer_font, 12))
        quick_lbl.setStyleSheet("color: #3498db; background: transparent;")
        contact_layout.addWidget(quick_lbl)
        
        quick_contacts_layout = QHBoxLayout()
        quick_contacts_layout.setSpacing(6)
        
        quick_btn1 = QPushButton("@Phunsinouen")
        quick_btn1.setFixedHeight(32)
        quick_btn1.setFont(QFont("Segoe UI", 10, QFont.Bold))
        quick_btn1.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]
        quick_btn1.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border-radius: 6px;
                font-weight: bold;
                padding: 6px 12px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        quick_btn1.clicked.connect(lambda: self.telegram_input.setText("@Phunsinouen"))
        quick_contacts_layout.addWidget(quick_btn1)
        
        quick_btn2 = QPushButton("💾 Saved Messages")
        quick_btn2.setFixedHeight(32)
        quick_btn2.setFont(QFont("Segoe UI", 10, QFont.Bold))
        quick_btn2.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]
        quick_btn2.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
                border-radius: 6px;
                font-weight: bold;
                padding: 6px 12px;
            }
            QPushButton:hover {
                background-color: #8e44ad;
            }
        """)
        quick_btn2.clicked.connect(lambda: self.telegram_input.setText("Saved Messages"))
        quick_contacts_layout.addWidget(quick_btn2)
        
        quick_btn3 = QPushButton("👥 Group")
        quick_btn3.setFixedHeight(32)
        quick_btn3.setFont(QFont("Segoe UI", 10, QFont.Bold))
        quick_btn3.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]
        quick_btn3.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border-radius: 6px;
                font-weight: bold;
                padding: 6px 12px;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        quick_btn3.clicked.connect(lambda: self.telegram_input.setText("@YourGroupName"))
        quick_contacts_layout.addWidget(quick_btn3)
        
        contact_layout.addLayout(quick_contacts_layout)
        
        layout.addWidget(contact_group)

        # Option 1: Full Database
        full_btn = QPushButton("📊 ទិន្នន័យទាំងស្រុង (Full Database)")
        full_btn.setFixedHeight(65)
        full_btn.setFont(QFont(self.khmer_font, 13, QFont.Bold))
        full_btn.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]
        full_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #667eea, stop:1 #764ba2);
                color: white;
                border-radius: 10px;
                padding: 12px;
                text-align: left;
                font-weight: bold;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #5a6fd6, stop:1 #6a4190);
            }
        """)
        full_btn.clicked.connect(lambda: self.select_type("full"))
        layout.addWidget(full_btn)

        # Option 2: Daily Report
        daily_btn = QPushButton("📅 របាយការណ៍ប្រចាំថ្ងៃ (Daily Report)")
        daily_btn.setFixedHeight(65)
        daily_btn.setFont(QFont(self.khmer_font, 13, QFont.Bold))
        daily_btn.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]
        daily_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #11998e, stop:1 #38ef7d);
                color: white;
                border-radius: 10px;
                padding: 12px;
                text-align: left;
                font-weight: bold;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #0f8a7f, stop:1 #2ed86f);
            }
        """)
        daily_btn.clicked.connect(lambda: self.select_type("daily"))
        layout.addWidget(daily_btn)

        # Info box
        info_box = QLabel("""
            💡 <b>ចំណាំសំខាន់ៗ:</b><br>
            • ទិន្នន័យនឹងត្រូវបង្ហាប់ (ZIP) មុនផ្ញើ<br>
            • អាចផ្ញើទៅ Telegram ណាក៏បាន (Username, Phone, Group, Channel ឬ Saved Messages)<br>
            • សម្រាប់ Group: វាយ @NameOfGroup ឬចុចប៊ូតុង 👥 Group<br>
            • អាចនឹងចំណាយពេលខ្លះអាស្រ័យលើទំហំទិន្នន័យ
        """)
        info_box.setFont(QFont(self.khmer_font, 12))
        info_box.setStyleSheet("""
            color: #f39c12;
            background-color: rgba(243, 156, 18, 0.1);
            padding: 12px;
            border-radius: 6px;
            border-left: 3px solid #f39c12;
        """)
        info_box.setWordWrap(True)
        layout.addWidget(info_box)

        layout.addStretch()

    def select_type(self, share_type):
        self.share_type = share_type
        self.telegram_contact = self.telegram_input.text().strip()
        if not self.telegram_contact:
            QMessageBox.warning(self, "Warning", "សូមបញ្ចូលអាសយដ្ឋាន Telegram ជាមុនសិន!")
            self.telegram_input.setFocus()
            return
        if self.parent_app and hasattr(self.parent_app, "_validate_telegram_contact"):
            is_valid, normalized_contact = self.parent_app._validate_telegram_contact(self.telegram_contact)
            if not is_valid:
                QMessageBox.warning(self, "Warning", normalized_contact)
                self.telegram_input.setFocus()
                self.telegram_input.selectAll()
                return
            self.telegram_contact = normalized_contact
        self.accept()

class LoginHistoryDialog(QDialog):
    def __init__(self, parent=None, mode_label="Upload"):
        super().__init__(parent)
        self.setWindowTitle("Login History")
        self.setMinimumSize(500, 600)
        self.setStyleSheet("""
            QDialog {
                background-color: #1e272e;
                color: #f5f6fa;
            }
            QLineEdit {
                background-color: #2f3640;
                color: #ffffff;
                border: 1px solid #0fbcf9;
                border-radius: 5px;
                padding: 8px;
                min-height: 24px;
            }
            QLineEdit::placeholder {
                color: #b2bec3;
            }
            QTableWidget {
                background-color: #111820;
                alternate-background-color: #1e272e;
                color: #ffffff;
                gridline-color: #485460;
                selection-background-color: #0fbcf9;
                selection-color: #000000;
            }
            QTableWidget::item {
                color: #ffffff;
                padding: 6px;
            }
            QTableWidget::item:selected {
                color: #000000;
            }
            QHeaderView::section {
                background-color: #0fbcf9;
                color: #000000;
                border: none;
                padding: 8px;
                font-weight: bold;
            }
            QDialogButtonBox QPushButton {
                background-color: #0fbcf9;
                color: #000000;
                border: none;
                border-radius: 5px;
                padding: 8px 18px;
                min-width: 80px;
                font-weight: bold;
            }
            QDialogButtonBox QPushButton:hover {
                background-color: #00a8ff;
            }
        """)

        layout = QVBoxLayout(self)

        self.search_bar = QLineEdit()
        self.search_bar.setPlaceholderText("Filter by username or date (e.g., admin, 2024-01-27)...")
        self.search_bar.textChanged.connect(self.filter_history)

        self.table = QTableWidget()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(["Username", "Login Time"])
        self.table.setEditTriggers(QTableWidget.NoEditTriggers) # type: ignore
        self.table.setAlternatingRowColors(True)
        header = self.table.horizontalHeader()
        if header:
            header.setStretchLastSection(True)

        history = db.get_login_history() or []
        self.table.setRowCount(len(history))

        for row_idx, (username, login_time) in enumerate(history):
            self.table.setItem(row_idx, 0, QTableWidgetItem(username))
            self.table.setItem(row_idx, 1, QTableWidgetItem(login_time))

        button_box = QDialogButtonBox(QDialogButtonBox.Ok) # type: ignore
        button_box.accepted.connect(self.accept)

        layout.addWidget(self.search_bar)
        layout.addWidget(self.table)
        layout.addWidget(button_box)

    def filter_history(self, text):
        search_text = text.lower()
        for i in range(self.table.rowCount()):
            username_item = self.table.item(i, 0)
            time_item = self.table.item(i, 1)

            # Check if items exist before accessing text
            match_found = False
            if (username_item and search_text in username_item.text().lower()) or \
               (time_item and search_text in time_item.text().lower()):
                match_found = True
            
            self.table.setRowHidden(i, not match_found)

class SettingsDialog(QDialog):
    def __init__(self, config_parser, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Settings - Manage Categories")
        self.setMinimumSize(600, 500)
        self.config = config_parser
        self.text_edits = {}
        self.theme_combo = QComboBox()
        self.font_combo = QComboBox()

        layout = QVBoxLayout(self)
        
        # Theme Selection at the top
        theme_layout = QHBoxLayout()
        theme_layout.addWidget(QLabel("🎨 ជ្រើសរើសរចនាបទកម្មវិធី (Theme):"))
        self.theme_combo.addItems(["Modern Dark", "Classic Light", "Medical Blue", "Nature Green"])
        self.theme_combo.setCurrentText(self.config.get('CATEGORIES', 'theme', fallback="Modern Dark"))
        theme_layout.addWidget(self.theme_combo)
        layout.addLayout(theme_layout)

        font_layout = QHBoxLayout()
        font_layout.addWidget(QLabel("🔤 ជ្រើសរើស Font Khmer សម្រាប់កម្មវិធីទាំងមូល:"))
        self.font_combo.addItems(available_khmer_fonts())
        self.font_combo.setCurrentText(self.config.get('CATEGORIES', 'khmer_font', fallback="Khmer OS Battambang"))
        font_layout.addWidget(self.font_combo)
        layout.addLayout(font_layout)

        self.tabs_widget = QTabWidget()
        self.tabs_widget.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #485460;
                background-color: #1e272e;
            }
            QTabBar::tab:selected {
                background-color: #0fbcf9;
                color: black;
            }
        """)

        # Define friendly names for sections
        category_labels = {
            'age': 'Age Groups',
            'sex': 'Sex / Gender',
            'area': 'Areas',
            'disease': 'Disease Cases',
            'imci': 'IMCI Options',
            'nutrition': 'Nutrition Status',
            'service': 'Service Types',
            'diagnosis': 'Diagnosis (រោគវិនិច្ឆ័យ)',
            'medicine': 'Medicines (Auto-complete)'
        }

        for key, label in category_labels.items():
            page = QWidget()
            page_layout = QVBoxLayout(page)
            
            # Load current values from config
            current_val = self.config.get('CATEGORIES', key, fallback="")
            
            editor = QTextEdit()
            editor.setPlainText(current_val)
            editor.setPlaceholderText("Enter items separated by commas (e.g., Item1, Item2, Item3)")
            
            page_layout.addWidget(QLabel(f"Edit list for {label}:"))
            page_layout.addWidget(editor)
            
            self.text_edits[key] = editor
            self.tabs_widget.addTab(page, label)

        btn_box = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel) # type: ignore
        btn_box.accepted.connect(self.save_settings)
        btn_box.rejected.connect(self.reject)

        layout.addWidget(self.tabs_widget)
        layout.addWidget(btn_box)

    def save_settings(self):
        # Save Theme
        self.config.set('CATEGORIES', 'theme', self.theme_combo.currentText())
        self.config.set('CATEGORIES', 'khmer_font', self.font_combo.currentText())
        
        for key, editor in self.text_edits.items():
            # Clean up the input: remove newlines, extra spaces
            raw_text = editor.toPlainText().replace('\n', ',')
            cleaned_items = [item.strip() for item in raw_text.split(',') if item.strip()]
            self.config.set('CATEGORIES', key, ', '.join(cleaned_items))
        
        self.accept()

class App(QWidget):
    def __init__(self, username, user_context=None):
        super().__init__()
        self.setWindowTitle(f"Clinic Management System - v{APP_VERSION}")
        self._configure_window_for_screen()
        self.current_user = username
        self.user_context = user_context or db.get_user_context(username)
        self.is_admin = bool(self.user_context.get("is_admin"))
        self.branch_code = self.user_context.get("branch_code", "MAIN")
        self.active_branch_code = None if self.is_admin else self.branch_code

        # កំណត់ Icon សម្រាប់កម្មវិធី
        self._set_app_icon()

        # កំណត់ផ្លូវឯកសារឱ្យច្បាស់លាស់ (ទោះបីជាដំណើរការចេញពី .exe ក៏ដោយ)
        if getattr(sys, 'frozen', False):
            self.base_dir = os.path.dirname(sys.executable)
        else:
            self.base_dir = os.path.dirname(os.path.abspath(__file__))

        # --- Load Categories from settings.ini ---
        self.config = configparser.ConfigParser()
        self.settings_file = os.path.join(self.base_dir, 'settings.ini')
        self._load_or_create_settings()

        self._reload_category_settings()

        # Set backup directory to AppData to avoid permission issues
        import ctypes.wintypes
        
        # Use AppData for backups
        CSIDL_APPDATA = 26  # Roaming AppData
        SHGFP_TYPE_CURRENT = 0
        buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
        ctypes.windll.shell32.SHGetFolderPathW(None, CSIDL_APPDATA, None, SHGFP_TYPE_CURRENT, buf)
        backup_base = os.path.join(buf.value, "ClinicManager", "backups")
        
        # Try to create backup directory, but don't fail if it doesn't work
        try:
            if not os.path.exists(backup_base):
                os.makedirs(backup_base)
            self.backup_dir = backup_base
        except Exception as e:
            # If we can't create backup folder, use temp folder
            temp_dir = os.getenv('TEMP') or os.getenv('TMP') or os.getcwd()
            self.backup_dir = os.path.join(temp_dir, "ClinicManager", "backups")
            if not os.path.exists(self.backup_dir):
                os.makedirs(self.backup_dir)
        
        print(f"Backup directory: {self.backup_dir}")

        # Store the ID of the currently selected patient for editing/deleting
        self.selected_id = None
        self.btn_add = None # រក្សាទុកប៊ូតុង ADD សម្រាប់ប្រើក្នុង Tab Order

        # Track currently active patient type based on tab
        # Default to Child
        self.current_patient_type = "Child"
        self._preserve_entry_on_tab_switch = False
        self._loading_table = False
        self._populating_form = False

        # --- Main Layout with Tabs ---
        main_layout = QVBoxLayout()

        title = QLabel("ប្រព័ន្ធគ្រប់គ្រងគ្លីនិក (CLINIC MANAGEMENT SYSTEM)")
        title.setAlignment(Qt.AlignCenter) # type: ignore
        title.setStyleSheet("""
            font-size:20px;
            font-weight:bold;
            color:#00cec9;
            padding:6px;
        """)
        main_layout.addWidget(title)

        self.tabs = QTabWidget()
        self.tabs.setStyleSheet("""
            QTabWidget::pane { border: 1px solid #485460; }
            QTabBar::tab {
                background: #2f3640;
                color: white;
                padding: 7px 14px;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
            }
            QTabBar::tab:selected {
                background: #00cec9;
                color: black;
                font-weight: bold;
            }
        """)

        # --- Sub-tabs for Data Entry (Child vs Adult) ---
        self.entry_tabs = QTabWidget()
        self.tab_child = QWidget()
        self.tab_adult = QWidget()
        
        # Initialize the shared form on the first tab (Child) initially
        # We will use a trick: The form is the same, but we move it between tabs or just change the context
        # Better approach: Keep the form in a main widget, and use tabs just to toggle the 'Type'
        
        self.init_entry_tab() 

        self.init_report_tab()
        self.init_help_tab()

        # --- បន្ថែម Footer ព័ត៌មានអ្នកបង្កើតនៅខាងក្រោម Tab ---
        footer_layout = QHBoxLayout()
        creator_info = QLabel('រក្សាសិទ្ធិដោយ៖ <b>នូរ សារ៉ាត់ (NOU SARAT)</b>')
        creator_info.setStyleSheet("color: #f5f6fa; font-size: 13px; font-weight: bold;")
        tg_link = QLabel('<a href="https://t.me/nousarat" style="color: #0fbcf9; text-decoration: none;">✈️ Telegram</a>')
        yt_link = QLabel('<a href="https://www.youtube.com/@TechFree2026" style="color: #ff4757; text-decoration: none;">📺 YouTube</a>')
        separator_1 = QLabel("|")
        separator_2 = QLabel("|")
        separator_1.setStyleSheet("color: #d2dae2; font-weight: bold;")
        separator_2.setStyleSheet("color: #d2dae2; font-weight: bold;")
        
        for lbl in [tg_link, yt_link]:
            lbl.setOpenExternalLinks(True)
        
        footer_layout.addStretch()
        footer_layout.addWidget(creator_info)
        footer_layout.addWidget(separator_1)
        footer_layout.addWidget(tg_link)
        footer_layout.addWidget(separator_2)
        footer_layout.addWidget(yt_link)
        footer_layout.addStretch()
        
        main_layout.addLayout(footer_layout)
        main_layout.addWidget(self.tabs)

        # Add Status Bar for user feedback
        self.statusBar = QStatusBar()
        branch_label = "ALL" if self.is_admin else self.branch_code
        self.statusBar.showMessage(f"សូមស្វាគមន៍, {self.current_user}! Branch: {branch_label}")
        main_layout.addWidget(self.statusBar)

        self.apply_app_font(self.current_font)

        self.setLayout(main_layout)

        # Apply the selected theme
        self.apply_theme(self.current_theme)
        
        self.view()

        # កំណត់ឱ្យកម្មវិធីប្តូរភាសាក្តារចុចដោយស្វ័យប្រវត្តិតាមប្រអប់បញ្ចូល
        # Setup automatic keyboard layout switching
        # Use isinstance check to satisfy Pylance type checker
        app_inst = QApplication.instance()
        if isinstance(app_inst, QApplication):
            app_inst.focusChanged.connect(self._on_focus_changed)

    def _reload_category_settings(self):
        self.CAT_AGE = self._get_setting('age')
        self.CAT_SEX = self._get_setting('sex')
        self.CAT_AREA = self._get_setting('area')
        self.CAT_DISEASE = self._get_setting('disease')
        self.CAT_IMCI = self._get_setting('imci')
        self.CAT_NUTRITION = self._get_setting('nutrition')
        self.CAT_SERVICE = self._get_setting('service')
        self.CAT_DIAGNOSIS = self._get_setting('diagnosis')
        self.CAT_MEDICINE = self._get_setting('medicine')
        self.current_theme = self.config.get('CATEGORIES', 'theme', fallback="Modern Dark")
        self.current_font = self._get_configured_font()

    def _configure_window_for_screen(self):
        """Size the main window to the usable screen area on high-DPI laptops."""
        screen = QApplication.primaryScreen()
        if not screen:
            self.resize(1200, 750)
            return

        available = screen.availableGeometry()
        width = min(1200, max(820, int(available.width() * 0.96)), available.width())
        height = min(750, max(560, int(available.height() * 0.92)), available.height())
        self.resize(width, height)
        self.setMinimumSize(min(820, available.width()), min(560, available.height()))

        frame = self.frameGeometry()
        frame.moveCenter(available.center())
        self.move(frame.topLeft())

        if available.width() < 1400 or available.height() < 800:
            self.showMaximized()

    def _set_app_icon(self):
        """កំណត់ Icon សម្រាប់កម្មវិធី (Taskbar, Title Bar, etc.)"""
        try:
            from PyQt5.QtGui import QIcon

            icon_path = None

            if getattr(sys, 'frozen', False):
                # ពេល build ជា .exe → icon មានក្នុង _MEIPASS temp folder
                meipass = getattr(sys, '_MEIPASS', None)
                if meipass:
                    icon_path = os.path.join(meipass, "healthcare.ico")

                # បើមិនមានក្នុង _MEIPASS ព្យាយាមរកក្នុង exe directory
                if not icon_path or not os.path.exists(icon_path):
                    icon_path = os.path.join(self.base_dir, "healthcare.ico")
            else:
                # ពេល run ពី source → រកក្នុង script directory
                icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "healthcare.ico")

            # បើមាន icon file → កំណត់
            if icon_path and os.path.exists(icon_path):
                icon = QIcon(icon_path)
                self.setWindowIcon(icon)

            # កំណត់ icon សម្រាប់ Taskbar (Windows)
            if sys.platform == 'win32':
                try:
                    myappid = 'nousarat.clinic.manager.1.0'  # arbitrary string
                    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
                except Exception:
                    pass
        except Exception:
            # បើ fail មិនបង្ហាញ error ទេ (just skip)
            pass

    def apply_theme(self, theme_name):
        """កំណត់រចនាបទកម្មវិធីទាំងមូល (Apply Global Stylesheet)"""
        themes = {
            "Modern Dark": {
                "bg": "#0f1720", "card": "#1e272e", "text": "white", 
                "accent": "#0fbcf9", "input": "#485460", "header": "#0fbcf9", "header_text": "black"
            },
            "Classic Light": {
                "bg": "#f5f6fa", "card": "#ffffff", "text": "#000000",
                "accent": "#0984e3", "input": "#ffffff", "header": "#dcdde1", "header_text": "black"
            },
            "Medical Blue": {
                "bg": "#1a2a6c", "card": "#1e3799", "text": "white", 
                "accent": "#4a69bd", "input": "#4b7bec", "header": "#f8c291", "header_text": "black"
            },
            "Nature Green": {
                "bg": "#1b3022", "card": "#2d4a3e", "text": "white", 
                "accent": "#2ecc71", "input": "#218c74", "header": "#a7c957", "header_text": "black"
            }
        }
        
        t = themes.get(theme_name, themes["Modern Dark"])
        self.current_theme = theme_name
        dialog_input_bg = "#f5f6fa" if theme_name != "Classic Light" else "#ffffff"
        dialog_input_text = "#000000"
        dialog_button_bg = t['accent']
        dialog_button_text = "black" if theme_name != "Classic Light" else "white"
        dialog_border = t['accent']

        # Global Stylesheet
        self.setStyleSheet(f"""
            QWidget {{
                background-color: {t['bg']};
                color: {t['text']};
                font-size: 12px;
            }}
            QLabel {{
                min-height: 24px;
            }}
            QGroupBox {{
                background-color: {t['card']};
                border: 1px solid {t['accent']};
                color: {t['accent']};
                margin-top: 18px;
                padding: 8px;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }}
            QLineEdit, QTextEdit, QPlainTextEdit, QComboBox, QDateEdit {{
                background-color: {t['input']};
                color: {t['text']};
                border-radius: 5px;
                padding: 5px 6px;
                min-height: 24px;
            }}
            QAbstractItemView {{
                background-color: {dialog_input_bg};
                color: {dialog_input_text};
                selection-background-color: {t['accent']};
                selection-color: black;
            }}
            QTabWidget::pane {{
                border: 1px solid {t['accent']};
            }}
            QTabBar::tab {{
                background: {t['card']};
                color: {t['text']};
            }}
            QTabBar::tab:selected {{
                background: {t['accent']};
                color: {"white" if theme_name != "Classic Light" else "black"};
            }}
            QTableWidget {{
                background-color: {t['card']};
                gridline-color: {t['accent']};
            }}
            QHeaderView::section {{
                background-color: {t['header']};
                color: {"white" if t['header'] != "#f8c291" else "black"};
            }}
            QMessageBox {{
                background-color: {t['card']};
            }}
            QMessageBox QLabel {{
                color: {t['text']};
                min-height: 24px;
                min-width: 260px;
                padding: 6px 10px;
                line-height: 1.3;
            }}
            QMessageBox QPushButton {{
                background-color: {dialog_button_bg};
                color: {dialog_button_text};
                border: 1px solid {dialog_border};
                border-radius: 5px;
                padding: 8px 18px;
                min-width: 90px;
                font-weight: bold;
            }}
            QMessageBox QPushButton:hover {{
                background-color: {t['header']};
                color: {"white" if t['header'] not in ["#f8c291", "#dcdde1"] else "black"};
            }}
            QDialog {{
                background-color: {t['card']};
                color: {t['text']};
            }}
            QDialog QLabel {{
                color: {t['text']};
                background: transparent;
            }}
            QDialog QLineEdit, QDialog QTextEdit, QDialog QPlainTextEdit,
            QDialog QComboBox, QDialog QDateEdit {{
                background-color: {dialog_input_bg};
                color: {dialog_input_text};
                border: 1px solid {dialog_border};
                border-radius: 5px;
                padding: 6px;
            }}
            QDialog QComboBox QAbstractItemView {{
                background-color: {dialog_input_bg};
                color: {dialog_input_text};
                selection-background-color: {t['accent']};
                selection-color: black;
                border: 1px solid {dialog_border};
            }}
            QDialog QPushButton, QDialogButtonBox QPushButton {{
                background-color: {dialog_button_bg};
                color: {dialog_button_text};
                border: 1px solid {dialog_border};
                border-radius: 5px;
                padding: 6px 16px;
                min-width: 90px;
                font-weight: bold;
            }}
            QDialog QPushButton:hover, QDialogButtonBox QPushButton:hover {{
                background-color: {t['header']};
                color: {"white" if t['header'] not in ["#f8c291", "#dcdde1"] else "black"};
            }}
            QMenu {{
                background-color: {t['card']};
                border: 1px solid {t['accent']};
                color: {t['text']};
            }}
            QMenu::item:selected {{
                background-color: {t['accent']};
                color: black;
            }}
        """)
        
        # Update Specific Elements that might have inline styles
        if hasattr(self, 'table'):
            # បញ្ឈប់ការប្រើ setStyleSheet ជាន់គ្នាដែលធ្វើឱ្យបាត់ពណ៌អក្សរ
            pass 

        if hasattr(self, 'rep_table'):
            pass

        if hasattr(self, 'search_input'):
            self.search_input.setStyleSheet(f"background: {t['input']}; border: 1px solid {t['accent']}; color: {t['text']}; padding: 8px;")

        # Refresh individual Tabs if necessary
        for i in range(self.tabs.count()):
            w = self.tabs.widget(i)
            if w:
                w.update()

    def init_report_tab(self):
        self.tab_report = QWidget()
        self.tabs.addTab(self.tab_report, "📊 វិភាគ និង សរុបទិន្នន័យ")
        layout = QVBoxLayout(self.tab_report)
        layout.setContentsMargins(12, 10, 12, 10)
        layout.setSpacing(10)

        # --- Analytics Settings Section ---
        analytics_group = QGroupBox("⚙️ កំណត់ការវិភាគ (Analytics Settings)")
        analytics_group.setStyleSheet("""
            QGroupBox {
                background-color: #1e272e;
                border: 1px solid #485460;
                border-radius: 6px;
                margin-top: 16px;
                padding: 10px;
                font-weight: bold;
                color: #0fbcf9;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                left: 10px;
                padding: 0 6px;
            }
            QLabel {
                color: #d2dae2;
                font-weight: bold;
                min-height: 24px;
            }
            QComboBox, QDateEdit {
                background-color: #353b48;
                color: #f5f6fa;
                border: 1px solid #485460;
                border-radius: 5px;
                padding: 5px 8px;
                min-height: 26px;
            }
            QComboBox:focus, QDateEdit:focus {
                border: 1px solid #0fbcf9;
            }
            QComboBox QAbstractItemView {
                background-color: #f5f6fa;
                color: #111820;
                selection-background-color: #0fbcf9;
                selection-color: #000000;
            }
        """)
        analytics_layout = QGridLayout(analytics_group)
        analytics_layout.setContentsMargins(10, 18, 10, 10)
        analytics_layout.setHorizontalSpacing(10)
        analytics_layout.setVerticalSpacing(8)
        analytics_layout.setColumnStretch(1, 1)
        analytics_layout.setColumnStretch(3, 1)

        # Group By Dropdown
        self.group_by_label = QLabel("ដាក់ជាក្រុមតាម (Group By):")
        self.group_by_combo = QComboBox()
        self.group_by_combo.addItems([
            "តំបន់ (Area)",
            "ភេទ (Sex)",
            "ករណីជំងឺ (Disease Case)",
            "អាយុ (Age Groups)",
            "សេវាកម្ម (Service)",
            "ប្រភេទអ្នកជំងឺ (Patient Type)",
            "រោគវិនិច្ឆ័យ (Diagnosis)"
        ])
        self.group_by_combo.setCurrentIndex(0)

        # Sub-Group By Dropdown (Breakdown)
        self.sub_group_by_label = QLabel("បំបែកក្រុមរងតាម (Sub-Group By):")
        self.sub_group_by_combo = QComboBox()
        self.sub_group_by_combo.addItem("— គ្មាន (None) —")
        self.sub_group_by_combo.addItems([
            "ភេទ (Sex)",
            "តំបន់ (Area)",
            "ករណីជំងឺ (Disease Case)",
            "អាយុ (Age Groups)",
            "សេវាកម្ម (Service)",
            "ប្រភេទអ្នកជំងឺ (Patient Type)",
            "រោគវិនិច្ឆ័យ (Diagnosis)"
        ])
        self.sub_group_by_combo.setCurrentIndex(0)

        # Sort By Dropdown
        self.sort_by_label = QLabel("តម្រៀបតាម (Sort By):")
        self.sort_by_combo = QComboBox()
        self.sort_by_combo.addItems([
            "ចំនួន ច្រើន→តិច (Count Desc)",
            "ចំនួន តិច→ច្រើន (Count Asc)",
            "ឈ្មោះ A→Z (Name A-Z)",
            "ឈ្មោះ Z→A (Name Z-A)"
        ])
        self.sort_by_combo.setCurrentIndex(0)

        # Filter Options
        self.rep_start_date = QDateEdit()
        self.rep_start_date.setCalendarPopup(True)
        self.rep_start_date.setDisplayFormat("dd/MM/yyyy")
        self.rep_start_date.setDate(QDate.currentDate().addMonths(-1))

        self.rep_end_date = QDateEdit()
        self.rep_end_date.setCalendarPopup(True)
        self.rep_end_date.setDisplayFormat("dd/MM/yyyy")
        self.rep_end_date.setDate(QDate.currentDate())

        self.rep_sex = QComboBox()
        self.rep_sex.addItem("All")
        self.rep_sex.addItems(self.CAT_SEX)

        self.rep_area = QComboBox()
        self.rep_area.addItem("All")
        self.rep_area.addItems(self.CAT_AREA)

        self.rep_disease = QComboBox()
        self.rep_disease.addItem("All")
        self.rep_disease.addItems(self.CAT_DISEASE)

        self.rep_patient_type = QComboBox()
        self.rep_patient_type.addItem("All")
        self.rep_patient_type.addItems(["Child", "Adult"])

        for control in [
            self.group_by_combo, self.sub_group_by_combo, self.sort_by_combo,
            self.rep_start_date, self.rep_end_date, self.rep_sex,
            self.rep_area, self.rep_disease, self.rep_patient_type,
        ]:
            control.setMinimumHeight(34)

        analytics_layout.addWidget(self.group_by_label, 0, 0)
        analytics_layout.addWidget(self.group_by_combo, 0, 1)
        analytics_layout.addWidget(self.sub_group_by_label, 0, 2)
        analytics_layout.addWidget(self.sub_group_by_combo, 0, 3)

        analytics_layout.addWidget(self.sort_by_label, 1, 0)
        analytics_layout.addWidget(self.sort_by_combo, 1, 1)
        analytics_layout.addWidget(QLabel("ពីថ្ងៃ:"), 1, 2)
        analytics_layout.addWidget(self.rep_start_date, 1, 3)

        analytics_layout.addWidget(QLabel("ដល់ថ្ងៃ:"), 2, 0)
        analytics_layout.addWidget(self.rep_end_date, 2, 1)
        analytics_layout.addWidget(QLabel("ភេទ:"), 2, 2)
        analytics_layout.addWidget(self.rep_sex, 2, 3)

        analytics_layout.addWidget(QLabel("តំបន់:"), 3, 0)
        analytics_layout.addWidget(self.rep_area, 3, 1)
        analytics_layout.addWidget(QLabel("ប្រភេទអ្នកជំងឺ:"), 3, 2)
        analytics_layout.addWidget(self.rep_patient_type, 3, 3)

        analytics_layout.addWidget(QLabel("ជំងឺ:"), 4, 0)
        analytics_layout.addWidget(self.rep_disease, 4, 1)

        # Generate Button
        self.btn_analytics = QPushButton("📊 បង្កើតរបាយការណ៍វិភាគ")
        self.btn_analytics.setMinimumHeight(38)
        self.btn_analytics.setStyleSheet("background-color: #00cec9; color: black; font-weight: bold; padding: 8px 12px; border-radius: 5px;")
        self.btn_analytics.clicked.connect(self.generate_analytics_report)
        analytics_layout.addWidget(self.btn_analytics, 5, 0, 1, 2)

        # Export Button
        self.btn_export_analytics = QPushButton("📊 Export របាយការណ៍នេះ")
        self.btn_export_analytics.setMinimumHeight(38)
        self.btn_export_analytics.setStyleSheet("background-color: #20bf6b; color: white; font-weight: bold; padding: 8px 12px; border-radius: 5px;")
        self.btn_export_analytics.clicked.connect(self.export_analytics_report)
        analytics_layout.addWidget(self.btn_export_analytics, 5, 2, 1, 2)

        # Clear Button
        self.btn_clear_analytics = QPushButton("🔄 លុបការកំណត់")
        self.btn_clear_analytics.setMinimumHeight(36)
        self.btn_clear_analytics.setStyleSheet("background-color: #57606f; color: white; font-weight: bold; padding: 8px 12px; border-radius: 5px;")
        self.btn_clear_analytics.clicked.connect(self.clear_analytics_report)
        analytics_layout.addWidget(self.btn_clear_analytics, 6, 0, 1, 2)

        # Advanced Query Button
        self.btn_advanced_query = QPushButton("🔍 ស្វែងរកកម្រិតខ្ពស់ (Advanced Query)")
        self.btn_advanced_query.setMinimumHeight(36)
        self.btn_advanced_query.setStyleSheet("background-color: #e74c3c; color: white; font-weight: bold; padding: 8px 12px; border-radius: 5px;")
        self.btn_advanced_query.clicked.connect(self.open_advanced_query)
        analytics_layout.addWidget(self.btn_advanced_query, 6, 2, 1, 2)

        self.btn_user_audit = QPushButton("👤 ពិនិត្យទិន្នន័យតាម User")
        self.btn_user_audit.setMinimumHeight(36)
        self.btn_user_audit.setStyleSheet("background-color: #f39c12; color: black; font-weight: bold; padding: 8px 12px; border-radius: 5px;")
        self.btn_user_audit.clicked.connect(self.open_user_audit_dialog)
        analytics_layout.addWidget(self.btn_user_audit, 7, 0, 1, 4)

        # Total Count Label
        self.lbl_analytics_count = QLabel("ក្រុមសរុប: 0")
        self.lbl_analytics_count.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        self.lbl_analytics_count.setStyleSheet("background-color: #111820; border: 1px solid #0fbcf9; border-radius: 5px; padding: 8px; font-weight: bold; color: #00cec9; font-size: 15px;")
        analytics_layout.addWidget(self.lbl_analytics_count, 8, 0, 1, 4)

        layout.addWidget(analytics_group)

        summary_group = QGroupBox("សរុបទិន្នន័យ")
        summary_group.setStyleSheet("""
            QGroupBox {
                background-color: #1e272e;
                border: 1px solid #485460;
                border-radius: 6px;
                margin-top: 14px;
                padding: 8px;
                color: #0fbcf9;
                font-weight: bold;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 6px;
            }
        """)
        summary_layout = QGridLayout(summary_group)
        summary_layout.setContentsMargins(8, 16, 8, 8)
        summary_layout.setHorizontalSpacing(8)
        self.lbl_summary_total = QLabel("ទិន្នន័យសរុប: 0")
        self.lbl_summary_male = QLabel("ប្រុស: 0")
        self.lbl_summary_female = QLabel("ស្រី: 0")
        self.lbl_summary_top = QLabel("ក្រុមខ្ពស់បំផុត: -")
        for widget in [self.lbl_summary_total, self.lbl_summary_male, self.lbl_summary_female, self.lbl_summary_top]:
            widget.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
            widget.setMinimumHeight(42)
            widget.setStyleSheet("background-color: #111820; border: 1px solid #485460; border-radius: 5px; padding: 8px; font-weight: bold; color: #f5f6fa;")
        summary_layout.addWidget(self.lbl_summary_total, 0, 0)
        summary_layout.addWidget(self.lbl_summary_male, 0, 1)
        summary_layout.addWidget(self.lbl_summary_female, 0, 2)
        summary_layout.addWidget(self.lbl_summary_top, 0, 3)
        layout.addWidget(summary_group)

        # --- Result Table ---
        self.analytics_table = QTableWidget()
        self.analytics_table.setEditTriggers(QTableWidget.NoEditTriggers) # type: ignore
        self.analytics_table.setSortingEnabled(True) # type: ignore
        self.analytics_table.setAlternatingRowColors(True)
        self.analytics_table.setSelectionBehavior(QTableWidget.SelectRows)  # type: ignore[attr-defined]
        self.analytics_table.verticalHeader().setVisible(False)
        self.analytics_table.horizontalHeader().setStretchLastSection(True)
        self.analytics_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.analytics_table.setStyleSheet("""
            QTableWidget {
                background-color: #111820;
                alternate-background-color: #1e272e;
                color: #f5f6fa;
                border: 1px solid #485460;
                border-radius: 6px;
                gridline-color: #485460;
                selection-background-color: #0fbcf9;
                selection-color: #000000;
            }
            QTableWidget::item {
                padding: 6px;
            }
            QHeaderView::section {
                background-color: #0fbcf9;
                color: #000000;
                border: none;
                padding: 8px;
                font-weight: bold;
            }
        """)
        self.analytics_table.setContextMenuPolicy(Qt.CustomContextMenu) # type: ignore
        self.analytics_table.customContextMenuRequested.connect(self.show_analytics_context_menu) # type: ignore
        self.analytics_table.doubleClicked.connect(self.drill_down_details) # type: ignore
        layout.addWidget(self.analytics_table)

        # Store analytics data
        self.analytics_data = []

    def init_help_tab(self):
        """បង្កើត Tab សម្រាប់បង្ហាញការណែនាំ និង Shortcut Keys"""
        self.tab_help = QWidget()
        self.tabs.addTab(self.tab_help, "ការណែនាំ")
        layout = QVBoxLayout(self.tab_help)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        content = QWidget()
        scroll_layout = QVBoxLayout(content)

        # --- ផ្នែក Shortcut Keys ---
        shortcut_group = QGroupBox("📌 គ្រាប់ចុចកាត់ (Shortcut Keys)")
        shortcut_group.setStyleSheet("QGroupBox { font-size: 18px; font-weight: bold; color: #0fbcf9; border: 1px solid #485460; margin-top: 20px; padding: 10px; }")
        s_layout = QGridLayout(shortcut_group)
        
        shortcuts = [
            ("Alt + A", "បន្ថែមទិន្នន័យ (ADD)"),
            ("Alt + E", "កែប្រែទិន្នន័យ (EDIT)"),
            ("Alt + D", "លុបទិន្នន័យ (DELETE)"),
            ("Alt + F", "ស្វែងរកឈ្មោះអ្នកជំងឺ (SEARCH)"),
            ("F5", "បង្ហាញទិន្នន័យទាំងអស់ឡើងវិញ (VIEW ALL)"),
            ("Alt + L", "សម្អាតប្រអប់បញ្ចូល (CLEAR)"),
            ("Alt + X", "ទាញទិន្នន័យចេញជា Excel"),
            ("Ctrl + P", "បោះពុម្ពប័ណ្ណអ្នកជំងឺ (PDF)"),
            ("Ctrl + 1", "ប្តូរទៅរបៀប កុមារ (Child Mode)"),
            ("Ctrl + 2", "ប្តូរទៅរបៀប មនុស្សចាស់ (Adult Mode)"),
            ("Ctrl + Q", "ចាកចេញពីកម្មវិធី (LOGOUT)"),
            ("Enter", "លោតទៅប្រអប់បញ្ចូលបន្ទាប់"),
        ]

        for i, (key, desc) in enumerate(shortcuts):
            lbl_key = QLabel(key)
            lbl_key.setStyleSheet("color: #ffdd59; font-weight: bold; font-size: 15px;")
            s_layout.addWidget(lbl_key, i, 0)
            lbl_desc = QLabel(f": {desc}")
            lbl_desc.setStyleSheet("font-size: 15px;")
            s_layout.addWidget(lbl_desc, i, 1)

        # --- ផ្នែកពន្យល់អំពីប៊ូតុង ---
        btn_desc_group = QGroupBox("🔘 មុខងាររបស់ប៊ូតុងសំខាន់ៗ")
        btn_desc_group.setStyleSheet("QGroupBox { font-size: 18px; font-weight: bold; color: #05c46b; border: 1px solid #485460; margin-top: 20px; padding: 10px; }")
        b_layout = QVBoxLayout(btn_desc_group)
        
        buttons_info = [
            ("<b>ADD:</b> រក្សាទុកទិន្នន័យថ្មីចូលក្នុងប្រព័ន្ធ។ លេខរៀងអ្នកជំងឺ (C-1, A-1) នឹងរត់អូតូ។", "#05c46b"),
            ("<b>EDIT:</b> កែប្រែព័ត៌មានអ្នកជំងឺដែលបានរើសក្នុងតារាង។", "#0fbcf9"),
            ("<b>DELETE:</b> លុបទិន្នន័យអ្នកជំងឺដែលបានរើសចេញ។ ប្រព័ន្ធនឹងរៀបលេខរៀងឡើងវិញអូតូ។", "#ff3f34"),
            ("<b>SEARCH:</b> ស្វែងរកតាម ឈ្មោះ, លេខរៀង, លេខប័ណ្ណ ឬ លេខទូរស័ព្ទ។", "#9b59b6"),
            ("<b>STATS:</b> បង្ហាញស្ថិតិអ្នកជំងឺសរុបប្រចាំថ្ងៃ និងតាមប្រភេទសេវា។", "#34495e"),
            ("<b>EXCEL:</b> ទាញទិន្នន័យចូលក្នុង Template Excel ផ្លូវការ (Child/Adult)។", "#00d2d3"),
            ("<b>MERGE DB:</b> បញ្ចូលទិន្នន័យពី File Database កុំព្យូទ័រផ្សេងទៀតចូលគ្នា។", "#6c5ce7"),
            ("<b>SETTINGS:</b> កំណត់បញ្ជីឈ្មោះថ្នាំ តំបន់ និងប្រភេទជំងឺក្នុង Dropdown។", "#576574")
        ]

        for info, color in buttons_info:
            lbl = QLabel(info)
            lbl.setWordWrap(True)
            lbl.setStyleSheet(f"border-left: 4px solid {color}; padding-left: 10px; margin-bottom: 5px; font-size: 15px;")
            b_layout.addWidget(lbl)

        # --- ផ្នែកគន្លឹះបញ្ចូលទិន្នន័យ ---
        tip_group = QGroupBox("💡 គន្លឹះបញ្ចូលទិន្នន័យឱ្យបានលឿន")
        tip_group.setStyleSheet("QGroupBox { font-size: 18px; font-weight: bold; color: #ffa801; border: 1px solid #485460; margin-top: 20px; padding: 10px; }")
        t_layout = QVBoxLayout(tip_group)
        
        tips = [
            "<b>• កាលបរិច្ឆេទ:</b> វាយតែលេខ DDMMYY (ឧ. 270125) រួចចុច Enter វានឹងដូរទៅ 27/01/2025 អូតូ។",
            "<b>• ការប្តូរភាសា:</b> កម្មវិធីនឹងប្តូរក្តារចុច ខ្មែរ/អង់គ្លេស ឱ្យអ្នកដោយស្វ័យប្រវត្តិតាមប្រអប់បញ្ចូល។",
            "<b>• Auto-complete:</b> ពេលវាយឈ្មោះថ្នាំ ឬអាស័យដ្ឋាន កម្មវិធីនឹងបង្ហាញពាក្យដែលធ្លាប់វាយពីមុនមកឱ្យរើស។",
            "<b>• Tab Order:</b> ប្រើគ្រាប់ចុច Tab ដើម្បីលោតពីប្រអប់មួយទៅមួយតាមលំដាប់លំដោយ។",
            "<b>• ប្រភេទអាយុ:</b> ត្រូវជ្រើសរើសប្រភេទអាយុឱ្យបានត្រឹមត្រូវ ដើម្បីឱ្យទិន្នន័យបង្ហាញក្នុង Excel តាមជួរឈរ។"
        ]

        for tip in tips:
            lbl_tip = QLabel(tip)
            lbl_tip.setWordWrap(True)
            lbl_tip.setStyleSheet("font-size: 15px;")
            t_layout.addWidget(lbl_tip)

        # --- បញ្ចូលគ្រប់ផ្នែកទៅក្នុង Layout ធំ ---
        scroll_layout.addWidget(shortcut_group)
        scroll_layout.addWidget(btn_desc_group)
        scroll_layout.addWidget(tip_group)
        scroll_layout.addStretch()

        scroll.setWidget(content)
        layout.addWidget(scroll)

        # Footer
        footer = QLabel("Clinic Management System v2.0 | © 2025")
        footer.setAlignment(Qt.AlignCenter) # type: ignore
        footer.setStyleSheet("color: #485460; font-size: 12px; padding: 10px;")
        layout.addWidget(footer)

    def set_tab_order_for_patient_type(self, patient_type):
        """
        Set tab order based on patient type (Child or Adult).
        For Child: Treatment → Nutrition W/A → Nutrition W/H → Refer To → Service → Remark → ADD
        For Adult: Treatment → Refer To → Service → Remark → ADD (skip Nutrition)
        """
        # FIRST: Update age field focus based on current age category
        self.on_age_category_changed(self.age_cat.currentText())
        
        # SECOND: Always skip auto-generated fields in tab order
        self.serial_no.setFocusPolicy(Qt.NoFocus)  # type: ignore
        self.area_val.setFocusPolicy(Qt.NoFocus)  # type: ignore
        self.disease_val.setFocusPolicy(Qt.NoFocus)  # type: ignore

        # THIRD: Set focus policy to exclude fields from tab order
        if patient_type == "Child":
            # For Children: IMCI fields should be skipped in tab order
            self.imci_val.setFocusPolicy(Qt.NoFocus)  # type: ignore
        else:
            # For Adults: IMCI and Nutrition fields should be skipped
            self.imci_val.setFocusPolicy(Qt.NoFocus)  # type: ignore
            self.nut_wa.setFocusPolicy(Qt.NoFocus)  # type: ignore
            self.nut_wh.setFocusPolicy(Qt.NoFocus)  # type: ignore
        
        # FOURTH: Set tab order
        if patient_type == "Child":
            # Full tab order for Children (includes Nutrition fields)
            tab_order = [
                self.date, self.card_id, self.name, self.guardian,
                self.age_cat, self.age_val, self.age_months, self.age_days,
                self.sex, self.address, self.phone, self.area_cat,
                self.pregnant, self.ref_from, self.disease_cat,
                self.symptoms, self.paraclinical, self.diagnosis,
                self.treatment, self.nut_wa, self.nut_wh,
                self.ref_to, self.service, self.remark, self.btn_add
            ]
        else:
            # Simplified tab order for Adults (skip Nutrition and IMCI fields)
            tab_order = [
                self.date, self.card_id, self.name, self.guardian,
                self.age_cat, self.age_val, self.age_months, self.age_days,
                self.sex, self.address, self.phone, self.area_cat,
                self.pregnant, self.ref_from, self.disease_cat,
                self.symptoms, self.paraclinical, self.diagnosis,
                self.treatment,
                self.ref_to, self.service, self.remark, self.btn_add
            ]
        
        # Set tab order sequentially
        for i in range(len(tab_order) - 1):
            self._safe_set_tab_order(tab_order[i], tab_order[i+1])

    def _safe_set_tab_order(self, first, second):
        """Set tab order only if both widgets belong to the same top-level window."""
        if first is None or second is None:
            return
        try:
            if first.window() is not second.window():
                return
        except Exception:
            return
        self.setTabOrder(first, second)

    def on_entry_tab_changed(self, index):
        if getattr(self, "_loading_table", False) or getattr(self, "_populating_form", False):
            return

        new_type = "Child" if index == 0 else "Adult"
        kh_type = "កុមារ (Child)" if index == 0 else "មនុស្សចាស់ (Adult)"
        preserve_inputs = self._preserve_entry_on_tab_switch

        # Check if any important fields are filled before showing confirmation
        is_dirty = any([self.name.text().strip(), self.card_id.text().strip(), self.diagnosis.currentText().strip()])
        
        if hasattr(self, 'current_patient_type') and self.current_patient_type != new_type:
            if is_dirty and not preserve_inputs:
                msg = f"តើអ្នកពិតជាចង់ប្តូរទៅកាន់របៀប '{kh_type}' មែនទេ?\n(ទិន្នន័យដែលកំពុងវាយនឹងត្រូវសម្អាត)"
                reply = QMessageBox.question(self, 'Confirm Mode Switch', msg, QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
                
                if reply == QMessageBox.No:
                    self.type_tabs.blockSignals(True)
                    self.type_tabs.setCurrentIndex(0 if self.current_patient_type == "Child" else 1)
                    self.type_tabs.blockSignals(False)
                    return

        self.current_patient_type = new_type
        self.statusBar.showMessage(f"Mode: {new_type} Patient Entry")

        # Show/Hide IMCI field based on patient type (IMCI is only for Children)
        if new_type == "Adult":
            self.imci_container.setVisible(False)
            self.imci_val.setFocusPolicy(Qt.NoFocus)  # type: ignore  # Skip in Tab
            self.nut_wa.setVisible(False)
            self.nut_wa.setFocusPolicy(Qt.NoFocus)  # type: ignore  # Skip in Tab
            self.nut_wh.setVisible(False)
            self.nut_wh.setFocusPolicy(Qt.NoFocus)  # type: ignore  # Skip in Tab
            # Show pregnant field for Adults only
            self.pregnant.setVisible(True)
            self.pregnant.setFocusPolicy(Qt.StrongFocus)  # type: ignore  # Allow focus
        else:
            self.imci_container.setVisible(True)
            self.imci_val.setFocusPolicy(Qt.NoFocus)  # type: ignore  # Read-only, skip in Tab
            self.nut_wa.setVisible(True)
            self.nut_wa.setFocusPolicy(Qt.StrongFocus)  # type: ignore
            self.nut_wh.setVisible(True)
            self.nut_wh.setFocusPolicy(Qt.StrongFocus)  # type: ignore
            # Hide pregnant field for Children (not applicable)
            self.pregnant.setVisible(False)
            self.pregnant.clear()
            self.pregnant.setFocusPolicy(Qt.NoFocus)  # type: ignore  # Skip in Tab

        # Update the age category options for the selected tab
        self.update_age_category_options()

        # Update Tab Order based on patient type (this also updates age field focus)
        self.set_tab_order_for_patient_type(new_type)

        if preserve_inputs:
            self.update_next_serial_no()
        else:
            # Clear inputs to avoid mixing data between types
            # But keep the date as it is usually the same for batch entry
            current_date = self.date.text()
            self.clear_inputs()
            self.date.setText(current_date)

        # Refresh the view to show only patients of the current type
        self.view()
        self._preserve_entry_on_tab_switch = False
        self._update_add_button_state()

    def _infer_patient_type_from_age_category(self, age_category):
        """Infer whether an age category belongs to Child or Adult tabs."""
        normalized_age = str(age_category or "").strip()
        if not normalized_age:
            return None

        child_age_groups = ['០-២៩ថ្ងៃ', '២៩ថ្ងៃ-១១ខែ', '១-៤ឆ្នាំ', '0-28', '0-29', '29-11', '1-4']
        adult_age_groups = ['៥-១៤ឆ្នាំ', '១៥-២៤ឆ្នាំ', '២៥-៤៩ឆ្នាំ', '៥០-៦៤ឆ្នាំ', '>=៦៥ឆ្នាំ',
                            '5-14', '15-24', '25-49', '50-64', '>=65']

        for child_group in child_age_groups:
            if child_group in normalized_age:
                return "Child"

        for adult_group in adult_age_groups:
            if adult_group in normalized_age:
                return "Adult"

        return None

    def _switch_entry_tab_preserving_inputs(self, target_type):
        """Switch Child/Adult tab without clearing the form."""
        self._preserve_entry_on_tab_switch = True
        self.type_tabs.setCurrentIndex(0 if target_type == "Child" else 1)

    def update_age_category_options(self):
        """Populate the age category dropdown with Child or Adult groups."""
        child_age_groups = [
            "",
            "០-២៩ថ្ងៃ",
            "២៩ថ្ងៃ-១១ខែ",
            "១-៤ឆ្នាំ"
        ]
        adult_age_groups = [
            "",
            "៥-១៤ឆ្នាំ",
            "១៥-២៤ឆ្នាំ",
            "២៥-៤៩ឆ្នាំ",
            "៥០-៦៤ឆ្នាំ",
            ">=៦៥ឆ្នាំ"
        ]

        target_type = getattr(self, 'current_patient_type', None)
        if target_type is None:
            target_type = "Adult" if self.type_tabs.currentIndex() == 1 else "Child"

        options = adult_age_groups if target_type == "Adult" else child_age_groups
        current_value = self.age_cat.currentText().strip() if hasattr(self.age_cat, 'currentText') else ""

        self.age_cat.blockSignals(True)
        self.age_cat.clear()
        self.age_cat.addItems(options)
        if current_value and current_value in options:
            self.age_cat.setCurrentText(current_value)
        else:
            self.age_cat.setCurrentIndex(0)
        self.age_cat.blockSignals(False)

        # Update helper text based on selected age category
        self.on_age_category_changed(self.age_cat.currentText())

    def _prompt_patient_type_for_age_mismatch(self):
        """Ask the user where to save when age range does not match the current tab."""
        inferred_type = self._infer_patient_type_from_age_category(self.age_cat.currentText())
        if not inferred_type or inferred_type == self.current_patient_type:
            return True

        current_label = "កុមារ" if self.current_patient_type == "Child" else "មនុស្សចាស់"
        target_label = "កុមារ" if inferred_type == "Child" else "មនុស្សចាស់"

        msg_box = QMessageBox(self)
        msg_box.setIcon(QMessageBox.Warning)
        msg_box.setWindowTitle("ជ្រើសរើសប្រភេទអ្នកជំងឺ")
        msg_box.setText(
            f"អាយុដែលបានជ្រើសស្ថិតក្នុងក្រុម '{target_label}' ប៉ុន្តែអ្នកកំពុងស្ថិតនៅ Tab '{current_label}'។\n\n"
            "សូមជ្រើសរើសថាចង់បញ្ចូលទៅ Tab មួយណា។"
        )

        switch_btn = msg_box.addButton(f"ប្ដូរទៅ {target_label}", QMessageBox.AcceptRole)
        cancel_btn = msg_box.addButton("បោះបង់", QMessageBox.RejectRole)
        msg_box.setDefaultButton(switch_btn)
        msg_box.exec_()

        clicked = msg_box.clickedButton()
        if clicked == cancel_btn:
            return False
        if clicked == switch_btn:
            self._switch_entry_tab_preserving_inputs(inferred_type)
            self.statusBar.showMessage(f"បានប្ដូរទៅ Tab {target_label} ដើម្បីបញ្ចូលទិន្នន័យបន្ត", 4000)
            return True

        return False

    def auto_switch_patient_type_by_age(self, age_category):
        """
        Automatically switch patient type based on selected age category.
        Child age groups: 0ថ្ងៃ-28ថ្ងៃ, 29ថ្ងៃ-១១ខែ, ១ឆ្នាំ-៤ឆ្នាំ
        Adult age groups: 5ឆ្នាំ-14ឆ្នាំ, 15ឆ្នាំ-24ឆ្នាំ, 25ឆ្នាំ-៤៩ឆ្នាំ, 50ឆ្នាំ-64ឆ្នាំ, >65ឆ្នាំ
        """
        child_age_groups = ['0ថ្ងៃ-28ថ្ងៃ', '29ថ្ងៃ-១១ខែ', '១ឆ្នាំ-៤ឆ្នាំ', '0-28', '29-11', '1-4']
        adult_age_groups = ['5ឆ្នាំ-14ឆ្នាំ', '15ឆ្នាំ-24ឆ្នាំ', '25ឆ្នាំ-៤៩ឆ្នាំ', '50ឆ្នាំ-64ឆ្នាំ', '>65ឆ្នាំ',
                           '5-14', '15-24', '25-49', '50-64', '>65']
        
        # Determine target patient type based on age category
        target_type = None
        for child_group in child_age_groups:
            if child_group in age_category:
                target_type = "Child"
                break
        
        if target_type is None:
            for adult_group in adult_age_groups:
                if adult_group in age_category:
                    target_type = "Adult"
                    break
        
        if target_type and target_type != self.current_patient_type:
            # Switch to the appropriate tab
            tab_index = 0 if target_type == "Child" else 1
            self.type_tabs.setCurrentIndex(tab_index)
            self.statusBar.showMessage(f"ប្តូរទៅកាន់ {target_type} ដោយស្វ័យប្រវត្តិ (Age: {age_category})", 3000)

    def auto_switch_to_child_by_imci(self, imci_text):
        """
        Automatically switch to Child tab when IMCI data is entered.
        IMCI is only for children, so if user enters IMCI data while on Adult tab, switch to Child.
        """
        if imci_text.strip() and self.current_patient_type == "Adult":
            self.type_tabs.setCurrentIndex(0)  # Switch to Child tab
            self.statusBar.showMessage("ប្តូរទៅកាន់ កុមារ ដោយស្វ័យប្រវត្តិ (IMCI សម្រាប់កុមារប៉ុណ្ណោះ)", 3000)

    def on_age_category_changed(self, age_category):
        """
        Update the age value helper text based on selected age category.
        The category itself is selected from a fixed list, and the next field remains editable.
        """
        age_category = str(age_category or "").strip()
        self.age_val.clear()
        self.age_months.clear()
        self.age_days.clear()
        self.age_val.setVisible(False)
        self.age_val.setFocusPolicy(Qt.NoFocus)  # type: ignore
        self.age_months.setVisible(False)
        self.age_days.setVisible(False)
        self.age_months.setFocusPolicy(Qt.NoFocus)  # type: ignore
        self.age_days.setFocusPolicy(Qt.NoFocus)  # type: ignore

        if not age_category or age_category == "ជ្រើសរើសក្រុមអាយុ":
            self.age_val.setVisible(True)
            self.age_val.setFocusPolicy(Qt.StrongFocus)  # type: ignore
            self.age_val.setPlaceholderText("ឧ. 25ឆ្នាំ")
            self._update_add_button_state()
            return

        if '0ថ្ងៃ-28ថ្ងៃ' in age_category or '0-28' in age_category or '0-29' in age_category or '០-២៩ថ្ងៃ' in age_category:
            self.age_days.setVisible(True)
            self.age_days.setFocusPolicy(Qt.StrongFocus)  # type: ignore
            self.age_days.setPlaceholderText("ឧ. ១៥ថ្ងៃ")
        elif '29ថ្ងៃ-១១ខែ' in age_category or '២៩ថ្ងៃ-១១ខែ' in age_category or '29-11' in age_category:
            self.age_months.setVisible(True)
            self.age_days.setVisible(True)
            self.age_months.setFocusPolicy(Qt.StrongFocus)  # type: ignore
            self.age_days.setFocusPolicy(Qt.StrongFocus)  # type: ignore
            self.age_months.setPlaceholderText("ខែ")
            self.age_days.setPlaceholderText("ថ្ងៃ")
        else:
            self.age_val.setVisible(True)
            self.age_val.setFocusPolicy(Qt.StrongFocus)  # type: ignore
            self.age_val.setPlaceholderText("ឧ. ២ឆ្នាំ")
        self._update_add_button_state()

    def init_entry_tab(self):
        # Create the main container for Data Entry
        entry_container = QWidget()
        layout = QVBoxLayout(entry_container)
        self.tabs.addTab(entry_container, "បញ្ចូលទិន្នន័យ")

        # Date Input with Calendar Shortcut
        self.date_container = QWidget()
        date_box_layout = QHBoxLayout(self.date_container)
        date_box_layout.setContentsMargins(0,0,0,0)
        date_box_layout.setSpacing(2)
        
        self.date = QLineEdit()
        self.date.setReadOnly(True)
        self.date.setPlaceholderText("ជ្រើសរើសកាលបរិច្ឆេទ")
        self.date.setToolTip("ចុចដើម្បីជ្រើសរើសកាលបរិច្ឆេទពីប្រតិទិន")
        self.date.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]
        self.date.mousePressEvent = lambda event: self.pick_date()  # type: ignore[method-assign]
        self.date.setStyleSheet("background-color: #353b48; color: #ffffff; font-weight: bold;")
        
        self.btn_cal = QPushButton("📅")
        self.btn_cal.setFixedWidth(35)
        self.btn_cal.setCursor(Qt.PointingHandCursor) # type: ignore
        self.btn_cal.setFocusPolicy(Qt.NoFocus) # type: ignore
        self.btn_cal.setStyleSheet("background-color: #34495e; border-radius: 5px; font-size: 16px;")
        self.btn_cal.clicked.connect(self.pick_date)
        
        date_box_layout.addWidget(self.date)
        date_box_layout.addWidget(self.btn_cal)

        self.serial_no = QLineEdit()
        self.serial_no.setReadOnly(True) # បិទមិនឱ្យវាយបញ្ចូល
        self.serial_no.setStyleSheet("background-color: #353b48; color: #d2dae2; font-weight: bold;")
        self.serial_no.setFocusPolicy(Qt.NoFocus) # type: ignore # រំលងការ Tab ចូលទៅក្នុងប្រអប់ដែលចេញលេខអូតូ
        self.serial_no.setPlaceholderText("Auto")
        
        # --- Type Selection Tabs ---
        self.type_tabs = QTabWidget()
        self.type_tabs.addTab(QWidget(), "កុមារ")
        self.type_tabs.addTab(QWidget(), "មនុស្សចាស់")
        self.type_tabs.setFixedHeight(40) # Small height just for the headers
        self.type_tabs.currentChanged.connect(self.on_entry_tab_changed)
        # We don't put content in these tabs, just use them as a switch
        
        layout.addWidget(self.type_tabs)

        self.card_id = QLineEdit()
        self.name = QLineEdit()
        self.guardian = QLineEdit()
        
        # Age Composite Input (Dropdown + Text + Months + Days for infants)
        self.age_container = QWidget()
        age_layout = QHBoxLayout(self.age_container)
        age_layout.setContentsMargins(0,0,0,0)
        age_layout.setSpacing(5)
        
        self.age_cat = QComboBox()
        self.age_cat.setFixedWidth(160)
        self.age_cat.setEditable(False)
        self.age_cat.setInsertPolicy(QComboBox.NoInsert)
        self.age_cat.setToolTip("ជ្រើសរើសក្រុមអាយុ")
        
        self.age_val = QLineEdit()
        self.age_val.setPlaceholderText("ឧ. 25ឆ្នាំ")
        self.age_val.setFixedWidth(120)
        
        # Additional fields for infants (months and days)
        self.age_months = QLineEdit()
        self.age_months.setPlaceholderText("ខែ")
        self.age_months.setFixedWidth(50)
        self.age_months.setVisible(False)
        
        self.age_days = QLineEdit()
        self.age_days.setPlaceholderText("ថ្ងៃ")
        self.age_days.setFixedWidth(50)
        self.age_days.setVisible(False)
        
        age_layout.addWidget(self.age_cat)
        age_layout.addWidget(self.age_val)
        age_layout.addWidget(self.age_months)
        age_layout.addWidget(self.age_days)

        # Connect age category change to update the age input helper text
        self.age_cat.currentTextChanged.connect(self.on_age_category_changed)
        self.update_age_category_options()

        self.sex = QComboBox()
        self.sex.setEditable(False)
        self._set_combo_items_blank_first(self.sex, ["ស្រី", "ប្រុស"])
        # Auto-hide pregnant field when sex is male
        self.sex.currentTextChanged.connect(self.on_sex_changed)

        # Area Composite Input (Dropdown for category + Auto counter)
        self.area_container = QWidget()
        area_layout = QHBoxLayout(self.area_container)
        area_layout.setContentsMargins(0,0,0,0)

        # Fixed QComboBox for Area: only the configured choices are allowed.
        self.area_cat = QComboBox()
        self.area_cat.setEditable(False)
        self.area_cat.setInsertPolicy(QComboBox.NoInsert)
        self.area_cat.setStyleSheet("background-color: #353b48; color: #d2dae2; font-weight: bold; padding: 5px;")

        self.area_val = QLineEdit()
        self.area_val.setReadOnly(True)
        self.area_val.setPlaceholderText("Auto")
        self.area_val.setStyleSheet("background-color: #353b48; color: #d2dae2; font-weight: bold;")
        self.area_val.setFocusPolicy(Qt.NoFocus)  # type: ignore

        area_layout.addWidget(self.area_cat)
        area_layout.addWidget(self.area_val)

        # Auto-generate Area counter when category is selected
        self.area_cat.currentTextChanged.connect(self.update_area_serial)

        # Populate Area categories, starting with a blank selection.
        self._set_combo_items_blank_first(self.area_cat, self.CAT_AREA)

        self.pregnant = QLineEdit()
        self.pregnant.setPlaceholderText("សប្ដាហ៍")
        # Initially hidden for Child patients (default patient type)
        self.pregnant.setVisible(False)
        self.pregnant.setFocusPolicy(Qt.NoFocus)  # type: ignore  # Skip by default for Child

        self.address = QLineEdit()
        self.phone = QLineEdit()
        self.ref_from = QLineEdit()
        
        # Disease Case Composite Input (Editable with autocomplete + Auto counter)
        self.disease_container = QWidget()
        dis_layout = QHBoxLayout(self.disease_container)
        dis_layout.setContentsMargins(0,0,0,0)

        # Fixed QComboBox for Disease to prevent invalid case categories.
        self.disease_cat = QComboBox()
        self.disease_cat.setEditable(False)
        self.disease_cat.setStyleSheet("background-color: #353b48; color: #d2dae2; font-weight: bold; padding: 5px;")

        self.disease_val = QLineEdit()
        self.disease_val.setReadOnly(True)
        self.disease_val.setPlaceholderText("Auto")
        self.disease_val.setStyleSheet("background-color: #353b48; color: #d2dae2; font-weight: bold;")
        self.disease_val.setFocusPolicy(Qt.NoFocus)  # type: ignore

        dis_layout.addWidget(self.disease_cat)
        dis_layout.addWidget(self.disease_val)

        # Auto-generate Disease counter when category is selected/typed
        self.disease_cat.currentTextChanged.connect(self.update_disease_serial)

        # Populate Disease categories, starting with a blank selection.
        self._set_combo_items_blank_first(self.disease_cat, self.CAT_DISEASE)

        self.symptoms = QLineEdit()
        self.paraclinical = QLineEdit()
        self.diagnosis = QComboBox()
        self.diagnosis.setEditable(False)
        self.diagnosis.setStyleSheet("background-color: #353b48; color: #d2dae2; font-weight: bold; padding: 5px;")
        if hasattr(self, 'CAT_DIAGNOSIS') and self.CAT_DIAGNOSIS:
            self._set_combo_items_blank_first(self.diagnosis, self.CAT_DIAGNOSIS)
            self.diagnosis.setPlaceholderText("ជ្រើសរើសរោគវិនិច្ឆ័យ")

        self.treatment = QLineEdit()
        
        # IMCI Composite Input (Auto for Children only - No manual input needed)
        self.imci_container = QWidget()
        imci_layout = QHBoxLayout(self.imci_container)
        imci_layout.setContentsMargins(0,0,0,0)
        
        self.imci_label = QLabel("IMCI:")
        self.imci_label.setStyleSheet("font-weight: bold; color: #00cec9;")
        
        self.imci_val = QLineEdit()
        self.imci_val.setReadOnly(True)  # Read-only (auto-generated)
        self.imci_val.setPlaceholderText("ស្វ័យប្រវត្តិ (សម្រាប់កុមារ)")
        self.imci_val.setFocusPolicy(Qt.NoFocus) # type: ignore  # Skip in Tab order
        self.imci_val.setStyleSheet("background-color: #353b48; color: #05c46b; font-weight: bold;")
        
        imci_layout.addWidget(self.imci_label)
        imci_layout.addWidget(self.imci_val)
        imci_layout.addStretch()

        # NEW Nutrition fields
        self.nut_wa = QLineEdit() # ទម្ងន់/អាយុ
        self.nut_wh = QLineEdit() # ទម្ងន់/កំពស់

        # IMCI and Nutrition fields are only for Children - hide/show based on default patient type
        self.imci_container.setVisible(True)  # Default is Child
        self.nut_wa.setVisible(True)
        self.nut_wh.setVisible(True)

        self.ref_to = QLineEdit()
        self.service = QComboBox()
        self.service.setEditable(False)
        self._set_combo_items_blank_first(
            self.service,
            ["HEF", "HEF-R", "HEF-I", "PAY", "FREE", "NSSF-A", "NSSF-7", "NSSF-8", "Other"]
        )
        self.remark = QLineEdit()

        # --- UI Refactoring with QGroupBox for better organization ---
        screen = QApplication.primaryScreen()
        available_width = screen.availableGeometry().width() if screen else 1200
        compact_entry_layout = available_width < 1400
        main_form_layout = QVBoxLayout() if compact_entry_layout else QHBoxLayout()

        # Helper function to create a group box with a grid layout
        def create_group_box(title):
            group_box = QGroupBox(title)
            grid = QGridLayout(group_box)
            grid.setContentsMargins(8, 18, 8, 8)
            grid.setHorizontalSpacing(6)
            grid.setVerticalSpacing(6)
            grid.setColumnMinimumWidth(0, 118)
            return group_box, grid

        def add_form_row(grid, row, label_text, widget):
            label = QLabel(label_text)
            label.setMinimumWidth(118)
            label.setMinimumHeight(32)
            label.setWordWrap(True)
            label.setAlignment(Qt.AlignVCenter | Qt.AlignLeft)  # type: ignore[attr-defined]
            widget.setMinimumHeight(32)
            grid.setRowMinimumHeight(row, 36)
            grid.addWidget(label, row, 0)
            grid.addWidget(widget, row, 1)
            return row + 1

        # Group 1: Personal Info
        personal_group, personal_grid = create_group_box("ព័ត៌មានផ្ទាល់ខ្លួនអ្នកជំងឺ")
        personal_info_fields = [
            ("កាលបរិច្ឆេទ", self.date_container),
            ("លេខរៀង", self.serial_no),
            ("លេខប័ណ្ណ", self.card_id),
            ("ឈ្មោះអ្នកជំងឺ", self.name),
            ("ឈ្មោះអាណាព្យាបាល", self.guardian),
            ("អាយុ", self.age_container),
            ("ភេទ", self.sex),
            ("អាស័យដ្ឋាន", self.address),
            ("លេខទូរស័ព្ទ", self.phone),
        ]
        for i, (lbl, wid) in enumerate(personal_info_fields):
            add_form_row(personal_grid, i, lbl, wid)

        # Group 2: Medical Info
        medical_group, medical_grid = create_group_box("ព័ត៌មានវេជ្ជសាស្រ្ត")
        medical_info_fields = [
            ("តំបន់ទទួលខុសត្រូវ", self.area_container),
            ("មានផ្ទៃពោះ(សប្ដាហ៍)", self.pregnant),
            ("បញ្ជូនមកពី", self.ref_from),
            ("ករណីជំងឺ", self.disease_container),
            ("រោគសញ្ញា", self.symptoms),
            ("អមវេជ្ជសាស្រ្ត", self.paraclinical),
            ("រោគវិនិឆ្ឆ័យ", self.diagnosis),
            ("ការព្យាបាល", self.treatment),
            ("បានប្រើប្រាស់ទំរង់ IMCI", self.imci_container),
            ("ទម្ងន់/អាយុ", self.nut_wa),
            ("ទម្ងន់/កំពស់", self.nut_wh),
        ]
        for i, (lbl, wid) in enumerate(medical_info_fields):
            add_form_row(medical_grid, i, lbl, wid)

        # Group 3: Administrative Info
        admin_group, admin_grid = create_group_box("ព័ត៌មានរដ្ឋបាល និងផ្សេងៗ")
        admin_info_fields = [
            ("បញ្ជូនទៅ", self.ref_to),
            ("ប្រភេទនៃការបង់ថ្លៃសេវា", self.service),
            ("សម្គាល់", self.remark),
        ]
        for i, (label, widget) in enumerate(admin_info_fields):
            add_form_row(admin_grid, i, label, widget)
        
        # Add an empty spacer to push the admin group up
        admin_grid.setRowStretch(len(admin_info_fields), 1)

        # Add groups to the main form layout
        if compact_entry_layout:
            for group in (personal_group, medical_group, admin_group):
                group.setMinimumWidth(520)
                main_form_layout.addWidget(group)
        else:
            personal_group.setMinimumWidth(380)
            medical_group.setMinimumWidth(430)
            admin_group.setMinimumWidth(360)
            main_form_layout.addWidget(personal_group, 1)
            main_form_layout.addWidget(medical_group, 1)
            main_form_layout.addWidget(admin_group, 1)

        # --- បង្កើតប៊ូតុងនៅទីនេះ (មុន Tab Order) ---
        btn_grid = QGridLayout()
        btn_grid.setSpacing(4) # កំណត់ចន្លោះរវាងប៊ូតុងឱ្យតូច
        # --- រៀបចំប៊ូតុងថ្មី (Primary Actions & Tools Menu) ---
        action_layout = QHBoxLayout()
        action_layout.setSpacing(10)

        # ១. ប៊ូតុងសំខាន់ៗដែលប្រើញឹកញាប់ (Primary)
        primary_btns = [
            ("បន្ថែម", "#05c46b", self.add, "Alt+A"),
            ("កែប្រែ", "#0fbcf9", self.update_patient, "Alt+E"),
            ("លុប", "#ff3f34", self.delete, "Alt+D"),
            ("សម្អាត", "#7f8c8d", self.clear_inputs, "Alt+L"),
            ("ប័ណ្ណ", "#e84393", self.print_patient_card, "Ctrl+P"),
        ]

        for text, color, handler, shortcut in primary_btns:
            btn = self._create_button(text, color, handler, shortcut)
            if handler == self.add:
                self.btn_add = btn
            action_layout.addWidget(btn)

        # ២. បង្កើតប៊ូតុង Tools ដែលមាន Menu ជា List (Secondary)
        self.tools_btn = QPushButton("🛠 ឧបករណ៍ និងទិន្នន័យ ▾")
        self.tools_btn.setCursor(Qt.PointingHandCursor) # type: ignore
        self.tools_btn.setStyleSheet("background: #34495e; color: white; padding: 6px 15px; border-radius: 5px; font-weight: bold; min-height: 25px;")
        
        self.tools_menu = QMenu(self)
        secondary_actions = [
            ("🔍 ស្វែងរកទិន្នន័យ (Search)", self.search, "Alt+F"),
            ("📋 បង្ហាញទាំងអស់ (View All)", self.view, "F5"),
            ("📊 ពិនិត្យស្ថិតិអ្នកជំងឺ", self.show_statistics, None),
            ("👤 ពិនិត្យទិន្នន័យតាម User / Excel", self.open_user_audit_dialog, None),
            ("📗 ទាញទិន្នន័យចេញ Excel", self.export_excel, "Alt+X"),
            ("📥 នាំចូលទិន្នន័យពី Excel", self.import_excel, None),
            ("📕 ទាញទិន្នន័យចេញ PDF", self.export_pdf, None),
            ("🔗 បញ្ចូល Database (Merge)", self.merge_database, None),
            ("📤 Share Database ទៅ Telegram", self.share_to_telegram, None),
            ("💾 Backup Database", self.backup_database, None),
            ("🔄 Restore Database", self.restore_database, None),
            ("🔐 License / Register", self.open_license_status, None),
            ("🎨 ការកំណត់ និងប្តូរ Themes", self.open_settings, None),
            ("🆙 ពិនិត្យមើលការ Update", self.check_for_updates, None),
            ("📜 ប្រវត្តិប្រើប្រាស់ (History)", self.show_login_history, None),
            ("🔑 ប្តូរពាក្យសម្ងាត់", self.open_change_password, None),
            ("🧹 Reset ទិន្នន័យចោល", self.reset_database, None),
            ("🚪 ចាកចេញពីកម្មវិធី (Logout)", self.logout, "Ctrl+Q"),
        ]

        for text, handler, shortcut in secondary_actions:
            action = QAction(text, self)
            if shortcut: action.setShortcut(QKeySequence(shortcut))
            action.triggered.connect(handler)
            self.tools_menu.addAction(action)
        
        self.tools_btn.setMenu(self.tools_menu)
        action_layout.addWidget(self.tools_btn)

        self.cloud_btn = QPushButton("☁ Cloud Sync ▾")
        self.cloud_btn.setCursor(Qt.PointingHandCursor)  # type: ignore[attr-defined]
        self.cloud_btn.setStyleSheet("""
            QPushButton {
                background: #0984e3;
                color: white;
                padding: 6px 15px;
                border-radius: 5px;
                font-weight: bold;
                min-height: 25px;
            }
            QPushButton:hover { background: #1e90ff; }
            QPushButton::menu-indicator { image: none; width: 0px; }
        """)
        self.cloud_menu = QMenu(self)
        cloud_actions = [
            ("☁️ ទាញពី Cloud", lambda checked=False: self._run_cloud_helper_action("sync_data_initial", refresh_after=True)),
            ("⬆️ ផ្ញើទៅ Cloud", lambda checked=False: self._run_cloud_helper_action("upload_to_cloud")),
            ("✏️ កែទិន្នន័យ Cloud", lambda checked=False: self._run_cloud_helper_action("edit_uploaded_cloud_data")),
            ("🗑️ លុបទិន្នន័យ Cloud", lambda checked=False: self._run_cloud_helper_action("delete_uploaded_cloud_data")),
            ("⚙️ កំណត់ Cloud Repository", lambda checked=False: self._run_cloud_helper_action("show_github_setup")),
            ("🔑 កំណត់ GitHub Token", lambda checked=False: self._run_cloud_helper_action("show_github_token_setup")),
            ("❓ ជំនួយ Cloud", lambda checked=False: self._run_cloud_helper_action("show_cloud_sync_help")),
        ]
        for text, handler in cloud_actions:
            action = QAction(text, self)
            action.triggered.connect(handler)
            self.cloud_menu.addAction(action)
        self.cloud_btn.setMenu(self.cloud_menu)
        action_layout.addWidget(self.cloud_btn)
        action_layout.addStretch() # រុញប៊ូតុងឱ្យនៅខាងឆ្វេង

        # Set Tab Order dynamically based on patient type
        self.set_tab_order_for_patient_type(self.current_patient_type)
        self._connect_required_field_signals()
        self._update_add_button_state()

        # Enable Enter key to move focus to next field (Speed up entry)
        def focus_next():
            self.focusNextChild()

        # Note: Enter key handler is set globally, not per widget

        # --- Search Bar for Table ---
        search_layout = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("🔍 ស្វែងរកតាម ឈ្មោះ, លេខរៀង, លេខប័ណ្ណ ឬ លេខទូរស័ព្ទ...")
        self.search_input.setStyleSheet("""
            background: #2f3640; border: 1px solid #0fbcf9; 
            padding: 8px; border-radius: 5px; color: white;
        """)
        self.search_input.textChanged.connect(self.search) # ស្វែងរកអូតូពេលវាយអក្សរ
        search_layout.addWidget(self.search_input)

        self.table = QTableWidget()
        self.table.setStyleSheet("""
            QTableWidget{
                background:#1e272e;
                color:white;
                border-radius:10px;
                gridline-color:#485460;
                font-size: 10px;
                font-size: 12px;
                font-size: 14px;
            }
            QHeaderView::section{
                background:#0fbcf9;
                color:black;
                padding:4px;
                border:none;
            }
        """)
        # Keep sorting disabled on the main entry table. With mixed custom table
        # items and hidden continuation rows, Qt can crash during mouse/table
        # interaction in some PyQt builds.
        self.table.setSortingEnabled(False)
        
        # Use a splitter to allow resizing between Form and Table
        self.splitter = QSplitter(Qt.Vertical) # type: ignore
        
        top_widget = QWidget()
        top_layout = QVBoxLayout(top_widget)
        top_layout.addLayout(main_form_layout)
        top_layout.addLayout(action_layout)
        top_layout.setContentsMargins(0, 0, 0, 0)

        top_scroll = QScrollArea()
        top_scroll.setWidgetResizable(True)
        top_scroll.setFrameShape(QFrame.NoFrame)
        top_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)  # type: ignore
        top_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)  # type: ignore
        top_scroll.setWidget(top_widget)
        
        bottom_widget = QWidget()
        bottom_layout = QVBoxLayout(bottom_widget)
        bottom_layout.addLayout(search_layout)
        bottom_layout.addWidget(self.table)
        
        self.splitter.addWidget(top_scroll)
        self.splitter.addWidget(bottom_widget)
        self.splitter.setStretchFactor(1, 2)
        self.splitter.setSizes([360, 390])
        
        layout.addWidget(self.splitter)
        
        self.table.itemSelectionChanged.connect(self.populate_form_from_table)


    def _load_or_create_settings(self):
        """Loads settings from settings.ini and backfills any missing defaults."""
        default_categories = {
            'age': '០-២៩ថ្ងៃ, ២៩ថ្ងៃ-១១ខែ, ១-៤ឆ្នាំ, ៥-១៤ឆ្នាំ, ១៥-២៤ឆ្នាំ, ២៥-៤៩ឆ្នាំ, ៥០-៦៤ឆ្នាំ, >=៦៥ឆ្នាំ',
            'sex': 'ស្រី, ប្រុស',
            'area': 'ក, ខ, គ',
            'disease': 'ថ្មី, ចាស់',
            'imci': '',
            'nutrition': 'ទម្ងន់/អាយុ, ទម្ងន់/កំពស់',
            'service': 'HEF, HEF-R, HEF-I, PAY, FREE, NSSF-A, NSSF-7, NSSF-8, Other',
            'diagnosis': 'ជំងឺរលាកផ្លូវដង្ហើមលើ, ជំងឺរលាកផ្លូវដង្ហើមក្រោម, របេងកូនកណ្ដុរ, ជំងឺរលាកក្រពះស្រួចស្រាវ, រលាកក្រពះរ៉ាំរ៉ៃ, រាកគ្មានខ្សោះជាតិទឺក, រាកមួល, របេងសួត, ធ្លាក់ស, សើរស្បែកផ្សេងៗ, SAM, ជំងឺភ្នែកក្រហម, ក្អកលើស១៤ថ្ងៃ, ស្លេកស្លាំង, ខ្វះជាតិស្ករ, របួសផ្សេងៗក្រៅពីគ្រោះថ្នាក់ចរាចរណ៍, រដូវមិនទៀង, រលាកទ្វាមាស, របួសគ្រោះថ្នាក់ចរាចរណ៍, រលាកសន្លាក់គ្មានខ្ទុះ, ជំងឺពងបែកដៃជើង មាត់, រលាកមាត់ស្បួន, ពុលអាហារ, រលាកអញ្ជាញធ្មេញស្រួចស្រាវ, MAM, រលាកតម្រងនោម, ឈឺចាប់ពេលមករដូវ, ជំងឺទូរទៅផ្សេងៗ',
            'medicine': 'Paracetamol, Amoxicillin, Vitamin C, Ibuprofen, Omeprazole, Ciprofloxacin, Metronidazole, Ceftriaxone, Salbutamol, Domperidone',
            'theme': 'Modern Dark',
            'khmer_font': 'Khmer OS Battambang',
            'cloud_sync_repo_url': 'https://github.com/saratboy1988-a11y/Clinic-Cloud-Sync.git',
        }

        if not os.path.exists(self.settings_file):
            print(f"'{self.settings_file}' not found. Creating with default values.")
            self.config['CATEGORIES'] = default_categories.copy()
            # Use AppData for writable settings to avoid Program Files permission issues
            writable_settings_file = get_writable_settings_file(os.path.dirname(self.settings_file))
            
            with open(writable_settings_file, 'w', encoding='utf-8') as configfile:
                self.config.write(configfile)
            self.settings_file = writable_settings_file

        self.config.read(self.settings_file, encoding='utf-8-sig')

        if 'CATEGORIES' not in self.config:
            self.config['CATEGORIES'] = {}

        updated = False
        for key, value in default_categories.items():
            if not self.config['CATEGORIES'].get(key, '').strip():
                self.config['CATEGORIES'][key] = value
                updated = True

        service_items = [
            item.strip()
            for item in self.config['CATEGORIES'].get('service', '').split(',')
            if item.strip()
        ]
        if service_items:
            normalized_services = []
            for item in service_items:
                upper_item = item.upper()
                if upper_item == 'HEFR':
                    item = 'HEF-R'
                elif upper_item == 'HEFI':
                    item = 'HEF-I'
                if item not in normalized_services:
                    normalized_services.append(item)
            required_services = ['HEF', 'HEF-R', 'HEF-I', 'PAY', 'FREE', 'NSSF-A', 'NSSF-7', 'NSSF-8', 'Other']
            deprecated_services = {'NSSF', 'បសស'}
            extra_services = [item for item in normalized_services if item not in required_services and item not in deprecated_services]
            new_service_value = ', '.join(required_services + extra_services)
            if new_service_value != self.config['CATEGORIES'].get('service', ''):
                self.config['CATEGORIES']['service'] = new_service_value
                updated = True

        if updated:
            # Use AppData for writable settings
            writable_settings_file = get_writable_settings_file(os.path.dirname(self.settings_file))
            
            with open(writable_settings_file, 'w', encoding='utf-8') as configfile:
                self.config.write(configfile)
            self.settings_file = writable_settings_file

    def _get_setting(self, key, section='CATEGORIES'):
        """Safely gets a comma-separated list from the config file."""
        try:
            value = self.config.get(section, key)
            return [item.strip() for item in value.split(',') if item.strip()]
        except (configparser.NoSectionError, configparser.NoOptionError):
            return [] # Return an empty list if the key or section is missing

    def _get_configured_font(self):
        configured = self.config.get('CATEGORIES', 'khmer_font', fallback='Khmer OS Battambang').strip()
        available = available_khmer_fonts()
        return configured if configured in available else available[0]

    def apply_app_font(self, font_name=None):
        font_name = font_name or self._get_configured_font()
        app_font = QFont(font_name, 12)
        app_inst = QApplication.instance()
        if isinstance(app_inst, QApplication):
            app_inst.setFont(app_font)

        window_font = self.font()
        window_font.setFamily(font_name)
        self.setFont(window_font)
        self.current_font = font_name
        for widget in self.findChildren(QWidget):
            widget_font = widget.font()
            widget_font.setFamily(font_name)
            widget.setFont(widget_font)

    def _set_combo_items_blank_first(self, combo, items):
        """Populate a combo box with a blank first option so users must choose."""
        current_text = combo.currentText().strip() if hasattr(combo, "currentText") else ""
        combo.blockSignals(True)
        combo.clear()
        combo.addItem("")
        seen = set()
        for item in items or []:
            item = str(item).strip()
            if item and item not in seen:
                combo.addItem(item)
                seen.add(item)

        if current_text:
            index = combo.findText(current_text)
            if index >= 0:
                combo.setCurrentIndex(index)
            elif combo.isEditable():
                combo.setCurrentText(current_text)
            else:
                combo.setCurrentIndex(0)
        else:
            combo.setCurrentIndex(0)
        combo.blockSignals(False)

    def _set_combo_current_text_or_blank(self, combo, text):
        """Select existing text in a combo box, otherwise leave it blank."""
        text = str(text).strip()
        if not text:
            combo.setCurrentIndex(0)
            return
        index = combo.findText(text)
        if index >= 0:
            combo.setCurrentIndex(index)
        elif combo.isEditable():
            combo.setCurrentText(text)
        else:
            combo.setCurrentIndex(0)

    def _create_nutrition_string(self):
        parts = []
        wa_val = self.nut_wa.text().strip()
        wh_val = self.nut_wh.text().strip()
        if wa_val:
            parts.append(f"ទម្ងន់/អាយុ::{wa_val}")
        if wh_val:
            parts.append(f"ទម្ងន់/កំពស់::{wh_val}")
        return ";;".join(parts)

    def _parse_nutrition_string(self, text):
        self.nut_wa.clear()
        self.nut_wh.clear()
        if not text: return
        for part in text.split(';;'):
            if "::" in part:
                key, value = part.split('::', 1)
                if key == "ទម្ងន់/អាយុ": self.nut_wa.setText(value)
                elif key == "ទម្ងន់/កំពស់": self.nut_wh.setText(value)

    def _create_button(self, text, color, handler, shortcut=None):
        """Helper function to create and style a QPushButton."""
        btn = QPushButton(text)
        btn.setCursor(Qt.PointingHandCursor) # type: ignore
        if shortcut:
            btn.setShortcut(shortcut)
        btn.setStyleSheet(f"""
            QPushButton {{
                background: {color};
                color: white;
                padding: 4px 5px;
                border: none;
                border-radius: 4px;
                font-weight: bold;
                min-height: 20px;
            }}
            QPushButton:hover {{
                background: #57606f;
            }}
            QPushButton:disabled {{
                background: #485460;
                color: #95a5a6;
            }}
        """)
        btn.clicked.connect(handler)
        return btn

    def load_table(self, rows):
        self._loading_table = True
        old_sorting = self.table.isSortingEnabled()
        self.table.blockSignals(True)
        self.table.setSortingEnabled(False)

        # Prepare display rows handling split treatment lines
        display_rows = []
        for row in rows:
            treatment_text = str(row[db.PatientCol.TREATMENT]) if row[db.PatientCol.TREATMENT] else ""
            treat_lines = [t.strip() for t in treatment_text.split(',') if t.strip()]
            if not treat_lines: treat_lines = [""]

            for k, line in enumerate(treat_lines):
                if k == 0:
                    # Check if row has patient_type safely
                    p_type = "N/A"
                    if len(row) > db.PatientCol.TYPE:
                        p_type = row[db.PatientCol.TYPE]
                    
                    d_row = [
                        row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7],
                        row[8], row[9], row[10], row[11], row[12], row[13], row[14], row[15], 
                        row[16],
                        line, # Treatment line (17)
                        row[18], row[19], row[20], row[21], row[22],
                        p_type # Patient Type (23)
                    ]
                else:
                    d_row = [""] * (db.PatientCol.TYPE + 1)
                    d_row[0] = row[0]
                    d_row[db.PatientCol.TREATMENT] = line
                display_rows.append(d_row)

        self.table.setRowCount(len(display_rows))
        # Headers ជាភាសាខ្មែរ
        headers = ["ID", "កាលបរិច្ឆេទ", "លេខរៀង", "លេខប័ណ្ណ", "ឈ្មោះអ្នកជំងឺ", "ឈ្មោះអាណាព្យាបាល", 
                   "អាយុ", "ភេទ", "តំបន់", "មានផ្ទៃពោះ", "អាស័យដ្ឋាន", "លេខទូរស័ព្ទ", 
                   "បញ្ជូនមកពី", "ករណីជំងឺ", "រោគសញ្ញា", "អមវេជ្ជសាស្រ្ត", 
                   "រោគវិនិឆ្ឆ័យ", "ការព្យាបាល", "IMCI", "ស្ថានភាពអាហារ",
                   "បញ្ជូនទៅ", "សេវា", "សម្គាល់", "ប្រភេទ"]
        
        self.table.setColumnCount(len(headers))
        self.table.setHorizontalHeaderLabels(headers)
        
        # Hide ID column so Date is first visible column
        self.table.setColumnHidden(0, True)
        
        for i, row_data in enumerate(display_rows):
            for j, val in enumerate(row_data):
                val_str = str(val)
                
                # Format IMCI column for display (Yes::1 -> 1)
                if j == 18 and val_str and "::" in val_str:  # IMCI column
                    val_str = val_str.split('::')[1]  # Just the number
                elif j == 18 and val_str:  # Non-composite IMCI
                    # Extract number if it starts with IMCI-
                    if val_str.startswith("IMCI-"):
                        val_str = val_str.replace("IMCI-", "")
                
                # Use custom items for sorting Date (Col 1) and Serial (Col 2)
                if j == 1 and val_str:
                    item = DateTableWidgetItem(val_str)
                elif j == 2 and val_str:
                    item = NumericTableWidgetItem(val_str)
                else:
                    item = QTableWidgetItem(val_str)
                self.table.setItem(i, j, item)

        self.table.setSortingEnabled(old_sorting)
        self.table.blockSignals(False)
        self._loading_table = False

    def _validate_form(self):
        """ផ្ទៀងផ្ទាត់រាល់ប្រអប់បញ្ចូលដែលចាំបាច់មិនអាចរំលងបាន (Validate mandatory fields)."""
        # ១. កាលបរិច្ឆេទ (Date)
        if not self.date.text().strip():
            QMessageBox.warning(self, "Missing Information", "សូមបញ្ចូលកាលបរិច្ឆេទ!")
            self.date.setFocus()
            return False

        # ២. ឈ្មោះអ្នកជំងឺ (Name)
        if not self.name.text().strip():
            QMessageBox.warning(self, "Missing Information", "សូមបញ្ចូលឈ្មោះអ្នកជំងឺ!")
            self.name.setFocus()
            return False

        # ៣. ប្រភេទអាយុ (Age Category)
        if not self.age_cat.currentText().strip() or self.age_cat.currentIndex() == 0:
            QMessageBox.warning(self, "Missing Information", "សូមជ្រើសរើសប្រភេទអាយុ!")
            self.age_cat.setFocus()
            return False

        inferred_age_type = self._infer_patient_type_from_age_category(self.age_cat.currentText())
        if inferred_age_type and inferred_age_type != self.current_patient_type:
            expected_label = "កុមារ" if inferred_age_type == "Child" else "មនុស្សចាស់"
            current_label = "កុមារ" if self.current_patient_type == "Child" else "មនុស្សចាស់"
            QMessageBox.warning(
                self,
                "Age Category Mismatch",
                f"ប្រភេទអាយុនេះស្ថិតក្នុង Tab '{expected_label}' មិនមែន Tab '{current_label}' ទេ។\n"
                f"សូមប្ដូរទៅ Tab '{expected_label}' មុនពេលរក្សាទុក។"
            )
            self.age_cat.setFocus()
            return False
            
        # ៤. តម្លៃអាយុ (Age Value)
        if self.age_months.isVisible() and self.age_days.isVisible():
            if not self.age_months.text().strip() or not self.age_days.text().strip():
                QMessageBox.warning(self, "Missing Information", "សូមបញ្ចូលខែ និងថ្ងៃសម្រាប់អាយុនេះ!")
                self.age_months.setFocus()
                return False
        elif self.age_days.isVisible():
            if not self.age_days.text().strip():
                QMessageBox.warning(self, "Missing Information", "សូមបញ្ចូលថ្ងៃសម្រាប់អាយុនេះ!")
                self.age_days.setFocus()
                return False
        else:
            age_val_full = self.get_age_value()
            if not age_val_full.strip():
                QMessageBox.warning(self, "Missing Information", "សូមបញ្ចូលអាយុ!")
                if self.age_months.isVisible(): self.age_months.setFocus()
                elif self.age_days.isVisible(): self.age_days.setFocus()
                else: self.age_val.setFocus()
                return False

        # ៥. ភេទ (Sex)
        if not self.sex.currentText().strip():
            QMessageBox.warning(self, "Missing Information", "សូមជ្រើសរើសភេទ!")
            self.sex.setFocus()
            return False

        # ៦. តំបន់ទទួលខុសត្រូវ (Area)
        if not self.area_cat.currentText().strip():
            QMessageBox.warning(self, "Missing Information", "សូមជ្រើសរើសតំបន់ទទួលខុសត្រូវ!")
            self.area_cat.setFocus()
            return False
        if self.CAT_AREA and self.area_cat.currentText().strip() not in self.CAT_AREA:
            QMessageBox.warning(self, "Invalid Area", "សូមជ្រើសរើសតំបន់ត្រឹមត្រូវពីបញ្ជី: ក, ខ, គ")
            self.area_cat.setFocus()
            return False

        # ៧. អាស័យដ្ឋាន (Address)
        if not self.address.text().strip():
            QMessageBox.warning(self, "Missing Information", "សូមបញ្ចូលអាស័យដ្ឋានបច្ចុប្បន្ន!")
            self.address.setFocus()
            return False

        # ៨. ករណីជំងឺ (Disease Case)
        if not self.disease_cat.currentText().strip():
            QMessageBox.warning(self, "Missing Information", "សូមបញ្ចូលករណីជំងឺ!")
            self.disease_cat.setFocus()
            return False

        if self.disease_cat.currentText().strip() not in self.CAT_DISEASE:
            QMessageBox.warning(
                self,
                "Invalid Disease Case",
                "ករណីជំងឺត្រូវតែជ្រើសរើសពីបញ្ជីដែលបានកំណត់។"
            )
            self.disease_cat.setFocus()
            return False

        # ៩. រោគសញ្ញា (Symptoms)
        if not self.symptoms.text().strip():
            QMessageBox.warning(self, "Missing Information", "សូមបញ្ចូលរោគសញ្ញា!")
            self.symptoms.setFocus()
            return False

        # ១០. រោគវិនិច្ឆ័យ (Diagnosis)
        if not self.diagnosis.currentText().strip():
            QMessageBox.warning(self, "Missing Information", "សូមជ្រើសរើសរោគវិនិច្ឆ័យ!")
            self.diagnosis.setFocus()
            return False

        # ១១. ស្ថានភាពអាហាររូបត្ថម (Nutrition) - សម្រាប់តែកុមារ (Child) ត្រូវមានមានមួយដាច់ខាតក្នុងចំណោមទាំង២
        if self.current_patient_type == "Child":
            if not self.nut_wa.text().strip() and not self.nut_wh.text().strip():
                QMessageBox.warning(self, "Missing Information", "សម្រាប់កុមារ ត្រូវមាន ទម្ងន់/អាយុ ឬ ទម្ងន់/កំពស់!")
                self.nut_wa.setFocus()
                return False

        # ១២. ប្រភេទនៃការបង់ថ្លៃសេវា (Service)
        if not self.service.currentText().strip():
            QMessageBox.warning(self, "Missing Information", "សូមជ្រើសរើសប្រភេទនៃការបង់ថ្លៃសេវា!")
            self.service.setFocus()
            return False

        # ១៣. ផ្ទៀងផ្ទាត់ប្រភេទអ្នកជំងឺតាមអាយុ (Age-Type Mismatch Prompt)
        if not self._prompt_patient_type_for_age_mismatch():
            return False

        return True

    def _connect_required_field_signals(self):
        """Keep the Add button locked until every mandatory field is complete."""
        def refresh(*_args):
            self._update_add_button_state()

        for field in [
            self.date, self.name, self.age_val, self.age_months, self.age_days,
            self.address, self.symptoms, self.nut_wa, self.nut_wh,
        ]:
            field.textChanged.connect(refresh)

        for combo in [
            self.age_cat, self.sex, self.area_cat, self.disease_cat,
            self.diagnosis, self.service,
        ]:
            combo.currentTextChanged.connect(refresh)
            combo.currentIndexChanged.connect(refresh)

    def _required_form_fields_complete(self):
        if not self.date.text().strip():
            return False
        if not self.name.text().strip():
            return False
        if not self.age_cat.currentText().strip() or self.age_cat.currentIndex() == 0:
            return False

        inferred_age_type = self._infer_patient_type_from_age_category(self.age_cat.currentText())
        if inferred_age_type and inferred_age_type != self.current_patient_type:
            return False

        if self.age_months.isVisible() and self.age_days.isVisible():
            if not self.age_months.text().strip() or not self.age_days.text().strip():
                return False
        elif self.age_days.isVisible():
            if not self.age_days.text().strip():
                return False
        elif not self.get_age_value().strip():
            return False

        if not self.sex.currentText().strip():
            return False
        if not self.area_cat.currentText().strip():
            return False
        if self.CAT_AREA and self.area_cat.currentText().strip() not in self.CAT_AREA:
            return False
        if not self.address.text().strip():
            return False
        if not self.disease_cat.currentText().strip():
            return False
        if self.CAT_DISEASE and self.disease_cat.currentText().strip() not in self.CAT_DISEASE:
            return False
        if not self.symptoms.text().strip():
            return False
        if not self.diagnosis.currentText().strip():
            return False
        if self.current_patient_type == "Child":
            if not self.nut_wa.text().strip() and not self.nut_wh.text().strip():
                return False
        if not self.service.currentText().strip():
            return False

        return True

    def _update_add_button_state(self):
        if not getattr(self, "btn_add", None):
            return

        complete = self._required_form_fields_complete()
        self.btn_add.setEnabled(complete)
        self.btn_add.setToolTip(
            "" if complete else "សូមបំពេញចន្លោះចាំបាច់ទាំងអស់ មុនពេលចុចបន្ថែម"
        )

    def add(self):
        # --- ផ្ទៀងផ្ទាត់ទិន្នន័យចាំបាច់ (Input Validation) ---
        if not self._validate_form():
            return

        try:
            # --- Check if patient visited TODAY only (not blocking same patient on different days) ---
            branch_clause = ""
            branch_params = ()
            if self.active_branch_code:
                branch_clause = " AND (branch_code = ? OR branch_code IS NULL OR branch_code = '')"
                branch_params = (self.active_branch_code,)
            exists_today = db.execute_read(
                f"SELECT id, name FROM patient WHERE date=? AND name=? AND patient_type=?{branch_clause}",
                (self.date.text(), self.name.text(), self.current_patient_type, *branch_params),
                one=True
            )
            
            if exists_today:
                # Patient already registered TODAY - show warning but allow override
                msg = f"⚠️ អ្នកជំងឺឈ្មោះ {exists_today[1]} បានចុះឈ្មោះរួចហើយនៅថ្ងៃនេះ!\n\n"
                msg += "តើអ្នកចង់បញ្ចូលម្តងទៀតមែនទេ? (ឧ. មកពិគ្រោះម្តងទៀតក្នុងថ្ងៃតែមួយ)"
                reply = QMessageBox.question(self, "Duplicate Entry Today", msg, 
                    QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
                
                if reply == QMessageBox.No:
                    return  # User cancelled - don't add

            # --- Check if patient has visited BEFORE (show history but allow) ---
            patient_type_clause, patient_type_params = db._build_patient_type_filter(self.current_patient_type)
            patient_history = db.execute_read(
                f"""
                SELECT id, date, diagnosis, treatment
                FROM patient
                WHERE name=? AND {patient_type_clause}{branch_clause}
                ORDER BY (SUBSTR(date, 7, 4) || SUBSTR(date, 4, 2) || SUBSTR(date, 1, 2)) DESC
                LIMIT 5
                """,
                (self.name.text(), *patient_type_params, *branch_params),
            )
            
            if patient_history and len(patient_history) > 0:
                # Patient has visited before - show history
                history_msg = f"📋 អ្នកជំងឺនេះធ្លាប់មកពិគ្រោះចំនួន {len(patient_history)} ដងហើយ៖\n\n"
                for i, (pid, pdate, pdiag, ptreat) in enumerate(patient_history, 1):
                    history_msg += f"{i}. ថ្ងៃ {pdate} - {pdiag}\n"
                
                history_msg += "\nតើអ្នកចង់បន្តបញ្ចូលទេ?"
                reply = QMessageBox.question(self, "Patient History Found", history_msg,
                    QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
                
                if reply == QMessageBox.No:
                    return  # User cancelled

            # គណនាលេខរៀងស្វ័យប្រវត្តិ
            area_cat_txt = self.area_cat.currentText()
            area_next = self.get_next_counter(area_cat_txt, 8) 
            
            dis_cat_txt = self.disease_cat.currentText()
            dis_next = self.get_next_counter(dis_cat_txt, 13)

            # Get age value using the new helper method
            age_value = self.get_age_value()

            # Auto-generate IMCI for Children only (No manual input needed)
            if self.current_patient_type == "Child":
                # Automatically set IMCI to "Yes" with auto serial
                imci_cat_txt = "Yes"
                imci_next = self.get_next_counter(imci_cat_txt, 18)
                imci_db = f"{imci_cat_txt}::{imci_next}"
                # Display only the auto-generated IMCI serial number
                self.imci_val.setText(str(imci_next))
            else:
                # Adults: No IMCI
                imci_db = ""
                self.imci_val.clear()

            db.insert(
                self.date.text(), self.serial_no.text(), self.card_id.text(),
                self.name.text(), self.guardian.text(), f"{self.age_cat.currentText()}::{age_value}",
                self.sex.currentText(), f"{area_cat_txt}::{area_next}", self.pregnant.text(),
                self.address.text(), self.phone.text(), self.ref_from.text(),
                f"{dis_cat_txt}::{dis_next}", self.symptoms.text(), self.paraclinical.text(),
                self.diagnosis.currentText(), self.treatment.text(), imci_db,
                self._create_nutrition_string(), self.ref_to.text(), self.service.currentText(),
                self.remark.text(), self.current_patient_type, branch_code=self.branch_code,
                created_by=self.current_user
            )
            self.clear_inputs()
            self.view()
            self.statusBar.showMessage("រក្សាទុកទិន្នន័យបានជោគជ័យ!", 5000)
        except Exception as e:
            QMessageBox.critical(self, "Database Error", f"មិនអាចបញ្ចូលទិន្នន័យបានទេ: {str(e)}")

    def view(self):
        # Filter by current patient type to show separate serial sequences
        rows = db.view_by_patient_type(self.current_patient_type, self.active_branch_code)
        self.load_table(rows)
        # Update the serial number for the currently selected patient type tab
        self.update_next_serial_no()
        self.setup_autocomplete(rows)

    def view_all_patients(self):
        """View all patients without filtering (for export, reports, etc.)"""
        rows = db.view(self.active_branch_code)
        self.load_table(rows)
        self.setup_autocomplete(rows)

    def get_next_counter(self, category, col_idx):
        # New, faster method using a direct SQL query
        column_name = ""
        if col_idx == db.PatientCol.AREA:
            column_name = "area"
        elif col_idx == db.PatientCol.DISEASE:
            column_name = "disease_case"
        elif col_idx == db.PatientCol.IMCI:
            column_name = "imci"

        if not column_name or not category:
            return 1 # Default to 1 if category is empty or column is unknown

        # Debug output to verify patient_type is being passed
        print(f"[DEBUG COUNTER] Getting next counter for: category='{category}', patient_type='{self.current_patient_type}', column='{column_name}'")

        # Pass patient_type to get counter separately for Child/Adult
        max_val = db.get_max_counter_for_category(column_name, category, self.current_patient_type, self.active_branch_code)
        print(f"[DEBUG COUNTER] Max value found: {max_val}, Next will be: {max_val + 1}")
        return max_val + 1

    def get_age_value(self):
        """Get the age value from the next age field(s)."""
        if self.age_months.isVisible() and self.age_days.isVisible():
            months = self.age_months.text().strip()
            days = self.age_days.text().strip()
            parts = []
            if months:
                parts.append(f"{months}ខែ")
            if days:
                parts.append(f"{days}ថ្ងៃ")
            return " ".join(parts).strip()
        if self.age_days.isVisible():
            return self.age_days.text().strip()
        return self.age_val.text().strip()

    def setup_autocomplete(self, rows):
        # កំណត់ Field ដែលចង់ឱ្យមាន Suggestion (Excel-like)
        targets = {
            self.guardian: (db.PatientCol.GUARDIAN, False),
            self.age_cat: (db.PatientCol.AGE, True),
            self.sex: (db.PatientCol.SEX, False),
            self.address: (db.PatientCol.ADDRESS, False),
            self.ref_from: (db.PatientCol.REF_FROM, False),
            self.disease_cat: (db.PatientCol.DISEASE, True),
            self.symptoms: (db.PatientCol.SYMPTOMS, False),
            self.paraclinical: (db.PatientCol.PARACLINICAL, False),
            self.diagnosis: (db.PatientCol.DIAGNOSIS, False),
            self.treatment: (db.PatientCol.TREATMENT, False),
            self.ref_to: (db.PatientCol.REF_TO, False),
            self.service: (db.PatientCol.SERVICE, False),
            self.remark: (db.PatientCol.REMARK, False)
        }

        # ប្រមូលទិន្នន័យចាស់ៗដើម្បីធ្វើជា Suggestion
        data_store = { w: set() for w in targets }

        # បន្ថែម Defaults ដើម្បីឱ្យមាន Suggestion ភ្លាមៗទោះបី DB ទទេ
        defaults = {
            self.age_cat: self.CAT_AGE,
            self.sex: self.CAT_SEX,
            self.disease_cat: self.CAT_DISEASE,
            self.service: self.CAT_SERVICE,
            self.treatment: self.CAT_MEDICINE
        }
        for widget, items in defaults.items():
            data_store[widget].update(items)
        
        for row in rows:
            for widget, (idx, is_composite) in targets.items():
                val = str(row[idx]).strip() if row[idx] else ""
                if not val: continue
                
                # ប្រសិនបើជា Composite (Ex: A::1) យកតែផ្នែកខាងមុខ (A)
                if is_composite and "::" in val:
                    val = val.split("::")[0]
                
                if val:
                    data_store[widget].add(val)
                    if widget == self.diagnosis and self.diagnosis.findText(val) < 0:
                        self.diagnosis.addItem(val)
                    # សម្រាប់ការព្យាបាល បន្ថែមការ Suggest ថ្នាំដាច់ដោយឡែក (ឧ. Amox, Para -> Suggest ទាំង Amox និង Para)
                    if idx == db.PatientCol.TREATMENT and "," in val:
                        parts = [x.strip() for x in val.split(',')]
                        data_store[widget].update(parts)

        # បង្កើត Completer សម្រាប់ Field នីមួយៗ
        for widget, values in data_store.items():
            completer = QCompleter(list(values), self)
            completer.setCaseSensitivity(Qt.CaseInsensitive) # type: ignore # មិនប្រកាន់អក្សរតូចធំ
            completer.setFilterMode(Qt.MatchContains) # type: ignore       # រកពាក្យដែលមាននៅក្នុងឃ្លា (មិនចាំបាច់ដើមឃ្លា)
            completer.setCompletionMode(QCompleter.PopupCompletion) # type: ignore
            
            if isinstance(widget, QComboBox):
                if widget.isEditable():
                    line_edit = widget.lineEdit()
                    if line_edit is None:
                        widget.setLineEdit(QLineEdit())
                        line_edit = widget.lineEdit()
                    widget.setCompleter(completer)
                    if line_edit is not None:
                        completer.activated.connect(line_edit.setText)
                else:
                    # Non-editable QComboBox cannot accept a completer directly
                    continue
            elif isinstance(widget, QLineEdit):
                widget.setCompleter(completer)
                completer.activated.connect(widget.setText)
            else:
                continue

            # Auto-switch patient type based on age category
            if widget == self.age_cat:
                completer.activated.connect(self.auto_switch_patient_type_by_age)


    def delete(self):
        row = self.table.currentRow()
        if row >= 0 and self.selected_id:
            # បន្ថែមការសួរបញ្ជាក់មុននឹងលុប (Confirmation)
            reply = QMessageBox.question(self, 'Confirm Delete', 
                                         f'តើអ្នកពិតជាចង់លុបអ្នកជំងឺដែលមាន ID {self.selected_id} មែនទេ?',
                                         QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.Yes:
                db.delete(self.selected_id, self.active_branch_code)
                self.view()
                self.statusBar.showMessage(f"Patient ID {self.selected_id} deleted.", 5000)
                self.clear_inputs()
        else:
            QMessageBox.warning(self, "Warning", "Please select a row to delete")

    def update_patient(self):
        row = self.table.currentRow()
        if row >= 0:
            # ផ្ទៀងផ្ទាត់ទិន្នន័យចាំបាច់មុនកែប្រែ
            if not self._validate_form():
                return
            item = self.table.item(row, 0)
            if self.selected_id:
                try:
                    patient_data = db.get_patient_by_id(self.selected_id)
                    existing_serial_no = patient_data[2] if patient_data and len(patient_data) > 2 else self.serial_no.text()
                    existing_branch_code = patient_data[24] if patient_data and len(patient_data) > 24 else self.branch_code

                    # Get existing IMCI value from database to preserve it
                    existing_imci = ""
                    if self.current_patient_type == "Child":
                        # Fetch the current IMCI value for this patient
                        if patient_data and len(patient_data) > 18:
                            existing_imci = patient_data[18] or ""
                    
                    # Auto-generate IMCI for Children only if it doesn't exist
                    if self.current_patient_type == "Child":
                        if not existing_imci:
                            # Only generate new IMCI if none exists
                            imci_cat_txt = "Yes"
                            imci_next = self.get_next_counter(imci_cat_txt, 18)
                            imci_db = f"{imci_cat_txt}::{imci_next}"
                            self.imci_val.setText(str(imci_next))
                        else:
                            # Keep existing IMCI
                            imci_db = existing_imci
                            # Display existing IMCI number
                            if "::" in existing_imci:
                                imci_num = existing_imci.split("::")[1]
                                self.imci_val.setText(imci_num)
                            else:
                                self.imci_val.setText(existing_imci)
                    else:
                        # Adults: no IMCI
                        imci_db = ""
                        self.imci_val.clear()

                    # Get age value using the new helper method
                    age_value = self.get_age_value()

                    db.update(
                        self.selected_id, self.date.text(), existing_serial_no, self.card_id.text(),
                        self.name.text(), self.guardian.text(), f"{self.age_cat.currentText()}::{age_value}",
                        self.sex.currentText(), f"{self.area_cat.currentText()}::{self.area_val.text()}", self.pregnant.text(),
                        self.address.text(), self.phone.text(), self.ref_from.text(),
                        f"{self.disease_cat.currentText()}::{self.disease_val.text()}", self.symptoms.text(),
                        self.paraclinical.text(), self.diagnosis.currentText(), self.treatment.text(),
                        imci_db, self._create_nutrition_string(),
                        self.ref_to.text(), self.service.currentText(), self.remark.text(),
                        self.current_patient_type, existing_branch_code, updated_by=self.current_user
                    )
                    self.view()
                    self.statusBar.showMessage(f"កែប្រែទិន្នន័យ ID {self.selected_id} បានជោគជ័យ!", 5000)
                    self.clear_inputs()
                except Exception as e:
                    QMessageBox.critical(self, "Update Error", f"មិនអាចកែប្រែទិន្នន័យបានទេ: {str(e)}")
        else:
            QMessageBox.warning(self, "Warning", "Please select a row to update")

    def search(self):
        # Use the dedicated search bar input instead of the form's name field
        keyword = self.search_input.text().strip()
        # Filter search results by current patient type
        rows = db.search(keyword, self.current_patient_type, self.active_branch_code)
        self.load_table(rows)
        self.update_next_serial_no()

    def show_statistics(self):
        """Fetches statistics from the database and displays them in a dialog."""
        try:
            s = db.get_statistics(self.active_branch_code)
            c = s['child']  # Children stats
            a = s['adult']  # Adult stats
            o = s['overall']  # Overall stats

            # Theme-based colors for statistics report
            is_light_theme = self.current_theme == "Classic Light"

            # Header colors based on theme
            if is_light_theme:
                overall_header = "#e74c3c"  # Red for overall
                child_header = "#27ae60"  # Green for children
                adult_header = "#8e44ad"  # Purple for adults
                comparison_header = "#f39c12"  # Orange for comparison
                age_header = "#16a085"  # Teal for age groups
                section_bg = "#e8f8f5"
                section_text = "#145a32"
                adult_section_bg = "#f5eef8"
                adult_section_text = "#4a235a"
                table_bg = "#f0f3f4"
                row_alt = "#d5dbdb"
                row_normal = "#eaf2f8"
            else:
                overall_header = "#c0392b"  # Dark red
                child_header = "#1e8449"  # Dark green
                adult_header = "#6c3483"  # Dark purple
                comparison_header = "#d68910"  # Dark orange
                age_header = "#148f77"  # Dark teal
                section_bg = "#d5f5e3"
                section_text = "#145a32"
                adult_section_bg = "#e8daef"
                adult_section_text = "#4a235a"
                table_bg = "#34495e"
                row_alt = "#2c3e50"
                row_normal = "#3d566e"

            header_text = "black" if is_light_theme else "white"
            row_text = "#1a1a1a" if is_light_theme else "#ecf0f1"

            # Calculate percentages
            def calc_pct(part, total):
                if total == 0: return 0
                return round((part / total) * 100, 1)

            # Overall percentages
            total_patients = o['total']
            today_pct = calc_pct(o['today_count'], total_patients)
            new_total = c['new_total'] + a['new_total']
            old_total = c['old_total'] + a['old_total']
            new_pct = calc_pct(new_total, total_patients)
            old_pct = calc_pct(old_total, total_patients)

            # Area percentages
            area_a_total = c['area_a_new'] + a['area_a_new']
            area_b_total = c['area_b_new'] + a['area_b_new']
            area_c_total = c['area_c_new'] + a['area_c_new']
            total_area = area_a_total + area_b_total + area_c_total
            area_a_pct = calc_pct(area_a_total, total_area)
            area_b_pct = calc_pct(area_b_total, total_area)
            area_c_pct = calc_pct(area_c_total, total_area)

            # Area totals (New + Old cases)
            area_a_total_all = c['area_a_total_m'] + c['area_a_total_f'] + a['area_a_total_m'] + a['area_a_total_f']
            area_b_total_all = c['area_b_total_m'] + c['area_b_total_f'] + a['area_b_total_m'] + a['area_b_total_f']
            area_c_total_all = c['area_c_total_m'] + c['area_c_total_f'] + a['area_c_total_m'] + a['area_c_total_f']
            
            # Area new cases breakdown
            area_a_new_m = c['area_a_new_m'] + a['area_a_new_m']
            area_a_new_f = c['area_a_new_f'] + a['area_a_new_f']
            area_b_new_m = c['area_b_new_m'] + a['area_b_new_m']
            area_b_new_f = c['area_b_new_f'] + a['area_b_new_f']
            area_c_new_m = c['area_c_new_m'] + a['area_c_new_m']
            area_c_new_f = c['area_c_new_f'] + a['area_c_new_f']
            
            # Area total cases breakdown (New + Old)
            area_a_all_m = c['area_a_total_m'] + a['area_a_total_m']
            area_a_all_f = c['area_a_total_f'] + a['area_a_total_f']
            area_b_all_m = c['area_b_total_m'] + a['area_b_total_m']
            area_b_all_f = c['area_b_total_f'] + a['area_b_total_f']
            area_c_all_m = c['area_c_total_m'] + a['area_c_total_m']
            area_c_all_f = c['area_c_total_f'] + a['area_c_total_f']

            # Service totals
            hef_total = c['hef_m'] + c['hef_f'] + a['hef_m'] + a['hef_f']
            pay_total = c['pay_m'] + c['pay_f'] + a['pay_m'] + a['pay_f']
            free_total = c['free_m'] + c['free_f'] + a['free_m'] + a['free_f']
            hefr_total = c['hefr_m'] + c['hefr_f'] + a['hefr_m'] + a['hefr_f']
            hefi_total = c['hefi_m'] + c['hefi_f'] + a['hefi_m'] + a['hefi_f']
            nssf8_total = c['nssf8_m'] + c['nssf8_f'] + a['nssf8_m'] + a['nssf8_f']
            nssf7_total = c['nssf7_m'] + c['nssf7_f'] + a['nssf7_m'] + a['nssf7_f']
            nssfa_total = c['nssfa_m'] + c['nssfa_f'] + a['nssfa_m'] + a['nssfa_f']
            other_total = c['other_m'] + c['other_f'] + a['other_m'] + a['other_f']
            service_total = hef_total + hefr_total + hefi_total + pay_total + free_total + nssfa_total + nssf7_total + nssf8_total + other_total
            service_m_total = (
                c['hef_m'] + a['hef_m'] + c['hefr_m'] + a['hefr_m'] + c['hefi_m'] + a['hefi_m']
                + c['pay_m'] + a['pay_m'] + c['free_m'] + a['free_m'] + c['nssfa_m'] + a['nssfa_m']
                + c['nssf7_m'] + a['nssf7_m'] + c['nssf8_m'] + a['nssf8_m']
                + c['other_m'] + a['other_m']
            )
            service_f_total = (
                c['hef_f'] + a['hef_f'] + c['hefr_f'] + a['hefr_f'] + c['hefi_f'] + a['hefi_f']
                + c['pay_f'] + a['pay_f'] + c['free_f'] + a['free_f'] + c['nssfa_f'] + a['nssfa_f']
                + c['nssf7_f'] + a['nssf7_f'] + c['nssf8_f'] + a['nssf8_f']
                + c['other_f'] + a['other_f']
            )

            # Child vs Adult comparison
            child_total = c['total']
            adult_total = a['total']
            child_pct = calc_pct(child_total, total_patients)
            adult_pct = calc_pct(adult_total, total_patients)

            # Age group stats for children
            child_age_total = c['total']

            import html
            diagnosis_age_columns = [
                ("០-២៩ថ្ងៃ", "age_0_29_days"),
                ("២៩ថ្ងៃ-១១ខែ", "age_29days_11months"),
                ("១-៤ឆ្នាំ", "age_1_4_years"),
                ("៥-១៤ឆ្នាំ", "age_5_14_years"),
                ("១៥-២៤ឆ្នាំ", "age_15_24_years"),
                ("២៥-៤៩ឆ្នាំ", "age_25_49_years"),
                ("៥០-៦៤ឆ្នាំ", "age_50_64_years"),
                (">=៦៥ឆ្នាំ", "age_64_plus"),
            ]
            diagnosis_age_group_headers = "".join(
                f"<td colspan='2' style='color: {header_text}; text-align: center;'><b>{label}</b></td>"
                for label, _ in diagnosis_age_columns
            )
            diagnosis_age_sex_headers = "".join(
                f"<td style='color: {header_text}; text-align: center;'><b>ប្រុស</b></td>"
                f"<td style='color: {header_text}; text-align: center;'><b>ស្រី</b></td>"
                for _label, _key in diagnosis_age_columns
            )
            def build_diagnosis_rows_html(diagnosis_stats, percent_base, empty_text):
                if not diagnosis_stats:
                    return f"""
                        <tr bgcolor='{row_normal}'>
                            <td colspan='23' style='text-align: center; color: {row_text};'>{empty_text}</td>
                        </tr>
                    """

                diagnosis_rows = []
                diagnosis_totals = {
                    "total": 0,
                    "male": 0,
                    "female": 0,
                    "ref_to_total": 0,
                    "hef_total": 0,
                }
                for _label, key in diagnosis_age_columns:
                    diagnosis_totals[f"{key}_m"] = 0
                    diagnosis_totals[f"{key}_f"] = 0

                for index, item in enumerate(diagnosis_stats):
                    bg = row_normal if index % 2 == 0 else row_alt
                    diagnosis_name = html.escape(str(item["diagnosis"]))
                    diagnosis_totals["total"] += item["total"]
                    diagnosis_totals["male"] += item["male"]
                    diagnosis_totals["female"] += item["female"]
                    diagnosis_totals["ref_to_total"] += item["ref_to_total"]
                    diagnosis_totals["hef_total"] += item["hef_total"]
                    for _label, key in diagnosis_age_columns:
                        diagnosis_totals[f"{key}_m"] += item[f"{key}_m"]
                        diagnosis_totals[f"{key}_f"] += item[f"{key}_f"]

                    age_cells = "".join(
                        f"<td style='text-align: center; color: {row_text};'>{item[f'{key}_m']}</td>"
                        f"<td style='text-align: center; color: {row_text};'>{item[f'{key}_f']}</td>"
                        for _label, key in diagnosis_age_columns
                    )
                    diagnosis_rows.append(f"""
                        <tr bgcolor='{bg}'>
                            <td style='color: {row_text};'>{diagnosis_name}</td>
                            <td style='text-align: center; color: {row_text};'><b>{item["total"]}</b></td>
                            <td style='text-align: center; color: {row_text};'>{item["male"]}</td>
                            <td style='text-align: center; color: {row_text};'>{item["female"]}</td>
                            {age_cells}
                            <td style='text-align: center; color: {row_text};'>{item["ref_to_total"]}</td>
                            <td style='text-align: center; color: {row_text};'>{item["hef_total"]}</td>
                            <td style='text-align: center; color: {row_text};'>{calc_pct(item["total"], percent_base)}%</td>
                        </tr>
                    """)
                total_age_cells = "".join(
                    f"<td style='text-align: center; color: {header_text};'><b>{diagnosis_totals[f'{key}_m']}</b></td>"
                    f"<td style='text-align: center; color: {header_text};'><b>{diagnosis_totals[f'{key}_f']}</b></td>"
                    for _label, key in diagnosis_age_columns
                )
                diagnosis_rows.append(f"""
                    <tr bgcolor='{comparison_header}'>
                        <td style='color: {header_text};'><b>សរុប</b></td>
                        <td style='text-align: center; color: {header_text};'><b>{diagnosis_totals["total"]}</b></td>
                        <td style='text-align: center; color: {header_text};'><b>{diagnosis_totals["male"]}</b></td>
                        <td style='text-align: center; color: {header_text};'><b>{diagnosis_totals["female"]}</b></td>
                        {total_age_cells}
                        <td style='text-align: center; color: {header_text};'><b>{diagnosis_totals["ref_to_total"]}</b></td>
                        <td style='text-align: center; color: {header_text};'><b>{diagnosis_totals["hef_total"]}</b></td>
                        <td style='text-align: center; color: {header_text};'><b>{calc_pct(diagnosis_totals["total"], percent_base)}%</b></td>
                    </tr>
                """)
                return "".join(diagnosis_rows)

            diagnosis_rows_html = build_diagnosis_rows_html(
                db.get_diagnosis_statistics(branch_code=self.active_branch_code),
                total_patients,
                "គ្មានទិន្នន័យរោគវិនិច្ឆ័យ"
            )
            new_case_diagnosis_rows_html = build_diagnosis_rows_html(
                db.get_diagnosis_statistics(new_only=True, branch_code=self.active_branch_code),
                new_total,
                "គ្មានទិន្នន័យរោគវិនិច្ឆ័យសម្រាប់ករណីថ្មី"
            )


            def build_area_new_rows_html():
                rows = []
                area_rows = [
                    ("តំបន់ ក", "area_a", area_a_total, area_a_new_m, area_a_new_f, area_a_pct, row_normal, row_alt, row_normal),
                    ("តំបន់ ខ", "area_b", area_b_total, area_b_new_m, area_b_new_f, area_b_pct, row_alt, row_normal, row_alt),
                    ("តំបន់ គ", "area_c", area_c_total, area_c_new_m, area_c_new_f, area_c_pct, row_normal, row_alt, row_normal),
                ]
                for label, key, total, male, female, pct, total_bg, child_bg, adult_bg in area_rows:
                    rows.append(f"""
                <tr bgcolor='{total_bg}'>
                    <td>{label}</td>
                    <td style='text-align: center; color: {row_text};'>{total}</td>
                    <td style='text-align: center; color: {row_text};'>{male}</td>
                    <td style='text-align: center; color: {row_text};'>{female}</td>
                    <td style='text-align: center; color: {row_text};'>{pct}%</td>
                </tr>
                <tr bgcolor='{child_bg}'>
                    <td style='padding-left: 20px;'>↳ កុមារ</td>
                    <td style='text-align: center; color: {row_text};'>{c[f'{key}_new']}</td>
                    <td style='text-align: center; color: {row_text};'>{c[f'{key}_new_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c[f'{key}_new_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(c[f'{key}_new'], new_total)}%</td>
                </tr>
                <tr bgcolor='{adult_bg}'>
                    <td style='padding-left: 20px;'>↳ មនុស្សចាស់</td>
                    <td style='text-align: center; color: {row_text};'>{a[f'{key}_new']}</td>
                    <td style='text-align: center; color: {row_text};'>{a[f'{key}_new_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{a[f'{key}_new_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(a[f'{key}_new'], new_total)}%</td>
                </tr>
                    """)
                return "".join(rows)

            def build_area_total_rows_html():
                rows = []
                area_rows = [
                    ("តំបន់ ក (សរុប)", area_a_all_m, area_a_all_f, row_normal),
                    ("តំបន់ ខ (សរុប)", area_b_all_m, area_b_all_f, row_alt),
                    ("តំបន់ គ (សរុប)", area_c_all_m, area_c_all_f, row_normal),
                ]
                for label, male, female, bg in area_rows:
                    total = male + female
                    rows.append(f"""
                <tr bgcolor='{bg}'>
                    <td>{label}</td>
                    <td style='text-align: center; color: {row_text};'>{total}</td>
                    <td style='text-align: center; color: {row_text};'>{male}</td>
                    <td style='text-align: center; color: {row_text};'>{female}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(total, total_patients)}%</td>
                </tr>
                    """)
                return "".join(rows)

            area_new_rows_html = build_area_new_rows_html()
            area_total_rows_html = build_area_total_rows_html()
            service_use_summary = db.get_service_use_summary(self.active_branch_code)

            def fmt_num(value):
                return f"{int(value or 0):,}"

            service_use_summary_html = f"""
            <h3 style='color: {comparison_header}; font-size: 22px; margin: 18px 0 8px 0;'>I. ស្ថិតិអ្នកប្រើសេវា</h3>
            <table border='1' cellpadding='6' width='100%' bgcolor='{table_bg}' style='border-collapse: collapse; font-size: 12px; border-color: #000000;'>
                <tr bgcolor='{comparison_header}'>
                    <td style='color: {header_text}; width: 28%; border-color: #000000;'><b>១ - បរិយាយលទ្ធផលការផ្តល់សេវា</b></td>
                    <td style='color: {header_text}; text-align: center; border-color: #000000;'><b>តំបន់ ក</b></td>
                    <td style='color: {header_text}; text-align: center; border-color: #000000;'><b>តំបន់ ខ</b></td>
                    <td style='color: {header_text}; text-align: center; border-color: #000000;'><b>តំបន់ គ</b></td>
                    <td style='color: {header_text}; text-align: center; border-color: #000000;'><b>សរុប</b></td>
                    <td style='color: {header_text}; text-align: center; border-color: #000000;'><b>ប្រុស</b></td>
                    <td style='color: {header_text}; text-align: center; border-color: #000000;'><b>ស្រី</b></td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td style='color: {row_text}; border-color: #000000;'>សរុបចំនួនអ្នកប្រើសេវា</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(area_a_total_all)}</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(area_b_total_all)}</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(area_c_total_all)}</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'><b>{fmt_num(service_use_summary["total"])}</b></td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(service_use_summary["male"])}</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(service_use_summary["female"])}</td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td style='color: {row_text}; border-color: #000000;'>ក្នុងនោះ ករណីថ្មី:សរុប</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(area_a_total)}</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(area_b_total)}</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(area_c_total)}</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'><b>{fmt_num(service_use_summary["new_total"])}</b></td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(service_use_summary["new_male"])}</td>
                    <td style='text-align: center; color: {row_text}; border-color: #000000;'>{fmt_num(service_use_summary["new_female"])}</td>
                </tr>
            </table>
            """
            msg = f"""
            <html>
            <body style='background-color: {table_bg}; color: {row_text}; font-family: "{self.current_font}", Arial, sans-serif;'>
            <h2 style='color: #0fbcf9; text-align: center;'>📊 របាយការណ៍ស្ថិតិអ្នកជំងឺ (Detailed Statistics)</h2>
            <hr style='border: 1px solid #0fbcf9;'>

            <h3 style='color: {overall_header}; font-size: 18px;'>🔵 ផ្នែកទី ១: ស្ថិតិសរុបទាំងអស់ (Overall Total)</h3>
            <table border='1' cellpadding='8' width='100%' bgcolor='{table_bg}' style='border-collapse: collapse;'>
                <tr bgcolor='{overall_header}'>
                    <td style='color: {header_text};'><b>ពិពណ៌នា</b></td>
                    <td style='color: {header_text}; text-align: center;'><b>សរុប</b></td>
                    <td style='color: {header_text}; text-align: center;'><b>ប្រុស</b></td>
                    <td style='color: {header_text}; text-align: center;'><b>ស្រី</b></td>
                    <td style='color: {header_text}; text-align: center;'><b>ភាគរយ</b></td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td><b>អ្នកជំងឺសរុប</b></td>
                    <td style='text-align: center; color: {row_text};'><b>{o['total']}</b></td>
                    <td style='text-align: center; color: {row_text};'>{o['male']}</td>
                    <td style='text-align: center; color: {row_text};'>{o['female']}</td>
                    <td style='text-align: center; color: #e74c3c;'><b>100%</b></td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td>ថ្ងៃនេះ</td>
                    <td style='text-align: center; color: {row_text};'>{o['today_count']}</td>
                    <td style='text-align: center; color: {row_text};'>-</td>
                    <td style='text-align: center; color: {row_text};'>-</td>
                    <td style='text-align: center; color: {row_text};'>{today_pct}%</td>
                </tr>
                <tr bgcolor='{section_bg}'>
                    <td colspan='5' style='color: {section_text};'><b>📋 ករណីជំងឺ</b></td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td>↳ ករណីថ្មី</td>
                    <td style='text-align: center; color: {row_text};'>{new_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['new_m'] + a['new_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['new_f'] + a['new_f']}</td>
                    <td style='text-align: center; color: #27ae60;'>{new_pct}%</td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td>↳ ករណីចាស់</td>
                    <td style='text-align: center; color: {row_text};'>{old_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['old_m'] + a['old_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['old_f'] + a['old_f']}</td>
                    <td style='text-align: center; color: #f39c12;'>{old_pct}%</td>
                </tr>
                <tr bgcolor='{section_bg}'>
                    <td colspan='5' style='color: {section_text};'><b>📍 តាមតំបន់ (ករណីថ្មី)</b></td>
                </tr>
                {area_new_rows_html}
                <tr bgcolor='{comparison_header}'>
                    <td style='color: {header_text};'><b>ករណីថ្មីសរុប</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{new_total}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{c['new_m'] + a['new_m']}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{c['new_f'] + a['new_f']}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{new_pct}%</b></td>
                </tr>
                <tr bgcolor='{section_bg}'>
                    <td colspan='5' style='color: {section_text};'><b>📍 សរុបតាមតំបន់ (ករណីថ្មី + ចាស់)</b></td>
                </tr>
                {area_total_rows_html}
                <tr bgcolor='{comparison_header}'>
                    <td style='color: {header_text};'><b>ករណីសរុបទាំងអស់</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{total_patients}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{o['male']}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{o['female']}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>100%</b></td>
                </tr>
                <tr bgcolor='{section_bg}'>
                    <td colspan='5' style='color: {section_text};'><b>👶 និន្ន័យអាយុកុមារ (Child Age Trends)</b></td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td>០-២៩ ថ្ងៃ (Newborn)</td>
                    <td style='text-align: center; color: {row_text};'>{c['age_0_29_days']}</td>
                    <td style='text-align: center; color: #e74c3c;'>{calc_pct(c['age_0_29_days'], c['total'])}%</td>
                    <td style='text-align: center; color: {row_text};'>-</td>
                    <td style='text-align: center; color: {row_text};'>-</td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td>២៩ ថ្ងៃ - ១១ ខែ (Infant)</td>
                    <td style='text-align: center; color: {row_text};'>{c['age_29days_11months']}</td>
                    <td style='text-align: center; color: #e74c3c;'>{calc_pct(c['age_29days_11months'], c['total'])}%</td>
                    <td style='text-align: center; color: {row_text};'>-</td>
                    <td style='text-align: center; color: {row_text};'>-</td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td>១ - ៤ ឆ្នាំ (Toddler)</td>
                    <td style='text-align: center; color: {row_text};'>{c['age_1_4_years']}</td>
                    <td style='text-align: center; color: #e74c3c;'>{calc_pct(c['age_1_4_years'], c['total'])}%</td>
                    <td style='text-align: center; color: {row_text};'>-</td>
                    <td style='text-align: center; color: {row_text};'>-</td>
                </tr>
                <tr bgcolor='{comparison_header}'>
                    <td style='color: {header_text};'><b>កុមារសរុប (០-៤ ឆ្នាំ)</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{c['age_0_29_days'] + c['age_29days_11months'] + c['age_1_4_years']}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{calc_pct(c['age_0_29_days'] + c['age_29days_11months'] + c['age_1_4_years'], c['total'])}%</b></td>
                    <td style='text-align: center; color: {header_text};'>-</td>
                    <td style='text-align: center; color: {header_text};'>-</td>
                </tr>
                <tr bgcolor='{section_bg}'>
                    <td colspan='5' style='color: {section_text};'><b>💼 សរុបរួមនៃប្រភេទសេវាកម្ម</b></td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td>HEF</td>
                    <td style='text-align: center; color: {row_text};'>{hef_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['hef_m'] + a['hef_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['hef_f'] + a['hef_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(hef_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td>HEF-R</td>
                    <td style='text-align: center; color: {row_text};'>{hefr_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['hefr_m'] + a['hefr_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['hefr_f'] + a['hefr_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(hefr_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td>HEF-I</td>
                    <td style='text-align: center; color: {row_text};'>{hefi_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['hefi_m'] + a['hefi_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['hefi_f'] + a['hefi_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(hefi_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td>PAY</td>
                    <td style='text-align: center; color: {row_text};'>{pay_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['pay_m'] + a['pay_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['pay_f'] + a['pay_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(pay_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td>FREE</td>
                    <td style='text-align: center; color: {row_text};'>{free_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['free_m'] + a['free_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['free_f'] + a['free_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(free_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td>NSSF-A</td>
                    <td style='text-align: center; color: {row_text};'>{nssfa_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['nssfa_m'] + a['nssfa_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['nssfa_f'] + a['nssfa_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(nssfa_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td>NSSF-7</td>
                    <td style='text-align: center; color: {row_text};'>{nssf7_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['nssf7_m'] + a['nssf7_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['nssf7_f'] + a['nssf7_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(nssf7_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td>NSSF-8</td>
                    <td style='text-align: center; color: {row_text};'>{nssf8_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['nssf8_m'] + a['nssf8_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['nssf8_f'] + a['nssf8_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(nssf8_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td>Other</td>
                    <td style='text-align: center; color: {row_text};'>{other_total}</td>
                    <td style='text-align: center; color: {row_text};'>{c['other_m'] + a['other_m']}</td>
                    <td style='text-align: center; color: {row_text};'>{c['other_f'] + a['other_f']}</td>
                    <td style='text-align: center; color: {row_text};'>{calc_pct(other_total, service_total)}%</td>
                </tr>
                <tr bgcolor='{comparison_header}'>
                    <td style='color: {header_text};'><b>សរុបសេវាកម្ម</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{service_total}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{service_m_total}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>{service_f_total}</b></td>
                    <td style='text-align: center; color: {header_text};'><b>100%</b></td>
                </tr>
            </table>

            {service_use_summary_html}

            <h3 style='color: {comparison_header}; font-size: 18px; margin-top: 20px;'>សរុបចំនួនអ្នកប្រើសេវាតាមប្រភេទជំងឺ ចន្លោះអាយុ និងភេទ</h3>
            <table border='1' cellpadding='4' width='100%' bgcolor='{table_bg}' style='border-collapse: collapse; font-size: 11px;'>
                <tr bgcolor='{comparison_header}'>
                    <td rowspan='2' style='color: {header_text};'><b>ប្រភេទជំងឺ</b></td>
                    <td rowspan='2' style='color: {header_text}; text-align: center;'><b>សរុប</b></td>
                    <td colspan='2' style='color: {header_text}; text-align: center;'><b>ភេទសរុប</b></td>
                    {diagnosis_age_group_headers}
                    <td rowspan='2' style='color: {header_text}; text-align: center;'><b>បញ្ជូនទៅ</b></td>
                    <td rowspan='2' style='color: {header_text}; text-align: center;'><b>HEF សរុប</b></td>
                    <td rowspan='2' style='color: {header_text}; text-align: center;'><b>ភាគរយ</b></td>
                </tr>
                <tr bgcolor='{comparison_header}'>
                    <td style='color: {header_text}; text-align: center;'><b>ប្រុស</b></td>
                    <td style='color: {header_text}; text-align: center;'><b>ស្រី</b></td>
                    {diagnosis_age_sex_headers}
                </tr>
                {diagnosis_rows_html}
            </table>

            <h3 style='color: {comparison_header}; font-size: 18px; margin-top: 20px;'>សរុបចំនួនអ្នកប្រើសេវាតាមប្រភេទជំងឺ ចន្លោះអាយុ និងភេទ (ករណីថ្មី)</h3>
            <table border='1' cellpadding='4' width='100%' bgcolor='{table_bg}' style='border-collapse: collapse; font-size: 11px;'>
                <tr bgcolor='{comparison_header}'>
                    <td rowspan='2' style='color: {header_text};'><b>ប្រភេទជំងឺ</b></td>
                    <td rowspan='2' style='color: {header_text}; text-align: center;'><b>សរុប</b></td>
                    <td colspan='2' style='color: {header_text}; text-align: center;'><b>ភេទសរុប</b></td>
                    {diagnosis_age_group_headers}
                    <td rowspan='2' style='color: {header_text}; text-align: center;'><b>បញ្ជូនទៅ</b></td>
                    <td rowspan='2' style='color: {header_text}; text-align: center;'><b>HEF សរុប</b></td>
                    <td rowspan='2' style='color: {header_text}; text-align: center;'><b>ភាគរយ</b></td>
                </tr>
                <tr bgcolor='{comparison_header}'>
                    <td style='color: {header_text}; text-align: center;'><b>ប្រុស</b></td>
                    <td style='color: {header_text}; text-align: center;'><b>ស្រី</b></td>
                    {diagnosis_age_sex_headers}
                </tr>
                {new_case_diagnosis_rows_html}
            </table>

            <h3 style='color: {comparison_header}; font-size: 18px; margin-top: 20px;'>📊 ការប្រៀបធៀប កុមារ និង មនុស្សចាស់</h3>
            <table border='1' cellpadding='8' width='100%' bgcolor='{table_bg}' style='border-collapse: collapse;'>
                <tr bgcolor='{comparison_header}'>
                    <td style='color: {header_text}; width: 30%;'><b>ប្រភេទ</b></td>
                    <td style='color: {header_text}; text-align: center; width: 20%;'><b>ចំនួន</b></td>
                    <td style='color: {header_text}; text-align: center; width: 20%;'><b>ភាគរយ</b></td>
                    <td style='color: {header_text}; text-align: center; width: 30%;'><b>ក្រាហ្វិក</b></td>
                </tr>
                <tr bgcolor='{row_normal}'>
                    <td style='color: #27ae60;'><b>🟢 កុមារ</b></td>
                    <td style='text-align: center; color: {row_text};'><b>{child_total}</b></td>
                    <td style='text-align: center; color: #27ae60;'><b>{child_pct}%</b></td>
                    <td style='text-align: center;'>{'█' * int(child_pct/5)}</td>
                </tr>
                <tr bgcolor='{row_alt}'>
                    <td style='color: #8e44ad;'><b>🟣 មនុស្សចាស់</b></td>
                    <td style='text-align: center; color: {row_text};'><b>{adult_total}</b></td>
                    <td style='text-align: center; color: #8e44ad;'><b>{adult_pct}%</b></td>
                    <td style='text-align: center;'>{'█' * int(adult_pct/5)}</td>
                </tr>
                <tr bgcolor='{section_bg}'>
                    <td colspan='4' style='color: {section_text}; text-align: center;'><b>សមាមាត្រ: {child_total}:{adult_total} (កុមារ:មនុស្សចាស់)</b></td>
                </tr>
            </table>

            </body>
            </html>
            """

            dlg = QDialog(self)
            dlg.setWindowTitle("Detailed Patient Statistics")
            screen = QApplication.primaryScreen()
            available = screen.availableGeometry() if screen else None
            if available:
                dlg.resize(int(available.width() * 0.90), int(available.height() * 0.86))
                dlg.setMinimumSize(min(1180, available.width()), min(760, available.height()))
            else:
                dlg.resize(1280, 850)
                dlg.setMinimumSize(1180, 760)
            layout = QVBoxLayout(dlg)
            layout.setContentsMargins(8, 8, 8, 8)
            layout.setSpacing(8)

            report_view = QTextBrowser()
            report_view.setReadOnly(True)
            report_view.setOpenExternalLinks(False)
            report_view.setLineWrapMode(QTextEdit.NoWrap)
            report_view.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)  # type: ignore[attr-defined]
            report_view.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)  # type: ignore[attr-defined]
            report_view.verticalScrollBar().setSingleStep(28)
            report_view.verticalScrollBar().setPageStep(360)
            report_view.horizontalScrollBar().setSingleStep(36)
            report_view.document().setTextWidth(2200)
            report_view.setHtml(msg)
            report_view.setStyleSheet("""
                QTextBrowser {
                    background-color: #2c3e50;
                    color: #ecf0f1;
                    border: 1px solid #0fbcf9;
                    border-radius: 6px;
                    padding: 8px;
                }
            """)
            layout.addWidget(report_view, 1)

            # Export buttons
            btn_layout = QHBoxLayout()
            
            btn_copy = QPushButton("📋 Copy to Clipboard")
            btn_copy.setStyleSheet("background-color: #3498db; color: white; padding: 10px; font-weight: bold;")
            btn_copy.clicked.connect(lambda: (QApplication.clipboard().setText(msg), None)[1])  # type: ignore
            
            btn_export_excel = QPushButton("📊 Export Excel")
            btn_export_excel.setStyleSheet("background-color: #27ae60; color: white; padding: 10px; font-weight: bold;")
            btn_export_excel.clicked.connect(lambda: self.export_statistics_to_excel())
            
            btn_close = QPushButton("❌ Close")
            btn_close.setStyleSheet("background-color: #e74c3c; color: white; padding: 10px; font-weight: bold;")
            btn_close.clicked.connect(dlg.accept)
            
            btn_layout.addStretch()
            btn_layout.addWidget(btn_copy)
            btn_layout.addWidget(btn_export_excel)
            btn_layout.addWidget(btn_close)
            btn_layout.addStretch()
            
            layout.addLayout(btn_layout)
            dlg.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"មិនអាចទាញយកស្ថិតិបានទេ: {str(e)}")

    def show_login_history(self):
        dialog = LoginHistoryDialog(self)
        dialog.exec_()

    def restore_database(self):
        if not self.is_admin:
            QMessageBox.warning(self, "Permission Denied", "Only admin can restore a full database backup.")
            return
        # Critical warning to the user
        reply = QMessageBox.critical(self, 'Confirm Restore', 
                                     '<b>WARNING:</b> This will completely overwrite all current data with the selected backup.\n\n'
                                     'This action CANNOT be undone.\n\n'
                                     'Are you absolutely sure you want to proceed?',
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        
        if reply != QMessageBox.Yes:
            self.statusBar.showMessage("Restore operation cancelled.", 3000)
            return

        # Open file dialog to select a backup
        backup_path, _ = QFileDialog.getOpenFileName(self, "Select Backup to Restore", self.backup_dir, "SQLite Files (*.db)")
        if not backup_path:
            return

        try:
            shutil.copy(backup_path, db.DB_NAME)
            QMessageBox.information(self, "Success", "Database restored successfully. The application will now restart to apply changes.")
            self.logout() # Use the existing logout mechanism to restart the app
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Restore failed: {str(e)}")

    def open_settings(self):
        # --- Security Check: Ask for Admin Password ---
        pwd, ok = QInputDialog.getText(self, "Admin Access", "Enter Admin Password:", QLineEdit.Password)
        if not ok:
            return
        
        # Security: ឆែក Password ធៀបជាមួយគណនីដែលកំពុង Login បច្ចុប្បន្ន
        if not db.check_user(self.current_user, pwd.strip()):
            QMessageBox.warning(self, "Access Denied", "Incorrect password.")
            return

        dialog = SettingsDialog(self.config, self)
        if dialog.exec_() == QDialog.Accepted:
            # Save to file
            with open(self.settings_file, 'w', encoding='utf-8') as configfile:
                self.config.write(configfile)
            
            self._reload_category_settings()
            self.apply_app_font(self.current_font)

            # Refresh entry dropdowns with blank first selections
            self._set_combo_items_blank_first(self.sex, ["ស្រី", "ប្រុស"])
            self._set_combo_items_blank_first(self.service, self.CAT_SERVICE)
            self._set_combo_items_blank_first(self.area_cat, self.CAT_AREA)

            # Refresh Disease dropdown with new categories
            self._set_combo_items_blank_first(self.disease_cat, self.CAT_DISEASE)

            if hasattr(self, 'rep_area'):
                self.rep_area.clear()
                self.rep_area.addItem("All")
                self.rep_area.addItems(self.CAT_AREA)

            if hasattr(self, 'rep_disease'):
                self.rep_disease.clear()
                self.rep_disease.addItem("All")
                self.rep_disease.addItems(self.CAT_DISEASE)

            # Refresh Diagnosis dropdown
            self.diagnosis.clear()
            if self.CAT_DIAGNOSIS:
                self._set_combo_items_blank_first(self.diagnosis, self.CAT_DIAGNOSIS)

            # ប្តូររចនាបទភ្លាមៗ (Apply theme immediately)
            self.apply_theme(self.current_theme)
            
            # Re-setup autocomplete to reflect changes immediately
            self.view() 
            
            self.statusBar.showMessage(f"បានរក្សាទុក! រចនាបទបច្ចុប្បន្ន៖ {self.current_theme}", 5000)
            QMessageBox.information(self, "Settings Saved", f"ការកំណត់ត្រូវបានរក្សាទុក ហើយរចនាបទត្រូវបានប្តូរទៅជា '{self.current_theme}'")

    def open_license_status(self):
        dialog = LicenseStatusDialog(self)
        dialog.exec_()

    def _run_cloud_helper_action(self, method_name, refresh_after=False):
        helper = LoginDialog()
        helper.hide()
        helper.current_user = self.current_user
        helper.user_context = self.user_context
        helper.branch_code = self.branch_code
        helper.active_branch_code = self.active_branch_code
        helper.backup_dir = self.backup_dir

        try:
            method = getattr(helper, method_name)
            method()
            if refresh_after:
                self.view()
                self.update_next_serial_no()
        finally:
            helper.deleteLater()

    def open_user_audit_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("ពិនិត្យទិន្នន័យតាម User")
        dialog.resize(1180, 720)
        dialog.setStyleSheet("""
            QDialog { background-color: #111820; color: #f5f6fa; }
            QLabel { color: #dfe6e9; font-weight: bold; }
            QComboBox, QDateEdit {
                background-color: #1e272e;
                color: #f5f6fa;
                border: 1px solid #485460;
                border-radius: 5px;
                padding: 6px 8px;
                min-height: 30px;
            }
            QTableWidget {
                background-color: #111820;
                alternate-background-color: #1e272e;
                color: #f5f6fa;
                border: 1px solid #485460;
                gridline-color: #485460;
                selection-background-color: #0fbcf9;
                selection-color: #000000;
            }
            QTableWidget::item { padding: 5px; }
            QHeaderView::section {
                background-color: #0fbcf9;
                color: #000000;
                border: none;
                padding: 8px;
                font-weight: bold;
            }
        """)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        title = QLabel("👤 ពិនិត្យ និងទាញយកទិន្នន័យតាម User")
        title.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        title.setStyleSheet("color: #0fbcf9; font-size: 20px; font-weight: bold;")
        layout.addWidget(title)

        filter_row = QHBoxLayout()
        user_combo = QComboBox()
        user_combo.addItem("All")
        user_combo.addItems(db.get_usernames(self.active_branch_code))
        if not self.is_admin:
            user_combo.setCurrentText(self.current_user)
            user_combo.setEnabled(False)

        start_date = QDateEdit()
        start_date.setCalendarPopup(True)
        start_date.setDisplayFormat("dd/MM/yyyy")
        start_date.setDate(QDate.currentDate())

        end_date = QDateEdit()
        end_date.setCalendarPopup(True)
        end_date.setDisplayFormat("dd/MM/yyyy")
        end_date.setDate(QDate.currentDate())

        btn_search = QPushButton("🔍 ស្វែងរក")
        btn_search.setMinimumHeight(38)
        btn_search.setStyleSheet("background-color: #0fbcf9; color: black; font-weight: bold; border-radius: 5px; padding: 8px 18px;")
        btn_export = QPushButton("📊 ទាញយក Excel")
        btn_export.setMinimumHeight(38)
        btn_export.setStyleSheet("background-color: #20bf6b; color: white; font-weight: bold; border-radius: 5px; padding: 8px 18px;")

        filter_row.addWidget(QLabel("User:"))
        filter_row.addWidget(user_combo, 2)
        filter_row.addWidget(QLabel("ពីថ្ងៃ:"))
        filter_row.addWidget(start_date, 1)
        filter_row.addWidget(QLabel("ដល់ថ្ងៃ:"))
        filter_row.addWidget(end_date, 1)
        filter_row.addWidget(btn_search)
        filter_row.addWidget(btn_export)
        layout.addLayout(filter_row)

        summary = QLabel("សរុប: 0 | ត្រឹមត្រូវ: 0 | ត្រូវពិនិត្យ: 0")
        summary.setMinimumHeight(42)
        summary.setAlignment(Qt.AlignCenter)  # type: ignore[attr-defined]
        summary.setStyleSheet("background-color: #18212b; border: 1px solid #0fbcf9; border-radius: 6px; color: #f5f6fa; font-size: 15px;")
        layout.addWidget(summary)

        table = QTableWidget()
        table.setColumnCount(13)
        table.setHorizontalHeaderLabels([
            "ល.រ", "ថ្ងៃ", "លេខ", "ឈ្មោះ", "ភេទ", "អាយុ", "តំបន់",
            "ជំងឺ", "រោគវិនិច្ឆ័យ", "សេវា", "User", "ស្ថានភាព", "ចំណុចត្រូវពិនិត្យ"
        ])
        table.setEditTriggers(QTableWidget.NoEditTriggers)  # type: ignore[attr-defined]
        table.setAlternatingRowColors(True)
        table.setSelectionBehavior(QTableWidget.SelectRows)  # type: ignore[attr-defined]
        table.verticalHeader().setVisible(False)
        table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        table.horizontalHeader().setStretchLastSection(True)
        layout.addWidget(table, 1)

        audit_state = {"records": [], "result": None, "user": "", "start": "", "end": ""}

        def run_search():
            if start_date.date() > end_date.date():
                QMessageBox.warning(dialog, "ថ្ងៃមិនត្រឹមត្រូវ", "ថ្ងៃចាប់ផ្តើម មិនអាចធំជាង ថ្ងៃបញ្ចប់បានទេ។")
                return

            selected_user = user_combo.currentText()
            result = db.get_user_patient_audit(
                selected_user,
                start_date.date().toPyDate(),
                end_date.date().toPyDate(),
                self.active_branch_code
            )
            records = result["records"]
            audit_state["records"] = records
            audit_state["result"] = result
            audit_state["user"] = selected_user
            audit_state["start"] = start_date.date().toString("dd/MM/yyyy")
            audit_state["end"] = end_date.date().toString("dd/MM/yyyy")
            summary.setText(
                f"User: {selected_user} | សរុប: {result['total']} | "
                f"ត្រឹមត្រូវ: {result['valid']} | ត្រូវពិនិត្យ: {result['issues']}"
            )
            table.setRowCount(len(records))
            for row_idx, record in enumerate(records):
                values = [
                    row_idx + 1,
                    record["date"],
                    record["serial_no"],
                    record["name"],
                    record["sex"],
                    record["age"],
                    record["area"],
                    record["disease"],
                    record["diagnosis"],
                    record["service"],
                    record["created_by"],
                    record["status"],
                    "; ".join(record["issues"]),
                ]
                for col_idx, value in enumerate(values):
                    item = QTableWidgetItem(str(value or ""))
                    if record["issues"]:
                        item.setBackground(QColor("#5c2f24"))
                    table.setItem(row_idx, col_idx, item)
            self.statusBar.showMessage(f"User audit loaded: {len(records)} records", 4000)

        def export_audit():
            records = audit_state.get("records") or []
            if not records:
                QMessageBox.information(dialog, "Export", "សូមស្វែងរកទិន្នន័យជាមុនសិន។")
                return
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            wb = Workbook()
            ws = wb.active
            ws.title = "User Audit"
            result = audit_state.get("result") or {"total": len(records), "valid": 0, "issues": 0}
            selected_user = audit_state.get("user") or user_combo.currentText()
            start_text = audit_state.get("start") or start_date.date().toString("dd/MM/yyyy")
            end_text = audit_state.get("end") or end_date.date().toString("dd/MM/yyyy")

            ws.merge_cells("A1:P1")
            ws["A1"] = "របាយការណ៍ពិនិត្យទិន្នន័យតាម User"
            ws["A1"].font = Font(bold=True, size=16, color="0FBCF9")
            ws["A1"].alignment = Alignment(horizontal="center")
            ws["A2"] = "User"
            ws["B2"] = selected_user
            ws["D2"] = "ចន្លោះថ្ងៃ"
            ws["E2"] = f"{start_text} - {end_text}"
            ws["A3"] = "សរុប"
            ws["B3"] = result["total"]
            ws["D3"] = "ត្រឹមត្រូវ"
            ws["E3"] = result["valid"]
            ws["G3"] = "ត្រូវពិនិត្យ"
            ws["H3"] = result["issues"]
            for cell_ref in ("A2", "D2", "A3", "D3", "G3"):
                ws[cell_ref].font = Font(bold=True)

            headers = [
                "No", "Date", "Serial", "Name", "Sex", "Age", "Area", "Disease",
                "Diagnosis", "Service", "Created By", "Created At", "Updated By",
                "Updated At", "Status", "Issues"
            ]
            ws.append([])
            ws.append(headers)
            header_fill = PatternFill("solid", fgColor="0FBCF9")
            header_row = 5
            for cell in ws[header_row]:
                cell.font = Font(bold=True, color="000000")
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
            issue_fill = PatternFill("solid", fgColor="FADBD8")
            for index, record in enumerate(records, 1):
                ws.append([
                    index, record["date"], record["serial_no"], record["name"],
                    record["sex"], record["age"], record["area"], record["disease"],
                    record["diagnosis"], record["service"], record["created_by"],
                    record["created_at"], record["updated_by"], record["updated_at"],
                    record["status"], "; ".join(record["issues"])
                ])
                if record["issues"]:
                    for cell in ws[ws.max_row]:
                        cell.fill = issue_fill
            for column_cells in ws.columns:
                max_length = max(len(str(cell.value or "")) for cell in column_cells)
                ws.column_dimensions[column_cells[0].column_letter].width = min(max(max_length + 2, 10), 36)
            ws.freeze_panes = "A6"

            safe_user = re.sub(r"[^A-Za-z0-9_-]", "_", selected_user or "All")
            safe_start = start_text.replace("/", "-")
            safe_end = end_text.replace("/", "-")
            filename = f"UserAudit_{safe_user}_{safe_start}_to_{safe_end}.xlsx"
            save_excel_report(dialog, wb, filename, "Save User Audit Excel", "បាននាំចេញរបាយការណ៍ User Audit ទៅ Excel ជោគជ័យ!")

        btn_search.clicked.connect(run_search)
        btn_export.clicked.connect(export_audit)
        run_search()
        dialog.exec_()

    def generate_advanced_report(self):
        report_handler.generate_advanced_report(self)

    def generate_analytics_report(self):
        """Generate analytics report with Group By and Sub-Group By functionality"""
        try:
            from collections import defaultdict

            # Get filter criteria
            start_date = self.rep_start_date.date().toPyDate()
            end_date = self.rep_end_date.date().toPyDate()
            p_type = self.rep_patient_type.currentText()

            # Get all data
            all_rows = db.advanced_search(
                start_date=start_date, end_date=end_date,
                sex=self.rep_sex.currentText(),
                disease=self.rep_disease.currentText(),
                area=self.rep_area.currentText(),
                p_type=p_type,
                branch_code=self.active_branch_code
            ) or []

            # Get Group By and Sub-Group By selections
            group_selection = self.group_by_combo.currentIndex()
            sub_group_selection = self.sub_group_by_combo.currentIndex()
            sort_selection = self.sort_by_combo.currentIndex()

            # Check if we have sub-grouping
            has_sub_group = sub_group_selection > 0  # Index 0 is "None"

            if has_sub_group:
                # Two-level grouping: Group → Sub-Group
                grouped_data = defaultdict(lambda: defaultdict(lambda: {'total': 0, 'male': 0, 'female': 0}))

                for row in all_rows:
                    # Extract main group key
                    group_key = self._get_group_key(row, group_selection)
                    
                    # Extract sub-group key
                    sub_group_key = self._get_group_key(row, sub_group_selection - 1)  # -1 because index 0 is "None"
                    
                    # Count
                    grouped_data[group_key][sub_group_key]['total'] += 1
                    if row[db.PatientCol.SEX] == 'ប្រុស':
                        grouped_data[group_key][sub_group_key]['male'] += 1
                    elif row[db.PatientCol.SEX] == 'ស្រី':
                        grouped_data[group_key][sub_group_key]['female'] += 1

                # Convert to flat list for table display
                analytics_list = []
                for group_key, sub_groups in grouped_data.items():
                    group_total = 0
                    group_male = 0
                    group_female = 0
                    
                    for sub_key, data in sub_groups.items():
                        analytics_list.append({
                            'group': f"{group_key} → {sub_key}",
                            'total': data['total'],
                            'male': data['male'],
                            'female': data['female'],
                            'percentage': 0
                        })
                        group_total += data['total']
                        group_male += data['male']
                        group_female += data['female']
                    
                    # Add group total row
                    analytics_list.append({
                        'group': f"━━ {group_key} (Total)",
                        'total': group_total,
                        'male': group_male,
                        'female': group_female,
                        'percentage': 0,
                        'is_total': True
                    })
            else:
                # Single-level grouping (original behavior)
                grouped_data = defaultdict(lambda: {'total': 0, 'male': 0, 'female': 0})

                for row in all_rows:
                    group_key = self._get_group_key(row, group_selection)
                    
                    grouped_data[group_key]['total'] += 1
                    if row[db.PatientCol.SEX] == 'ប្រុស':
                        grouped_data[group_key]['male'] += 1
                    elif row[db.PatientCol.SEX] == 'ស្រី':
                        grouped_data[group_key]['female'] += 1

                # Convert to list
                analytics_list = []
                for key, data in grouped_data.items():
                    analytics_list.append({
                        'group': key,
                        'total': data['total'],
                        'male': data['male'],
                        'female': data['female'],
                        'percentage': 0
                    })

            # Calculate percentage
            total_count = sum(item['total'] for item in analytics_list if not item.get('is_total'))
            for item in analytics_list:
                if total_count > 0 and not item.get('is_total'):
                    item['percentage'] = round((item['total'] / total_count) * 100, 1)
                elif item.get('is_total'):
                    item['percentage'] = 100.0

            # Sort
            if sort_selection == 0:  # Count Desc
                analytics_list.sort(key=lambda x: x['total'], reverse=True)
            elif sort_selection == 1:  # Count Asc
                analytics_list.sort(key=lambda x: x['total'])
            elif sort_selection == 2:  # Name A-Z
                analytics_list.sort(key=lambda x: x['group'])
            elif sort_selection == 3:  # Name Z-A
                analytics_list.sort(key=lambda x: x['group'], reverse=True)

            # Store data
            self.analytics_data = analytics_list

            # Display in table
            self.display_analytics_table(analytics_list)
            self.update_analytics_summary(analytics_list, total_count)

            # Update count label
            self.lbl_analytics_count.setText(f"ក្រុមសរុប: {len(analytics_list)} | ទិន្នន័យសរុប: {total_count}")

            self.statusBar.showMessage(f"បានបង្កើតរបាយការណ៍វិភាគ: {len(analytics_list)} ក្រុម")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"មិនអាចបង្កើតរបាយការណ៍វិភាគបានទេ: {str(e)}")

    def _get_group_key(self, row, group_selection):
        """Helper function to extract group key from row based on selection"""
        if group_selection == 0:  # Area
            area_raw = str(row[db.PatientCol.AREA])
            return area_raw.split("::")[0] if "::" in area_raw else area_raw
        elif group_selection == 1:  # Sex
            return str(row[db.PatientCol.SEX])
        elif group_selection == 2:  # Disease Case
            disease_raw = str(row[db.PatientCol.DISEASE])
            return disease_raw.split("::")[0] if "::" in disease_raw else disease_raw
        elif group_selection == 3:  # Age Groups
            age_raw = str(row[db.PatientCol.AGE])
            return age_raw.split("::")[0] if "::" in age_raw else age_raw
        elif group_selection == 4:  # Service
            return str(row[db.PatientCol.SERVICE])
        elif group_selection == 5:  # Patient Type
            return str(row[db.PatientCol.TYPE])
        elif group_selection == 6:  # Diagnosis
            return str(row[db.PatientCol.DIAGNOSIS])
        else:
            return "Unknown"

    def display_analytics_table(self, data):
        """Display analytics data in table"""
        self.analytics_table.setColumnCount(5)
        self.analytics_table.setHorizontalHeaderLabels([
            "ក្រុម (Group)", "ចំនួនសរុប (Total)", "ប្រុស (Male)", "ស្រី (Female)", "ភាគរយ (%)"
        ])
        self.analytics_table.setRowCount(len(data))

        for i, item in enumerate(data):
            # Create items with styling for total rows
            group_item = QTableWidgetItem(item['group'])
            total_item = QTableWidgetItem(str(item['total']))
            male_item = QTableWidgetItem(str(item['male']))
            female_item = QTableWidgetItem(str(item['female']))
            pct_item = QTableWidgetItem(f"{item['percentage']}%")
            
            # Highlight total rows
            if item.get('is_total'):
                from PyQt5.QtGui import QFont, QColor
                font = QFont()
                font.setBold(True)
                for table_item in [group_item, total_item, male_item, female_item, pct_item]:
                    table_item.setFont(font)
                    table_item.setBackground(QColor("#00cec9"))  # Teal background for totals
            
            self.analytics_table.setItem(i, 0, group_item)
            self.analytics_table.setItem(i, 1, total_item)
            self.analytics_table.setItem(i, 2, male_item)
            self.analytics_table.setItem(i, 3, female_item)
            self.analytics_table.setItem(i, 4, pct_item)

    def update_analytics_summary(self, data, total_count):
        """Update summary labels for analytics results."""
        visible_rows = [item for item in data if not item.get('is_total')]
        total_male = sum(item['male'] for item in visible_rows)
        total_female = sum(item['female'] for item in visible_rows)
        top_group = max(visible_rows, key=lambda item: item['total'])['group'] if visible_rows else "-"

        self.lbl_summary_total.setText(f"ទិន្នន័យសរុប: {total_count}")
        self.lbl_summary_male.setText(f"ប្រុស: {total_male}")
        self.lbl_summary_female.setText(f"ស្រី: {total_female}")
        self.lbl_summary_top.setText(f"ក្រុមខ្ពស់បំផុត: {top_group}")

    def export_analytics_report(self):
        """Export analytics report to Excel"""
        if not self.analytics_data:
            QMessageBox.warning(self, "Warning", "សូមបង្កើតរបាយការណ៍វិភាគជាមុនសិន!")
            return

        try:
            wb, ws, styles = setup_excel_report_workbook(
                "Analytics Report",
                "📊 របាយការណ៍វិភាគ (Analytics Report)",
                "FF00CEC9",
                "FF00CEC9",
                {"A": 25, "B": 15, "C": 15, "D": 15, "E": 15},
            )
            row = 3
            add_excel_headers(ws, row, ["ក្រុម (Group)", "ចំនួនសរុប", "ប្រុស", "ស្រី", "ភាគរយ"], styles)
            row += 1
            for item in self.analytics_data:
                ws.cell(row=row, column=1, value=item['group']).alignment = styles["center_align"]  # type: ignore
                ws.cell(row=row, column=2, value=item['total']).alignment = styles["center_align"]  # type: ignore
                ws.cell(row=row, column=3, value=item['male']).alignment = styles["center_align"]  # type: ignore
                ws.cell(row=row, column=4, value=item['female']).alignment = styles["center_align"]  # type: ignore
                ws.cell(row=row, column=5, value=f"{item['percentage']}%").alignment = styles["center_align"]  # type: ignore
                row += 1

            from datetime import datetime
            filename = f"Analytics_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
            save_excel_report(self, wb, filename, "Save Analytics Excel", "បាននាំចេញរបាយការណ៍វិភាគទៅ Excel ជោគជ័យ!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"ការនាំចេញបរាជ័យ: {str(e)}")

    def clear_analytics_report(self):
        """Clear analytics report"""
        self.rep_start_date.setDate(QDate.currentDate().addMonths(-1))
        self.rep_end_date.setDate(QDate.currentDate())
        self.rep_sex.setCurrentIndex(0)
        self.rep_area.setCurrentIndex(0)
        self.rep_disease.setCurrentIndex(0)
        self.rep_patient_type.setCurrentIndex(0)
        self.group_by_combo.setCurrentIndex(0)
        self.sub_group_by_combo.setCurrentIndex(0)  # Reset sub-group
        self.sort_by_combo.setCurrentIndex(0)
        self.analytics_table.setRowCount(0)
        self.analytics_data.clear()
        self.lbl_analytics_count.setText("ក្រុមសរុប: 0")
        self.update_analytics_summary([], 0)

        self.statusBar.showMessage("បានលុបការកំណត់រួចរាល់", 3000)

    def open_advanced_query(self):
        """Open Advanced Query Dialog for multi-criteria search"""
        dialog = AdvancedQueryDialog(self)
        dialog.exec_()

    def show_analytics_context_menu(self, position):
        """Show context menu for analytics table"""
        row = self.analytics_table.rowAt(position.y())
        if row < 0:
            return
        
        menu = QMenu(self)
        view_details_action = menu.addAction("🔍 មើលលម្អិត (View Details)")
        export_group_action = menu.addAction("📊 Export ក្រុមនេះ")
        
        action = menu.exec_(self.analytics_table.mapToGlobal(position))
        
        if action == view_details_action:
            self.drill_down_details()
        elif action == export_group_action:
            self.export_selected_group(row)

    def drill_down_details(self):
        """Show detailed patient list for selected group"""
        selected_items = self.analytics_table.selectedItems()
        if not selected_items:
            QMessageBox.information(self, "Info", "សូមជ្រើសរើសក្រុមដែលអ្នកចង់មើលលម្អិត!")
            return

        row = selected_items[0].row()
        item = self.analytics_table.item(row, 0)
        if not item:
            return
        group_name = item.text()
        
        # Skip total rows
        if "(Total)" in group_name:
            QMessageBox.information(self, "Info", "សូមជ្រើសរើសក្រុមរង (មិនមែន Total) ដើម្បីមើលអ្នកជំងឺ!")
            return

        # Get the group data
        if row < len(self.analytics_data):
            group_data = self.analytics_data[row]

            # Create drill-down dialog
            drill_dialog = QDialog(self)
            drill_dialog.setWindowTitle(f"📋 លម្អិត: {group_name}")
            drill_dialog.setMinimumSize(1000, 600)

            layout = QVBoxLayout(drill_dialog)

            # Info label
            info_label = QLabel(f"<h3>📊 ក្រុម: {group_name}</h3>"
                              f"<p>ចំនួនសរុប: <b>{group_data['total']}</b> នាក់ | "
                              f"ប្រុស: <b>{group_data['male']}</b> | "
                              f"ស្រី: <b>{group_data['female']}</b> | "
                              f"ភាគរយ: <b>{group_data['percentage']}%</b></p>")
            info_label.setStyleSheet("background-color: #1e272e; padding: 10px; border-radius: 5px;")
            layout.addWidget(info_label)

            # Extract filter criteria from group name
            filters = self._parse_group_filters(group_name)

            # Get filtered patient data
            filtered_patients = self._get_patients_by_filters(filters)

            # Display patient list
            patient_table = QTableWidget()
            patient_table.setColumnCount(7)
            patient_table.setHorizontalHeaderLabels([
                "លេខរៀង", "ឈ្មោះ", "ភេទ", "អាយុ", "តំបន់", "ជំងឺ", "រោគវិនិច្ឆ័យ"
            ])
            patient_table.setRowCount(len(filtered_patients))
            patient_table.setEditTriggers(QTableWidget.NoEditTriggers)  # type: ignore
            
            for i, patient in enumerate(filtered_patients):
                patient_table.setItem(i, 0, QTableWidgetItem(str(patient.get('serial', ''))))
                patient_table.setItem(i, 1, QTableWidgetItem(str(patient.get('name', ''))))
                patient_table.setItem(i, 2, QTableWidgetItem(str(patient.get('sex', ''))))
                patient_table.setItem(i, 3, QTableWidgetItem(str(patient.get('age', ''))))
                patient_table.setItem(i, 4, QTableWidgetItem(str(patient.get('area', ''))))
                patient_table.setItem(i, 5, QTableWidgetItem(str(patient.get('disease', ''))))
                patient_table.setItem(i, 6, QTableWidgetItem(str(patient.get('diagnosis', ''))))
            
            patient_table.setSortingEnabled(True)  # type: ignore
            layout.addWidget(patient_table)

            # Close button
            close_btn = QPushButton("បិទ (Close)")
            close_btn.setStyleSheet("background-color: #e74c3c; color: white; padding: 10px; font-weight: bold;")
            close_btn.clicked.connect(drill_dialog.accept)
            layout.addWidget(close_btn)

            drill_dialog.exec_()

    def _parse_group_filters(self, group_name):
        """
        Parse group name to extract filter criteria
        Example: "ក → ប្រុស" → {'area': 'ក', 'sex': 'ប្រុស'}
        Example: "ថ្មី" → {'disease': 'ថ្មី'}
        """
        filters = {}
        
        # Check for two-level grouping (A → B)
        if " → " in group_name:
            parts = group_name.split(" → ", 1)
            main_part = parts[0].strip()
            sub_part = parts[1].strip()
            
            # Try to match main part with known categories
            if main_part in self.CAT_AREA:
                filters['area'] = main_part
            elif main_part in self.CAT_SEX:
                filters['sex'] = main_part
            elif main_part in self.CAT_DISEASE:
                filters['disease'] = main_part
            else:
                filters['diagnosis'] = main_part
            
            # Try to match sub part
            if sub_part in self.CAT_AREA:
                filters['area'] = sub_part
            elif sub_part in self.CAT_SEX:
                filters['sex'] = sub_part
            elif sub_part in self.CAT_DISEASE:
                filters['disease'] = sub_part
            else:
                filters['diagnosis'] = sub_part
        else:
            # Single-level grouping
            if group_name in self.CAT_AREA:
                filters['area'] = group_name
            elif group_name in self.CAT_SEX:
                filters['sex'] = group_name
            elif group_name in self.CAT_DISEASE:
                filters['disease'] = group_name
            else:
                filters['diagnosis'] = group_name
        
        return filters

    def _get_patients_by_filters(self, filters):
        """
        Get patient list based on filters
        Returns list of dicts with patient info
        """
        # Get patients visible to this user's branch.
        ptype = self.rep_patient_type.currentText()
        if ptype == "All":
            all_patients = db.view(self.active_branch_code) or []
        else:
            all_patients = db.view_by_patient_type(ptype, self.active_branch_code) or []
        
        filtered = []
        for row in all_patients:
            match = True
            
            # Apply area filter
            if 'area' in filters:
                area_raw = str(row[db.PatientCol.AREA])
                area_cat = area_raw.split("::")[0] if "::" in area_raw else area_raw
                if area_cat != filters['area']:
                    match = False
            
            # Apply sex filter
            if 'sex' in filters and row[db.PatientCol.SEX] != filters['sex']:
                match = False
            
            # Apply disease filter
            if 'disease' in filters:
                disease_raw = str(row[db.PatientCol.DISEASE])
                disease_cat = disease_raw.split("::")[0] if "::" in disease_raw else disease_raw
                if disease_cat != filters['disease']:
                    match = False

            # Apply diagnosis filter
            if 'diagnosis' in filters and row[db.PatientCol.DIAGNOSIS] != filters['diagnosis']:
                match = False
            
            if match:
                filtered.append({
                    'serial': row[db.PatientCol.SERIAL],
                    'name': row[db.PatientCol.NAME],
                    'sex': row[db.PatientCol.SEX],
                    'age': row[db.PatientCol.AGE],
                    'area': row[db.PatientCol.AREA],
                    'disease': row[db.PatientCol.DISEASE],
                    'diagnosis': row[db.PatientCol.DIAGNOSIS]
                })
        
        return filtered

    def export_selected_group(self, row):
        """Export selected group data"""
        if row < len(self.analytics_data):
            group_data = self.analytics_data[row]
            group_name = group_data['group']
            
            QMessageBox.information(self, "Export", 
                f"មុខងារនេះនឹងមានពេលក្រោយ!\n\n"
                f"ក្រុម: {group_name}\n"
                f"ចំនួនសរុប: {group_data['total']} នាក់")

    def clear_advanced_report(self):
        self.rep_start_date.setDate(QDate.currentDate().addMonths(-1))
        self.rep_end_date.setDate(QDate.currentDate())
        self.rep_sex.setCurrentIndex(0)
        self.rep_disease.setCurrentIndex(0)
        self.rep_area.setCurrentIndex(0)
        self.rep_patient_type.setCurrentIndex(0)
        self.rep_table.setRowCount(0)
        self.filtered_report_rows.clear()
        self.lbl_report_count.setText("Total Found: 0")
        report_handler.update_report_chart(self, {})

    def export_excel(self):
        excel_handler.export_generic_excel(self)

    def export_filtered_report(self):
        excel_handler.export_generic_excel(self, rows_to_export=self.filtered_report_rows, export_type_prompt=False)

    def export_statistics_to_excel(self):
        """Export statistics report to Excel"""
        try:
            wb, ws, styles = setup_excel_report_workbook(
                "Statistics Report",
                "📊 របាយការណ៍ស្ថិតិអ្នកជំងឺ (Detailed Statistics)",
                "FF0FBCF9",
                "FFE74C3C",
                {"A": 35, "B": 15, "C": 15, "D": 15, "E": 15},
            )
            s = db.get_statistics(self.active_branch_code)
            c = s['child']
            a = s['adult']
            o = s['overall']

            row = 3
            ws.merge_cells(f'A{row}:E{row}')  # type: ignore
            ws[f'A{row}'] = "🔵 ផ្នែកទី ១: ស្ថិតិសរុបទាំងអស់ (Overall Total)"  # type: ignore
            ws[f'A{row}'].font = styles["section_font"]  # type: ignore
            ws[f'A{row}'].fill = styles["section_fill"]  # type: ignore

            row += 1
            add_excel_headers(ws, row, ["ពិពណ៌នា", "សរុប", "ប្រុស", "ស្រី", "ភាគរយ"], styles)

            row += 1
            data = [
                ["អ្នកជំងឺសរុប", o['total'], o['male'], o['female'], "100%"],
                ["ថ្ងៃនេះ", o['today_count'], "-", "-", f"{self.calc_pct(o['today_count'], o['total'])}%"],
                ["ករណីថ្មី", c['new_total'] + a['new_total'], c['new_m'] + a['new_m'], c['new_f'] + a['new_f'], f"{self.calc_pct(c['new_total'] + a['new_total'], o['total'])}%"],
                ["ករណីចាស់", c['old_total'] + a['old_total'], c['old_m'] + a['old_m'], c['old_f'] + a['old_f'], f"{self.calc_pct(c['old_total'] + a['old_total'], o['total'])}%"],
            ]

            for item in data:
                for col, value in enumerate(item, 1):
                    ws.cell(row=row, column=col, value=value).alignment = styles["center_align"]  # type: ignore
                row += 1

            from datetime import datetime
            filename = f"Statistics_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
            save_excel_report(self, wb, filename, "Save Statistics Excel", "បាននាំចេញរបាយការណ៍ស្ថិតិទៅ Excel ជោគជ័យ!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"ការនាំចេញបរាជ័យ: {str(e)}")

    def calc_pct(self, part, total):
        """Calculate percentage"""
        if total == 0:
            return 0
        return round((part / total) * 100, 1)

    def export_pdf(self):
        try:
            file_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", "", "PDF Files (*.pdf)")
            if file_path:
                c = canvas.Canvas(file_path, pagesize=letter)
                width, height = letter
                
                # Function to draw headers (for reuse on new pages)
                def draw_header(y_pos):
                    c.setFont("Helvetica-Bold", 16)
                    c.drawString(30, y_pos, "Clinic Patient Report")
                    y_pos -= 30
                    
                    headers = ["Date", "ID", "Name", "Age", "Sex", "Service"] 
                    x_coords_h = [30, 80, 200, 250, 310, 400]
                    
                    c.setFont("Helvetica-Bold", 10)
                    for i, h in enumerate(headers):
                        c.drawString(x_coords_h[i], y_pos, h)
                    return y_pos - 20

                y = height - 50
                y = draw_header(y)
                
                c.setFont("Helvetica", 10)
                x_coords = [30, 80, 200, 250, 310, 400]

                # Get ALL patients for PDF export (not filtered by type)
                rows = db.view(self.active_branch_code) or []
                for row in rows:
                    # Clean up Age (row[6]) which might be "Category::Value"
                    age_raw = str(row[6])
                    age_display = age_raw
                    if "::" in age_raw:
                        parts = age_raw.split("::")
                        # Show Value if present, else Category
                        age_display = parts[1] if len(parts) > 1 and parts[1].strip() else parts[0]

                    # Show specific columns for PDF: Date(1), ID(0), Name(4), Age(Cleaned), Sex(7), Service(21)
                    display_data = [row[1], row[0], row[4], age_display, row[7], row[21]]
                    
                    for i, item in enumerate(display_data):
                        c.drawString(x_coords[i], y, str(item))
                    y -= 15
                    
                    if y < 50: # New page
                        c.showPage()
                        y = height - 50
                        y = draw_header(y)
                        c.setFont("Helvetica", 10) # Reset font for data
                
                c.save()
                self.statusBar.showMessage(f"PDF report generated: {os.path.basename(file_path)}", 7000)
                # QMessageBox.information(self, "Success", "Report PDF generated!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not generate PDF: {str(e)}")

    def print_patient_card(self):
        if not self.selected_id:
            QMessageBox.warning(self, "Warning", "Please select a patient to print.")
            return

        # Save Dialog - Word format
        filename = f"Patient_{self.serial_no.text()}.docx"
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Patient Card", filename, "Word Files (*.docx)")
        if not file_path:
            return

        try:
            # Create Word document
            doc = Document()
            
            # Set Khmer font style
            style = doc.styles['Normal']
            font = style.font  # type: ignore
            font.name = 'Khmer OS'
            font.size = Pt(11)

            # Title
            title = doc.add_heading('ប័ណ្ណសម្គាល់អ្នកជំងឺ', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph('ប្រព័ន្ធគ្រប់គ្រងគ្លីនិក')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(12)
            subtitle.runs[0].font.bold = True

            doc.add_paragraph('_' * 50)

            # Translate service types to Khmer
            service_text = self.service.currentText()
            service_kh = service_text
            if service_text == "HEF":
                service_kh = "ហេហ្វ (HEF)"
            elif service_text == "PAY":
                service_kh = "បង់ប្រាក់ (PAY)"
            elif service_text == "FREE":
                service_kh = "ឥតគិតថ្លៃ (FREE)"
            elif service_text == "E":
                service_kh = "អ៊ី (E)"
            elif service_text in ("HEF-R", "HEFR"):
                service_kh = "ហេហ្វ-អរ (HEF-R)"
            elif service_text in ("HEF-I", "HEFI"):
                service_kh = "ហេហ្វ-អ៊ី (HEF-I)"
            elif service_text == "NSSF-8":
                service_kh = "របបសន្តិសុខសង្គម-៨ (NSSF-8)"
            elif service_text == "NSSF-7":
                service_kh = "របបសន្តិសុខសង្គម-៧ (NSSF-7)"
            elif service_text == "NSSF-A":
                service_kh = "របបសន្តិសុខសង្គម-អា (NSSF-A)"

            # Translate patient type to Khmer
            type_kh = "កុមារ" if self.current_patient_type == "Child" else "មនុស្សធំ"

            # Patient Information Table
            table = doc.add_table(rows=11, cols=2)
            table.style = 'Table Grid'
            
            # Fill table with patient data
            data = [
                ("កាលបរិច្ឆេទ", self.date.text()),
                ("លេខរៀង", self.serial_no.text()),
                ("លេខប័ណ្ណ", self.card_id.text()),
                ("ឈ្មោះ", self.name.text()),
                ("អាយុ", f"{self.age_cat.currentText()} {self.age_val.text()}"),
                ("ភេទ", self.sex.currentText()),
                ("អាស័យដ្ឋាន", self.address.text()),
                ("លេខទូរស័ព្ទ", self.phone.text()),
                ("រោគវិនិច្ឆ័យ", self.diagnosis.currentText()),
                ("ការព្យាបាល", self.treatment.text()),
                ("អាហារូបត្ថម្ភ", f"(ទ/អ)៖ {self.nut_wa.text()} | (ទ/ក)៖ {self.nut_wh.text()}"),
            ]

            for i, (label, value) in enumerate(data):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                # Make labels bold
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[0].width = Cm(5)
                row.cells[1].width = Cm(10)

            # Add remaining info
            doc.add_paragraph(f"ប្រភេទអ្នកជំងឺ៖ {type_kh}")
            doc.add_paragraph(f"សេវាកម្ម៖ {service_kh}")
            doc.add_paragraph(f"សម្គាល់៖ {self.remark.text()}")

            doc.add_paragraph('_' * 50)
            doc.add_paragraph('ឯកសារនេះត្រូវបានបង្កើតដោយស្វ័យប្រវត្តិពីប្រព័ន្ធគ្រប់គ្រងគ្លីនិក')

            # Save document
            doc.save(file_path)
            self.statusBar.showMessage(f"ប័ណ្ណត្រូវបានរក្សាទុក៖ {os.path.basename(file_path)}", 5000)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not create Word document: {str(e)}")

    def import_excel(self):
        excel_handler.import_excel(self)

    def merge_database(self):
        # 1. Choose an external SQLite database file
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Database to Merge", "", "SQLite Files (*.db)")
        if not file_path:
            return

        if os.path.abspath(file_path) == os.path.abspath(db.DB_NAME):
            QMessageBox.warning(self, "Error", "Cannot merge the current database into itself.")
            return

        try:
            added_count, skipped_count = db.merge_database_file(file_path, self.branch_code)
            self.view()
            self.update_next_serial_no()
            self.statusBar.showMessage(
                f"Merge successful. Added: {added_count}, Skipped (Duplicate): {skipped_count}.",
                7000
            )
            msg_box = QMessageBox(self)
            msg_box.setIcon(QMessageBox.Information)
            msg_box.setWindowTitle("Merge Result")
            msg_box.setText("Merge completed successfully.")
            msg_box.setInformativeText(
                f"Added: {added_count}\n"
                f"Skipped (Duplicate): {skipped_count}"
            )
            msg_box.setStyleSheet("""
                QMessageBox {
                    background-color: #1e272e;
                }
                QMessageBox QLabel {
                    color: #ffffff;
                    min-height: 24px;
                    padding: 4px 8px;
                }
                QMessageBox QPushButton {
                    background-color: #0fbcf9;
                    color: #000000;
                    border: none;
                    border-radius: 5px;
                    padding: 8px 28px;
                    min-width: 90px;
                    font-weight: bold;
                }
                QMessageBox QPushButton:hover {
                    background-color: #00a8ff;
                }
            """)
            msg_box.setMinimumWidth(420)
            msg_box.exec_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Merge failed: {str(e)}")

    def backup_database(self):
        try:
            username_safe = re.sub(r"[^A-Za-z0-9_-]", "_", self.current_user or "unknown")
            filename = f"backup_clinic_{username_safe}_{datetime.now().strftime('%Y%m%d_%H%M')}.db"
            save_path, _ = QFileDialog.getSaveFileName(self, "Backup Database", filename, "SQLite Files (*.db)")
            if save_path:
                if self.is_admin:
                    shutil.copy(db.DB_NAME, save_path)
                else:
                    build_patient_share_database(save_path, db.view(self.active_branch_code) or [])
                self.statusBar.showMessage(f"Database backed up successfully to {os.path.basename(save_path)}", 7000)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Backup failed: {str(e)}")

    def share_to_telegram(self):
        """Share database to Telegram for merging"""
        dialog = TelegramShareDialog(self)
        if dialog.exec_():
            share_type = dialog.share_type
            telegram_contact = dialog.telegram_contact
            
            if share_type == "full":
                self.share_full_database(telegram_contact)
            elif share_type == "daily":
                self.share_daily_report(telegram_contact)

    def _get_default_telegram_contact(self):
        """ទាញយក Telegram contact លំនាំដើមពី settings"""
        return self.config.get('CATEGORIES', 'telegram_contact', fallback="@Phunsinouen")

    def _build_patient_share_database(self, target_db_path, patient_rows):
        build_patient_share_database(target_db_path, patient_rows)

    def _validate_telegram_contact(self, contact):
        """ផ្ទៀងផ្ទាត់ទម្រង់ Telegram contact"""
        if not contact or not contact.strip():
            return False, "សូមបញ្ចូល Telegram contact!"

        contact = contact.strip()

        # Username format: @something
        if contact.startswith("@"):
            username = contact[1:]
            if not re.match(r'^[A-Za-z0-9_]{5,32}$', username):
                return False, f"ទម្រង់ Username មិនត្រឹមត្រូវ: @{username}\n\nសូមប្រើតួអក្សរ A-Z, 0-9, _ (៥-៣២ តួអក្សរ)"
            return True, contact

        # Phone number format: +855...
        if contact.startswith("+"):
            if not re.match(r'^\+\d{9,15}$', contact):
                return False, f"ទម្រង់លេខទូរស័ព្ទមិនត្រឹមត្រូវ: {contact}\n\nឧ. +85512345678"
            return True, contact

        # Saved Messages
        if contact.lower() == "saved messages":
            return True, contact

        # Group/Channel name (generic)
        if len(contact) < 3:
            return False, "ឈ្មោះ Group/Channel ខ្លីពេក (យ៉ាងតិច ៣ តួអក្សរ)"

        return True, contact

    def _open_telegram(self, telegram_contact):
        """បើក Telegram ជាមួយ contact"""
        contact_lower = telegram_contact.lower()

        if contact_lower == "saved messages":
            QDesktopServices.openUrl(QUrl("https://web.telegram.org"))
        elif "group" in contact_lower or "channel" in contact_lower:
            QDesktopServices.openUrl(QUrl("https://web.telegram.org"))
        elif telegram_contact.startswith("@"):
            telegram_url = f"https://t.me/{telegram_contact[1:]}"
            QDesktopServices.openUrl(QUrl(telegram_url))
        elif telegram_contact.startswith("+"):
            telegram_url = f"tg://resolve?phone={telegram_contact[1:]}"
            QDesktopServices.openUrl(QUrl(telegram_url))
        else:
            telegram_url = f"https://t.me/{telegram_contact}"
            QDesktopServices.openUrl(QUrl(telegram_url))

    def _open_file_explorer(self, file_path):
        """បើក File Explorer ជាមួយ file (គាំទ្រគ្រប់ OS)"""
        folder_path = os.path.dirname(file_path)

        if sys.platform == "win32":
            os.startfile(folder_path)
        elif sys.platform == "darwin":  # macOS
            subprocess.Popen(["open", folder_path])
        else:  # Linux
            subprocess.Popen(["xdg-open", folder_path])

    def _get_telegram_share_dir(self):
        """
        ទាញយក folder សម្រាប់ដាក់ file ផ្ញើ Telegram
        ប្រើ AppData ជំនួសឱ្យ Program Files (ដែលមិនអាចសរសេរបាន)
        """
        try:
            import ctypes.wintypes

            # ប្រើ AppData\Roaming\ClinicManager\TelegramShares
            CSIDL_APPDATA = 26
            SHGFP_TYPE_CURRENT = 0
            buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
            ctypes.windll.shell32.SHGetFolderPathW(None, CSIDL_APPDATA, None, SHGFP_TYPE_CURRENT, buf)

            share_dir = os.path.join(buf.value, "ClinicManager", "TelegramShares")
            if not os.path.exists(share_dir):
                os.makedirs(share_dir)

            return share_dir
        except Exception:
            # Fallback ទៅ Temp folder
            temp_dir = os.getenv('TEMP') or os.getenv('TMP') or os.getcwd()
            share_dir = os.path.join(temp_dir, "ClinicManager", "TelegramShares")
            if not os.path.exists(share_dir):
                os.makedirs(share_dir)
            return share_dir

    def _open_telegram_url(self, telegram_contact):
        """
        បើក Telegram ជាមួយ fallback ច្រើនជម្រើស
        """
        contact_lower = telegram_contact.lower()

        # កំណត់ URL ត្រឹមត្រូវ
        if contact_lower == "saved messages":
            url = "https://web.telegram.org"
        elif "group" in contact_lower or "channel" in contact_lower:
            url = "https://web.telegram.org"
        elif telegram_contact.startswith("@"):
            url = f"https://t.me/{telegram_contact[1:]}"
        elif telegram_contact.startswith("+"):
            url = f"tg://resolve?phone={telegram_contact[1:]}"
        else:
            url = f"https://t.me/{telegram_contact}"

        # ព្យាយាមបើក Telegram
        try:
            result = QDesktopServices.openUrl(QUrl(url))

            # បើ fail (មិនមាន browser/telegram)
            if not result:
                QMessageBox.warning(
                    self,
                    "⚠️ មិនអាចបើក Telegram បានទេ",
                    f"មិនអាចបើក Telegram បានទេ!\n\n"
                    f"សូមអនុវត្តដោយដៃ៖\n"
                    f"1. បើក Telegram Desktop ឬ https://web.telegram.org\n"
                    f"2. ស្វែងរក {telegram_contact}\n"
                    f"3. អូស file ចូលក្នុង chat\n\n"
                    f"💡 File path ត្រូវបានចម្លងទៅ Clipboard ហើយ!"
                )
        except Exception as e:
            QMessageBox.warning(
                self,
                "⚠️ កំហុស",
                f"មិនអាចបើក Telegram បានទេ: {str(e)}\n\n"
                f"សូមបើកដោយដៃ រួចអូស file ចូល។"
            )

    def _get_valid_telegram_contact_or_none(self, telegram_contact):
        if telegram_contact is None:
            telegram_contact = self._get_default_telegram_contact()
        is_valid, msg = self._validate_telegram_contact(telegram_contact)
        if not is_valid:
            QMessageBox.warning(self, "⚠️ ទម្រង់មិនត្រឹមត្រូវ", msg)
            return None
        return telegram_contact

    def _cleanup_temp_file(self, file_path):
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                pass

    def _zip_temp_database_for_telegram(self, temp_db_path, zip_filename, archive_name):
        share_dir = self._get_telegram_share_dir()
        zip_path = os.path.join(share_dir, zip_filename)
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.write(temp_db_path, archive_name)
        file_size = os.path.getsize(zip_path) / (1024 * 1024)
        return zip_path, file_size

    def _open_prepared_telegram_share(self, zip_path, telegram_contact, instruction_callback):
        copy_to_clipboard(zip_path)
        self._open_telegram_url(telegram_contact)
        QTimer.singleShot(1500, lambda: self._open_file_explorer(zip_path))
        QTimer.singleShot(1700, instruction_callback)

    def share_full_database(self, telegram_contact=None):
        """ផ្ញើ Database ទាំងមូលទៅ Telegram"""
        telegram_contact = self._get_valid_telegram_contact_or_none(telegram_contact)
        if not telegram_contact:
            return


        try:
            # បង្កើត temporary file សម្រាប់ copy database
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            db_path = db.DB_NAME

            # ប្រើ tempfile ដើម្បីឱ្យ auto-cleanup
            temp_db_copy = None
            zip_path = None

            try:
                # បង្កើត temp file សម្រាប់ copy database
                fd, temp_db_copy = tempfile.mkstemp(suffix=".db", prefix="clinic_backup_")
                os.close(fd)
                if self.is_admin:
                    shutil.copy2(db_path, temp_db_copy)
                else:
                    build_patient_share_database(temp_db_copy, db.view(self.active_branch_code) or [])

                # បង្កើត ZIP file ក្នុង AppData (មិនមែន Program Files)
                zip_filename = f"ClinicDB_Full_{timestamp}.zip"
                zip_path, file_size = self._zip_temp_database_for_telegram(
                    temp_db_copy,
                    zip_filename,
                    f"ClinicDB_{timestamp}.db",
                )

                msg = self._build_telegram_ready_message(
                    "✅ ឯកសារ Database ត្រូវបានគណនារួចរាល់!",
                    zip_filename,
                    file_size,
                    zip_path,
                    telegram_contact,
                )

                reply = QMessageBox.information(self, "✅ Share Complete", msg,
                                               QMessageBox.Ok | QMessageBox.Cancel)

                if reply == QMessageBox.Ok:
                    self._open_prepared_telegram_share(
                        zip_path,
                        telegram_contact,
                        lambda: self._show_send_instructions(zip_path, file_size, telegram_contact),
                    )

            finally:
                # សម្អា់ temp file ជានិច្ច
                self._cleanup_temp_file(temp_db_copy)

        except PermissionError:
            QMessageBox.critical(self, "❌ កំហុស",
                "មិនអាចចម្លង Database បានទេ!\n\n"
                "សូមបិទកម្មវិធីផ្សេងៗដែលកំពុងប្រើ Database រួចសាកល្បងម្តងទៀត។")
        except Exception as e:
            QMessageBox.critical(self, "❌ កំហុស", f"មិនអាចផ្ញើ Database បានទេ: {str(e)}")

    def _build_telegram_ready_message(self, heading, zip_filename, file_size, zip_path, telegram_contact, extra_lines=None):
        lines = [
            heading,
            "",
            *(extra_lines or []),
            f"📦 ឈ្មោះ: {zip_filename}",
            f"📊 ទំហំ: {file_size:.2f} MB",
            "",
            f"📂 ទីតាំង: {zip_path}",
            "",
            f"📤 ផ្ញើទៅ: {telegram_contact}",
            "",
            "សូមអនុវត្តតាមជំហានខាងក្រោម:",
            "1. ចុច OK ដើម្បីបើក Telegram",
            f"2. ស្វែងរក {telegram_contact}",
            "3. អូស file ចូលក្នុង Telegram",
            "",
            "💡 ឬចម្លង file ពី folder ដោយផ្ទាល់។",
        ]
        return "\n".join(lines)

    def _show_telegram_share_instructions(self, title, zip_path, file_size, telegram_contact, extra_lines=None):
        instructions = QMessageBox(self)
        instructions.setIcon(QMessageBox.Information)
        instructions.setWindowTitle(title)
        detail_lines = [
            f"📂 ឈ្មោះ file: {os.path.basename(zip_path)}",
            f"📊 ទំហំ: {file_size:.2f} MB",
        ]
        detail_lines.extend(extra_lines or [])
        detail_lines.append(f"📤 ផ្ញើទៅ: {telegram_contact}")
        instructions.setText(
            "✅ File path ត្រូវបានចម្លងទៅ Clipboard ហើយ!\n\n"
            "📋 វិធីផ្ញើ៖\n"
            "1. ចុច Ctrl + V ដើម្បី Paste file path ចូល Telegram\n"
            "2. ឬអូស file ពី folder ចូលក្នុង Telegram\n\n"
            + "\n".join(detail_lines)
        )
        instructions.addButton("យល់ហើយ 👍", QMessageBox.AcceptRole)
        instructions.exec_()

    def _show_send_instructions(self, zip_path, file_size, telegram_contact):
        self._show_telegram_share_instructions("📤 ជំហានផ្ញើ File", zip_path, file_size, telegram_contact)

    def share_daily_report(self, telegram_contact=None):
        """ផ្ញើរបាយការណ៍ប្រចាំថ្ងៃទៅ Telegram"""
        telegram_contact = self._get_valid_telegram_contact_or_none(telegram_contact)
        if not telegram_contact:
            return


        try:
            today = datetime.now().strftime("%d/%m/%Y")

            # ទាញយកអ្នកជំងឺថ្ងៃនេះ
            branch_clause = ""
            branch_params = ()
            if self.active_branch_code:
                branch_clause = " AND (branch_code = ? OR branch_code IS NULL OR branch_code = '')"
                branch_params = (self.active_branch_code,)
            today_patients = db.execute_read(
                f"SELECT * FROM patient WHERE date=?{branch_clause}",
                (today, *branch_params)
            ) or []

            if not today_patients:
                QMessageBox.information(self, "📭 គ្មានទិន្នន័យ",
                    f"មិនមានទិន្នន័យសម្រាប់ថ្ងៃ {today} ទេ។")
                return

            # ប្រើ tempfile សម្រាប់ temporary database
            temp_db_path = None
            zip_path = None

            try:
                # បង្កើត temp database
                fd, temp_db_path = tempfile.mkstemp(suffix=".db", prefix="clinic_daily_")
                os.close(fd)

                # ចម្លង structure + ទិន្នន័យថ្ងៃនេះ
                self._build_patient_share_database(temp_db_path, today_patients)
                # បង្កើត ZIP file ក្នុង AppData (មិនមែន Program Files)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                zip_filename = f"ClinicDB_Daily_{today.replace('/', '')}_{timestamp}.zip"
                zip_path, file_size = self._zip_temp_database_for_telegram(
                    temp_db_path,
                    zip_filename,
                    f"ClinicDB_Daily_{today.replace('/', '')}.db",
                )

                msg = self._build_telegram_ready_message(
                    "✅ របាយការណ៍ប្រចាំថ្ងៃត្រូវបានគណនារួចរាល់!",
                    zip_filename,
                    file_size,
                    zip_path,
                    telegram_contact,
                    extra_lines=[
                        f"📅 ថ្ងៃ: {today}",
                        f"👥 ចំនួនអ្នកជំងឺ: {len(today_patients)} នាក់",
                    ],
                )

                reply = QMessageBox.information(self, "✅ Daily Report Ready", msg,
                                               QMessageBox.Ok | QMessageBox.Cancel)

                if reply == QMessageBox.Ok:
                    self._open_prepared_telegram_share(
                        zip_path,
                        telegram_contact,
                        lambda: self._show_daily_report_instructions(zip_path, file_size, len(today_patients), telegram_contact),
                    )

            finally:
                # សម្អា់ temp file ជានិច្ច
                self._cleanup_temp_file(temp_db_path)

        except PermissionError:
            QMessageBox.critical(self, "❌ កំហុស",
                "មិនអាចចម្លង Database បានទេ!\n\n"
                "សូមបិទកម្មវិធីផ្សេងៗដែលកំពុងប្រើ Database រួចសាកល្បងម្តងទៀត។")
        except Exception as e:
            QMessageBox.critical(self, "❌ កំហុស", f"មិនអាចបង្កើតរបាយការណ៍ប្រចាំថ្ងៃបានទេ: {str(e)}")

    def _show_daily_report_instructions(self, zip_path, file_size, patient_count, telegram_contact):
        self._show_telegram_share_instructions(
            "📤 ជំហានផ្ញើរបាយការណ៍",
            zip_path,
            file_size,
            telegram_contact,
            [f"👥 ចំនួនអ្នកជំងឺ: {patient_count} នាក់"],
        )
    def renumber_all_serials(self):
        db.renumber_all_patient_numbering()
        self.update_next_serial_no()

    def update_next_serial_no(self):
        if self.selected_id:
            return
        max_serial = db.get_max_serial_for_type(self.current_patient_type, self.date.text(), self.active_branch_code)
        prefix = 'C-' if self.current_patient_type == 'Child' else 'A-'
        self.serial_no.setText(f"{prefix}{max_serial + 1}")

    def on_sex_changed(self, sex_text):
        """Auto-hide pregnant field when sex is male (ប្រុស)"""
        sex = sex_text.strip().lower()
        # Check for male in Khmer (ប្រុស) or English (male, m)
        if sex in ['ប្រុស', 'male', 'm']:
            # Hide pregnant field for males
            self.pregnant.setVisible(False)
            self.pregnant.clear()
            self.pregnant.setFocusPolicy(Qt.NoFocus)  # type: ignore
        else:
            # Show pregnant field for females (if adult)
            if self.current_patient_type == "Adult":
                self.pregnant.setVisible(True)
                self.pregnant.setFocusPolicy(Qt.StrongFocus)  # type: ignore
            else:
                # Keep hidden for children
                self.pregnant.setVisible(False)
                self.pregnant.setFocusPolicy(Qt.NoFocus)  # type: ignore
        self._update_add_button_state()

    def update_area_serial(self):
        """Auto-update Area serial when category changes"""
        area_cat_txt = self.area_cat.currentText().strip()
        if area_cat_txt:
            area_next = self.get_next_counter(area_cat_txt, 8)
            self.area_val.setText(str(area_next))
        else:
            self.area_val.clear()
    
    def update_disease_serial(self):
        """Auto-update Disease serial when category changes"""
        dis_cat_txt = self.disease_cat.currentText().strip()
        if dis_cat_txt:
            dis_next = self.get_next_counter(dis_cat_txt, 13)
            self.disease_val.setText(str(dis_next))
        else:
            self.disease_val.clear()
        
    def renumber_categories(self):
        """
        Renumber categories (Area, Disease, IMCI) separately for each patient type
        រៀបលេខ Category ដាច់ដោយឡែករវាងកុមារ និងមនុស្សចាស់
        """
        db.renumber_all_patient_numbering()

    def _renumber_categories_for_rows(self, rows, patient_type):
        """
        Helper function to renumber categories for a given set of rows
        
        Args:
            rows: List of patient records to renumber
            patient_type: 'Child' or 'Adult'
        """
        if not rows:
            return
            
        area_counters = {}
        disease_counters = {}
        imci_counters = {}
        batch_update_data = []

        for row in rows:
            pid = row[0]
            p_type = row[23] if len(row) > 23 else "Child"

            # 1. រៀបលេខ Area (Col 8)
            area_raw = str(row[8])
            area_cat = area_raw.split("::")[0] if "::" in area_raw else area_raw
            area_counters[area_cat] = area_counters.get(area_cat, 0) + 1
            new_area = f"{area_cat}::{area_counters[area_cat]}"

            # 2. រៀបលេខ Disease (Col 13)
            dis_raw = str(row[13])
            dis_cat = dis_raw.split("::")[0] if "::" in dis_raw else dis_raw
            disease_counters[dis_cat] = disease_counters.get(dis_cat, 0) + 1
            new_disease = f"{dis_cat}::{disease_counters[dis_cat]}"

            # 3. រៀបលេខ IMCI (Col 18) - សម្រាប់តែ Child
            imci_raw = str(row[18])
            if p_type == 'Child':
                imci_cat = imci_raw.split("::")[0] if "::" in imci_raw else imci_raw
                imci_cat = imci_cat.strip()  # Remove whitespace
                
                # Auto-assign IMCI "Yes" for all children if empty
                if not imci_cat or imci_cat.lower() in ['', 'none', 'n/a', 'n/a']:
                    imci_cat = "Yes"
                
                if imci_cat:
                    imci_counters[imci_cat] = imci_counters.get(imci_cat, 0) + 1
                    new_imci = f"{imci_cat}::{imci_counters[imci_cat]}"
                else:
                    new_imci = ""
            else:
                new_imci = imci_raw

            batch_update_data.append((new_area, new_disease, new_imci, pid))

        # Update ចូល Database តែម្តង (Batch)
        if batch_update_data:
            db.update_counters_batch(batch_update_data)

    def format_date_input(self):
        text = self.date.text().replace("/", "").replace(".", "").replace("-", "").strip()
        if not text.isdigit():
            return
            
        formatted = ""
        if len(text) == 8: # ឧ. 27012025
            formatted = f"{text[:2]}/{text[2:4]}/{text[4:]}"
        elif len(text) == 6: # ឧ. 270125
            formatted = f"{text[:2]}/{text[2:4]}/20{text[4:]}"
        elif len(text) == 5: # ឧ. 27125 (ចេញ 27/01/2025 តាមសំណើ)
            formatted = f"{text[:2]}/0{text[2]}/20{text[3:]}"
            
        if formatted:
            # --- New: Validate the date ---
            try:
                datetime.strptime(formatted, "%d/%m/%Y")
                self.date.setText(formatted)
                self.update_next_serial_no()
            except ValueError:
                QMessageBox.warning(self, "Invalid Date", f"'{formatted}' is not a valid date. Please correct it.")
                self.date.clear()
                self.date.setFocus()
        elif text: # If user entered something that couldn't be formatted
            QMessageBox.warning(self, "Invalid Format", "Please enter the date as DDMMYY (e.g., 270125).")

    def pick_date(self):
        """បើកផ្ទាំងប្រតិទិនដើម្បីជ្រើសរើសថ្ងៃខែឆ្នាំ"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Select Date")
        dialog.setMinimumSize(400, 350)
        vbox = QVBoxLayout(dialog)
        
        cal = QCalendarWidget()
        cal.setGridVisible(True)
        cal.setLocale(QLocale(QLocale.Khmer, QLocale.Cambodia))
        cal.setFirstDayOfWeek(Qt.Monday)  # type: ignore[attr-defined]
        # រចនាប្រតិទិនឱ្យស៊ីជាមួយពណ៌កម្មវិធី
        cal.setStyleSheet("""
            QCalendarWidget QWidget {
                background-color: #2f3640;
                color: #f5f6fa;
                alternate-background-color: #222b35;
            }
            QCalendarWidget QAbstractItemView:enabled { 
                background-color: #2f3640;
                color: #f5f6fa;
                selection-background-color: #00cec9;
                selection-color: #000000;
            }
            QCalendarWidget QAbstractItemView:disabled {
                color: #8b97a3;
                background-color: #26313b;
            }
            QCalendarWidget QHeaderView::section {
                background-color: #f5f6fa;
                color: #111820;
                font-weight: bold;
                border: 1px solid #dfe6e9;
                padding: 6px 4px;
            }
            QCalendarWidget QTableView {
                alternate-background-color: #26313b;
                gridline-color: #485460;
                outline: 0;
            }
            QCalendarWidget QToolButton {
                color: #f5f6fa;
                background-color: #1e272e;
                border: 1px solid #3d5368;
                border-radius: 4px;
                padding: 5px 8px;
                font-weight: bold;
            }
            QCalendarWidget QToolButton:hover {
                background-color: #34495e;
            }
            QCalendarWidget QMenu {
                background-color: #1e272e;
                color: #f5f6fa;
                border: 1px solid #485460;
            }
            QCalendarWidget QSpinBox {
                background-color: #1e272e;
                color: #f5f6fa;
                selection-background-color: #00cec9;
                selection-color: #000000;
                border: 1px solid #485460;
                border-radius: 4px;
                padding: 3px;
            }
        """)

        # ព្យាយាមកំណត់ប្រតិទិនទៅតាមថ្ងៃដែលបានវាយក្នុង LineEdit (បើមាន)
        try:
            d = datetime.strptime(self.date.text(), "%d/%m/%Y")
            cal.setSelectedDate(QDate(d.year, d.month, d.day))
        except:
            cal.setSelectedDate(QDate.currentDate())

        vbox.addWidget(cal)
        
        def select_date():
            qdate = cal.selectedDate()
            self.date.setText(qdate.toString("dd/MM/yyyy"))
            self.update_next_serial_no()
            dialog.accept()

        cal.activated.connect(select_date) # ជ្រើសរើសដោយ Double Click
        btn_confirm = QPushButton("Confirm Selection")
        btn_confirm.setStyleSheet("background-color: #0fbcf9; color: black; font-weight: bold; padding: 10px;")
        btn_confirm.clicked.connect(select_date)
        vbox.addWidget(btn_confirm)
        dialog.exec_()

    def populate_form_from_table(self):
        if getattr(self, "_loading_table", False) or getattr(self, "_populating_form", False):
            return

        selected_items = self.table.selectedItems()
        if not selected_items:
            return
        self._populating_form = True

        row_idx = selected_items[0].row()

        # Helper to get text from table cell safely
        def get_item_text(r, c):
            item = self.table.item(r, c)
            return item.text() if item else ""

        # Helper to parse composite fields (e.g., "Category::Value")
        def parse_composite(text, cat_widget, val_widget):
            if "::" in text:
                cat, val = text.split("::", 1)
                if isinstance(cat_widget, QComboBox):
                    index = cat_widget.findText(cat)
                    if index >= 0:
                        cat_widget.setCurrentIndex(index)
                    else:
                        cat_widget.setCurrentIndex(0)
                else:
                    cat_widget.setText(cat)
                val_widget.setText(val)
            else:
                if isinstance(cat_widget, QComboBox):
                    index = cat_widget.findText(text)
                    if index >= 0:
                        cat_widget.setCurrentIndex(index)
                    else:
                        cat_widget.setCurrentIndex(0)
                else:
                    cat_widget.setText(text)
                val_widget.clear()

        # Get the hidden ID and store it
        self.selected_id = get_item_text(row_idx, 0)

        # Find all rows for this patient ID
        patient_rows = []
        for r in range(self.table.rowCount()):
            if get_item_text(r, 0) == self.selected_id:
                patient_rows.append(r)

        if not patient_rows:
            self._populating_form = False
            return

        # Find the main row (the one with serial number)
        main_row_idx = None
        for r in patient_rows:
            if get_item_text(r, 2):  # Serial column
                main_row_idx = r
                break

        if main_row_idx is None:
            self._populating_form = False
            return

        p_type = get_item_text(main_row_idx, 23) or "Child"
        self.current_patient_type = p_type
        self.type_tabs.blockSignals(True)
        self.type_tabs.setCurrentIndex(1 if p_type == "Adult" else 0)
        self.type_tabs.blockSignals(False)
        self.update_age_category_options()

        # Populate fields from the main row
        self.date.setText(get_item_text(main_row_idx, 1))
        self.serial_no.setText(get_item_text(main_row_idx, 2))
        self.card_id.setText(get_item_text(main_row_idx, 3))
        self.name.setText(get_item_text(main_row_idx, 4))
        self.guardian.setText(get_item_text(main_row_idx, 5))
        parse_composite(get_item_text(main_row_idx, 6), self.age_cat, self.age_val)
        self._set_combo_current_text_or_blank(self.sex, get_item_text(main_row_idx, 7))
        parse_composite(get_item_text(main_row_idx, 8), self.area_cat, self.area_val)
        self.pregnant.setText(get_item_text(main_row_idx, 9))
        
        # Update pregnant field visibility based on sex
        self.on_sex_changed(self.sex.currentText())
        
        self.address.setText(get_item_text(main_row_idx, 10))
        self.phone.setText(get_item_text(main_row_idx, 11))
        self.ref_from.setText(get_item_text(main_row_idx, 12))
        parse_composite(get_item_text(main_row_idx, 13), self.disease_cat, self.disease_val)
        self.symptoms.setText(get_item_text(main_row_idx, 14))
        self.paraclinical.setText(get_item_text(main_row_idx, 15))
        self._set_combo_current_text_or_blank(self.diagnosis, get_item_text(main_row_idx, 16))
        # IMCI is now auto-generated, just display the value
        imci_text = get_item_text(main_row_idx, 18)
        if imci_text and "::" in imci_text:
            imci_parts = imci_text.split("::")
            self.imci_val.setText(imci_parts[1])  # Show only the number
        elif imci_text and imci_text.startswith("IMCI-"):
            self.imci_val.setText(imci_text.replace("IMCI-", ""))  # Show only the number
        else:
            self.imci_val.clear()
        self._parse_nutrition_string(get_item_text(main_row_idx, 19))
        self.ref_to.setText(get_item_text(main_row_idx, 20))
        self._set_combo_current_text_or_blank(self.service, get_item_text(main_row_idx, 21))
        self.remark.setText(get_item_text(main_row_idx, 22))

        # Set Patient Type
        p_type = get_item_text(main_row_idx, 23) or "Child"
        self.current_patient_type = p_type
        self.type_tabs.blockSignals(True)
        if p_type == "Adult":
            self.type_tabs.setCurrentIndex(1)
            # Show pregnant field for Adults
            self.pregnant.setVisible(True)
            self.pregnant.setFocusPolicy(Qt.StrongFocus)  # type: ignore
        else:
            self.type_tabs.setCurrentIndex(0)
            # Hide pregnant field for Children
            self.pregnant.setVisible(False)
            self.pregnant.clear()
            self.pregnant.setFocusPolicy(Qt.NoFocus)  # type: ignore
        self.type_tabs.blockSignals(False)

        # Reconstruct the full treatment string from all related rows
        treatments = [get_item_text(r, 17) for r in patient_rows if get_item_text(r, 17)]
        self.treatment.setText(", ".join(treatments))
        self._populating_form = False
        self._update_add_button_state()

    def clear_inputs(self):
        self.date.clear()
        # The serial number is now set by update_next_serial_no()
        self.card_id.clear()
        self.name.clear()
        self.guardian.clear()
        if isinstance(self.age_cat, QComboBox):
            self.age_cat.setCurrentIndex(0)
        else:
            self.age_cat.clear()
        self.age_val.clear()
        self.sex.setCurrentIndex(0)

        # Reset dropdowns to blank selections.
        self.area_val.clear()
        self._set_combo_items_blank_first(self.area_cat, self.CAT_AREA)
        self.area_cat.setCurrentIndex(0)

        self.pregnant.clear()
        self.address.clear()
        self.phone.clear()
        self.ref_from.clear()

        # Reset Disease to a blank selection.
        self.disease_val.clear()
        self._set_combo_items_blank_first(self.disease_cat, self.CAT_DISEASE)
        self.disease_cat.setCurrentIndex(0)

        self.symptoms.clear()
        self.paraclinical.clear()
        self._set_combo_items_blank_first(self.diagnosis, self.CAT_DIAGNOSIS)
        self.diagnosis.setCurrentIndex(0)
        if self.treatment: self.treatment.clear()
        self.imci_val.clear()
        self.nut_wa.clear()
        self.nut_wh.clear()
        self.ref_to.clear()
        self.service.setCurrentIndex(0)
        self.remark.clear()

        if self.type_tabs.currentIndex() == 0:
             self.current_patient_type = "Child"
        else:
             self.current_patient_type = "Adult"

        self.selected_id = None # Clear selected ID before generating the next new serial
        self.update_next_serial_no()

        self.table.blockSignals(True)
        self.table.clearSelection() # Remove visual selection from table
        self.table.blockSignals(False)
        self._update_add_button_state()

    def reset_database(self):
        reply = QMessageBox.question(self, 'Reset Database', 
                                     'តើអ្នកពិតជាចង់លុបទិន្នន័យទាំងអស់ចេញពីប្រព័ន្ធមែនទេ?\n(សកម្មភាពនេះមិនអាចត្រឡប់វិញបានទេ!)',
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            db.delete_all(self.active_branch_code)
            self.view()
            self.clear_inputs()
            scope = "all branches" if self.is_admin else f"branch {self.branch_code}"
            self.statusBar.showMessage(f"Database has been reset for {scope}.", 5000)
            # QMessageBox.information(self, "Success", "ទិន្នន័យទាំងអស់ត្រូវបានលុបសម្អាត!")

    def open_change_password(self):
        """Open the change password dialog"""
        dialog = ChangePasswordDialog(self.current_user)
        if dialog.exec_() == QDialog.Accepted:
            self.statusBar.showMessage("ពាក្យសម្ងាត់ត្រូវបានប្តូរជោគជ័យ! (Password changed successfully)", 5000)

    def _on_focus_changed(self, old, new):
        """ប្តូរភាសាក្តារចុច Windows ដោយស្វ័យប្រវត្តិតាមការកំណត់របស់អ្នកប្រើប្រាស់"""
        if sys.platform != "win32" or not new:
            return

        # បញ្ជីប្រអប់បញ្ចូលដែលត្រូវប្រើភាសាខ្មែរ
        khmer_fields = [
            self.name, self.guardian, self.age_cat, self.sex,
            self.area_cat, self.address, self.ref_from, self.disease_cat,
            self.symptoms, self.remark
        ]

        if new in khmer_fields:
            # ប្តូរទៅភាសាខ្មែរ (Khmer Layout ID: 0x0453)
            ctypes.windll.user32.ActivateKeyboardLayout(0x0453, 0)
        elif isinstance(new, (QLineEdit, QTextEdit)):
            # ប្តូរទៅភាសាអង់គ្លេស (English Layout ID: 0x0409) សម្រាប់ប្រអប់ផ្សេងៗទៀត
            ctypes.windll.user32.ActivateKeyboardLayout(0x0409, 0)

    def check_for_updates(self):
        """Check GitHub for a newer version and update safely."""
        update_url = "https://raw.githubusercontent.com/saratboy1988-a11y/Clinic-Update/main/version.json"

        self.statusBar.showMessage("Checking for updates...")
        try:
            request = urllib.request.Request(
                update_url,
                headers={"User-Agent": f"ClinicManager/{APP_VERSION}"}
            )
            with _urlopen_update_request(request, timeout=10) as response:
                data = json.loads(response.read().decode("utf-8"))

            remote_version = str(data.get("version", "")).strip()
            files_to_update = data.get("files", {})
            installer_info = data.get("installer", {})

            if not remote_version:
                raise ValueError("Update manifest is missing a version.")
            if not isinstance(files_to_update, dict) or not files_to_update:
                if any(key in data for key in ("upload_date", "patient_count", "period", "upload_by")):
                    raise ValueError(
                        "The remote version.json is a Cloud Sync metadata file, not an update manifest. "
                        "Restore the updater manifest on GitHub."
                    )
                raise ValueError("Invalid update manifest.")

            remote_v = self._parse_version(remote_version)
            local_v = self._parse_version(APP_VERSION)

            if remote_v <= local_v:
                QMessageBox.information(
                    self,
                    "No Update",
                    f"Current version is already the latest.\n\nVersion: {APP_VERSION}"
                )
                self.statusBar.showMessage("Application is already up to date.")
                return

            if getattr(sys, "frozen", False):
                installer_url = ""
                installer_name = f"ClinicManager_Setup_{remote_version}.exe"
                silent_args = ["/SP-", "/VERYSILENT", "/SUPPRESSMSGBOXES", "/NORESTART"]
                if isinstance(installer_info, dict):
                    installer_url = str(installer_info.get("url", "")).strip()
                    installer_name = os.path.basename(str(installer_info.get("file_name", installer_name))) or installer_name
                    custom_args = installer_info.get("silent_args")
                    if isinstance(custom_args, list) and custom_args:
                        silent_args = [str(arg) for arg in custom_args]

                if not installer_url:
                    QMessageBox.information(
                        self,
                        "Update Not Applicable",
                        "A newer version was found, but no installer is available for this EXE installation.\n\n"
                        f"Current: {APP_VERSION}\nNew: {remote_version}"
                    )
                    self.statusBar.showMessage("Update was found but no installer is available.")
                    return

                details = [
                    f"New version found: {remote_version}",
                    f"Current version: {APP_VERSION}",
                    "",
                    "The installer will be downloaded and started automatically.",
                    "The application will close during installation.",
                    "",
                    "Do you want to update now?"
                ]
                reply = QMessageBox.question(
                    self,
                    "Update Available",
                    "\n".join(details),
                    QMessageBox.Yes | QMessageBox.No
                )
                if reply == QMessageBox.Yes:
                    self.perform_installer_update(installer_url, installer_name, remote_version, silent_args)
                return

            applicable_files, skipped_files = self._filter_update_files(files_to_update)
            if not applicable_files:
                skipped_text = ", ".join(skipped_files) if skipped_files else "all files"
                QMessageBox.information(
                    self,
                    "Update Not Applicable",
                    "A newer version was found, but this installation cannot apply the available files.\n\n"
                    f"Current: {APP_VERSION}\nNew: {remote_version}\nSkipped: {skipped_text}"
                )
                self.statusBar.showMessage("Update was found but is not applicable to this installation.")
                return

            details = [
                f"New version found: {remote_version}",
                f"Current version: {APP_VERSION}",
                "",
                "Files to update:",
                "\n".join(f"- {name}" for name in applicable_files.keys()),
            ]
            if skipped_files:
                details.extend([
                    "",
                    "Files skipped to protect this installation:",
                    "\n".join(f"- {name}" for name in skipped_files),
                ])

            details.append("")
            details.append("Do you want to update now?")

            reply = QMessageBox.question(
                self,
                "Update Available",
                "\n".join(details),
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                self.perform_update(applicable_files, remote_version, skipped_files)
        except Exception as e:
            QMessageBox.warning(self, "Update Error", f"Could not check for updates:\n{str(e)}")
            self.statusBar.showMessage("Update check failed.")

    def _parse_version(self, version_text):
        parts = [int(part) for part in re.findall(r"\d+", str(version_text))]
        if not parts:
            raise ValueError(f"Invalid version: {version_text}")
        while len(parts) < 3:
            parts.append(0)
        return tuple(parts[:3])

    def _filter_update_files(self, files_dict):
        applicable_files = {}
        skipped_files = []
        is_frozen = getattr(sys, "frozen", False)

        for file_name, url in files_dict.items():
            normalized_name = os.path.basename(str(file_name))
            lower_name = normalized_name.lower()

            if lower_name == "settings.ini":
                skipped_files.append(normalized_name)
                continue

            if is_frozen and lower_name.endswith(".py"):
                skipped_files.append(normalized_name)
                continue

            applicable_files[normalized_name] = url

        return applicable_files, skipped_files

    def _get_update_target_path(self, file_name):
        return os.path.join(self.base_dir, file_name)

    def _download_update_file(self, url, target_path):
        request = urllib.request.Request(
            url,
            headers={"User-Agent": f"ClinicManager/{APP_VERSION}"}
        )
        with _urlopen_update_request(request, timeout=20) as response:
            with open(target_path, "wb") as handle:
                shutil.copyfileobj(response, handle)

    def _restart_after_update(self):
        if getattr(sys, "frozen", False):
            subprocess.Popen([sys.executable], cwd=self.base_dir)
        else:
            subprocess.Popen([sys.executable, os.path.abspath(sys.argv[0])], cwd=self.base_dir)
        QApplication.quit()

    def perform_installer_update(self, installer_url, installer_name, remote_version, silent_args=None):
        """Download a setup EXE and run it after this process exits."""
        silent_args = silent_args or ["/SP-", "/VERYSILENT", "/SUPPRESSMSGBOXES", "/NORESTART"]
        temp_dir = tempfile.mkdtemp(prefix="clinic_installer_update_", dir=tempfile.gettempdir())
        installer_path = os.path.join(temp_dir, installer_name)
        runner_path = os.path.join(temp_dir, "run_update.bat")

        try:
            self.statusBar.showMessage(f"Downloading installer {installer_name}...")
            self._download_update_file(installer_url, installer_path)

            if not os.path.exists(installer_path) or os.path.getsize(installer_path) == 0:
                raise ValueError("Downloaded installer is empty or missing.")

            exe_path = sys.executable
            args_text = " ".join(f'"{arg}"' for arg in silent_args)
            runner = f"""@echo off
timeout /t 2 /nobreak >nul
start "" /wait "{installer_path}" {args_text}
if exist "{exe_path}" start "" "{exe_path}"
rmdir /s /q "{temp_dir}"
"""
            with open(runner_path, "w", encoding="utf-8") as handle:
                handle.write(runner)

            QMessageBox.information(
                self,
                "Update Ready",
                f"Installer for version {remote_version} has been downloaded.\n\n"
                "The application will now close and the installer will start."
            )
            subprocess.Popen(
                ["cmd", "/c", runner_path],
                cwd=temp_dir,
                creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0)
            )
            QApplication.quit()
        except Exception as e:
            shutil.rmtree(temp_dir, ignore_errors=True)
            QMessageBox.critical(self, "Update Failed", f"Installer update failed:\n{str(e)}")
            self.statusBar.showMessage("Installer update failed.")

    def perform_update(self, files_dict, remote_version, skipped_files=None):
        """Download all files first, then replace them with rollback protection."""
        skipped_files = skipped_files or []
        temp_dir = tempfile.mkdtemp(prefix="clinic_update_", dir=self.base_dir)
        backup_dir = os.path.join(
            self.base_dir,
            "update_backup",
            datetime.now().strftime("%Y%m%d_%H%M%S")
        )
        downloaded_files = {}
        replaced_files = []

        try:
            for file_name, url in files_dict.items():
                self.statusBar.showMessage(f"Downloading {file_name}...")
                temp_path = os.path.join(temp_dir, file_name)
                self._download_update_file(url, temp_path)
                downloaded_files[file_name] = temp_path

            os.makedirs(backup_dir, exist_ok=True)

            for file_name, temp_path in downloaded_files.items():
                target_path = self._get_update_target_path(file_name)
                backup_path = os.path.join(backup_dir, file_name)

                if os.path.exists(target_path):
                    shutil.copy2(target_path, backup_path)

                os.replace(temp_path, target_path)
                replaced_files.append((target_path, backup_path))

            message_lines = [
                f"Updated to version {remote_version} successfully.",
                "",
                "The application will now restart."
            ]
            if skipped_files:
                message_lines.extend([
                    "",
                    "Skipped files:",
                    "\n".join(f"- {name}" for name in skipped_files)
                ])

            QMessageBox.information(self, "Update Complete", "\n".join(message_lines))
            self._restart_after_update()
        except Exception as e:
            for target_path, backup_path in reversed(replaced_files):
                if os.path.exists(backup_path):
                    shutil.copy2(backup_path, target_path)

            QMessageBox.critical(self, "Update Failed", f"Update failed:\n{str(e)}")
            self.statusBar.showMessage("Update failed.")
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def logout(self):
        # Improved logout: just close the main window. The main loop will handle re-showing the login dialog.
        self.close()

    def closeEvent(self, a0):
        print("Main window is closing.")
        # --- Auto Backup Logic ---
        try:
            if not os.path.exists(self.backup_dir):
                os.makedirs(self.backup_dir)

            backup_name = os.path.join(self.backup_dir, f"auto_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db")
            shutil.copy(db.DB_NAME, backup_name)
            print(f"Auto-backup created: {backup_name}")

            # --- Clean up old backups (Keep only last 10) ---
            backups = sorted(
                [os.path.join(self.backup_dir, f) for f in os.listdir(self.backup_dir) if f.endswith(".db")],
                key=os.path.getmtime
            )
            if len(backups) > 10:
                files_to_delete = backups[:-10]
                for f in files_to_delete:
                    os.remove(f)
                    print(f"Deleted old backup: {f}")

        except Exception as e:
            print(f"Auto-backup failed: {e}")
        if a0:
            a0.accept()

# ============================================================================
# Advanced Query Dialog Class
# ============================================================================

class AdvancedQueryDialog(QDialog):
    """
    Advanced Query Dialog for multi-criteria patient search
    Allows filtering by: Disease, Area, Patient Type, Sex
    """

    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_app = parent
        self.setWindowTitle("🔍 ស្វែងរកកម្រិតខ្ពស់ (Advanced Query)")
        self.setMinimumSize(900, 700)
        self.init_ui()

    def init_ui(self):
        """Initialize the dialog UI"""
        layout = QVBoxLayout()
        layout.setSpacing(10)

        # Title
        title_label = QLabel("🔍 ស្វែងរកទិន្នន័យតាមលក្ខខណ្ឌច្រើន")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; color: #00cec9; padding: 10px;")
        layout.addWidget(title_label)

        # Filter Group Box
        filter_group = QGroupBox("⚙️ លក្ខខណ្ឌស្វែងរក (Filter Criteria)")
        filter_group.setStyleSheet("""
            QGroupBox {
                background-color: #2f3640;
                border: 2px solid #00cec9;
                border-radius: 10px;
                margin-top: 10px;
                padding: 15px;
            }
            QGroupBox::title {
                color: #00cec9;
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 5px;
            }
        """)
        filter_layout = QGridLayout()

        # Disease Filter
        filter_layout.addWidget(QLabel("ករណីជំងឺ (Disease):"), 0, 0)
        self.disease_combo = QComboBox()
        self.disease_combo.addItem("— ទាំងអស់ (All) —")
        if self.parent_app and hasattr(self.parent_app, 'CAT_DISEASE'):
            self.disease_combo.addItems(self.parent_app.CAT_DISEASE)
        filter_layout.addWidget(self.disease_combo, 0, 1)

        # Area Filter
        filter_layout.addWidget(QLabel("តំបន់ (Area):"), 1, 0)
        self.area_combo = QComboBox()
        self.area_combo.addItem("— ទាំងអស់ (All) —")
        if self.parent_app and hasattr(self.parent_app, 'CAT_AREA'):
            self.area_combo.addItems(self.parent_app.CAT_AREA)
        filter_layout.addWidget(self.area_combo, 1, 1)

        # Patient Type Filter
        filter_layout.addWidget(QLabel("ប្រភេទ (Type):"), 2, 0)
        self.type_combo = QComboBox()
        self.type_combo.addItem("— ទាំងអស់ (All) —")
        self.type_combo.addItems(["Child (កុមារ)", "Adult (មនុស្សចាស់)"])
        filter_layout.addWidget(self.type_combo, 2, 1)

        # Sex Filter
        filter_layout.addWidget(QLabel("ភេទ (Sex):"), 3, 0)
        self.sex_combo = QComboBox()
        self.sex_combo.addItem("— ទាំងអស់ (All) —")
        if self.parent_app and hasattr(self.parent_app, 'CAT_SEX'):
            self.sex_combo.addItems(self.parent_app.CAT_SEX)
        filter_layout.addWidget(self.sex_combo, 3, 1)

        # Date Range (Optional)
        filter_layout.addWidget(QLabel("ពីថ្ងៃ (From):"), 0, 2)
        self.from_date = QDateEdit()
        self.from_date.setCalendarPopup(True)
        self.from_date.setDisplayFormat("dd/MM/yyyy")
        self.from_date.setDate(QDate.currentDate().addMonths(-1))
        filter_layout.addWidget(self.from_date, 0, 3)

        filter_layout.addWidget(QLabel("ដល់ថ្ងៃ (To):"), 1, 2)
        self.to_date = QDateEdit()
        self.to_date.setCalendarPopup(True)
        self.to_date.setDisplayFormat("dd/MM/yyyy")
        self.to_date.setDate(QDate.currentDate())
        filter_layout.addWidget(self.to_date, 1, 3)

        filter_group.setLayout(filter_layout)
        layout.addWidget(filter_group)

        # Search Button
        search_btn = QPushButton("🔍 ស្វែងរកឥឡូវ (Search Now)")
        search_btn.setStyleSheet("""
            QPushButton {
                background-color: #00cec9;
                color: black;
                font-weight: bold;
                font-size: 16px;
                padding: 12px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #00b894;
            }
        """)
        search_btn.clicked.connect(self.execute_query)
        layout.addWidget(search_btn)

        # Result Label
        self.result_label = QLabel("ស្រេច: 0 នាក់")
        self.result_label.setStyleSheet("font-size: 20px; font-weight: bold; color: #fdcb6e; padding: 10px; text-align: center;")
        self.result_label.setAlignment(Qt.AlignCenter)  # type: ignore
        layout.addWidget(self.result_label)

        # Results Table
        results_label = QLabel("📋 បញ្ជីអ្នកជំងឺ (Patient List):")
        results_label.setStyleSheet("font-size: 16px; font-weight: bold; color: #00cec9; padding: 5px;")
        layout.addWidget(results_label)

        self.results_table = QTableWidget()
        self.results_table.setColumnCount(9)
        self.results_table.setHorizontalHeaderLabels([
            "លេខរៀង", "ឈ្មោះ", "ភេទ", "អាយុ", "តំបន់", "ជំងឺ", "រោគវិនិច្ឆ័យ", "កាលបរិច្ឆេទ", "សេវា"
        ])
        self.results_table.setEditTriggers(QTableWidget.NoEditTriggers)  # type: ignore
        self.results_table.setSortingEnabled(True)  # type: ignore
        self.results_table.setAlternatingRowColors(True)  # type: ignore
        # Fix text color visibility
        self.results_table.setStyleSheet("""
            QTableWidget {
                background-color: #2f3640;
                alternate-background-color: #353b48;
                color: #ffffff;
                gridline-color: #576574;
            }
            QTableWidget::item {
                color: #ffffff;
                padding: 5px;
            }
            QHeaderView::section {
                background-color: #00cec9;
                color: #000000;
                padding: 8px;
                border: 1px solid #576574;
                font-weight: bold;
            }
        """)
        layout.addWidget(self.results_table)

        # Export Button
        export_btn = QPushButton("📊 Export ទៅ Excel")
        export_btn.setStyleSheet("""
            QPushButton {
                background-color: #20bf6b;
                color: white;
                font-weight: bold;
                font-size: 14px;
                padding: 10px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #26de81;
            }
        """)
        export_btn.clicked.connect(self.export_results)
        layout.addWidget(export_btn)

        # Close Button
        close_btn = QPushButton("បិទ (Close)")
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                font-weight: bold;
                font-size: 14px;
                padding: 10px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

        self.setLayout(layout)
        self.setStyleSheet("""
            QDialog {
                background-color: #1e272e;
            }
            QLabel {
                color: #f5f6fa;
                background: transparent;
            }
            QComboBox, QDateEdit {
                background-color: #f5f6fa;
                color: #000000;
                border: 1px solid #00cec9;
                border-radius: 5px;
                padding: 6px 8px;
                min-height: 18px;
                selection-background-color: #dfe6e9;
                selection-color: #000000;
            }
            QComboBox QAbstractItemView {
                background-color: #ffffff;
                color: #000000;
                selection-background-color: #74b9ff;
                selection-color: #000000;
                border: 1px solid #00cec9;
            }
            QDateEdit::drop-down, QComboBox::drop-down {
                border: none;
                background: transparent;
                width: 24px;
            }
            QDateEdit::down-arrow, QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 6px solid #2d3436;
                margin-right: 8px;
            }
        """)

    def execute_query(self):
        """Execute the advanced query and display results"""
        try:
            # Get filter values
            disease = self.disease_combo.currentText()
            area = self.area_combo.currentText()
            patient_type = self.type_combo.currentText()
            sex = self.sex_combo.currentText()
            branch_code = getattr(self.parent_app, "active_branch_code", None)

            # Get all patients
            if patient_type.startswith("Child"):
                all_patients = db.view_by_patient_type("Child", branch_code) or []
            elif patient_type.startswith("Adult"):
                all_patients = db.view_by_patient_type("Adult", branch_code) or []
            else:
                all_patients = db.view(branch_code) or []

            # Apply filters
            filtered = []
            for row in all_patients:
                match = True

                # Filter by disease
                if disease != "— ទាំងអស់ (All) —":
                    disease_raw = str(row[db.PatientCol.DISEASE])
                    disease_cat = disease_raw.split("::")[0] if "::" in disease_raw else disease_raw
                    if disease_cat != disease:
                        match = False

                # Filter by area
                if match and area != "— ទាំងអស់ (All) —":
                    area_raw = str(row[db.PatientCol.AREA])
                    area_cat = area_raw.split("::")[0] if "::" in area_raw else area_raw
                    if area_cat != area:
                        match = False

                # Filter by sex
                if match and sex != "— ទាំងអស់ (All) —":
                    if row[db.PatientCol.SEX] != sex:
                        match = False

                # Filter by date range
                if match:
                    try:
                        date_str = row[db.PatientCol.DATE]
                        from PyQt5.QtCore import QDate
                        patient_date = QDate.fromString(date_str, "dd/MM/yyyy")
                        from_date = self.from_date.date()
                        to_date = self.to_date.date()
                        
                        if patient_date.isValid() and from_date.isValid() and to_date.isValid():
                            if patient_date < from_date or patient_date > to_date:
                                match = False
                    except:
                        pass  # Skip date filter if invalid

                if match:
                    filtered.append(row)

            # Update result label
            self.result_label.setText(f"✅ ស្រេច: {len(filtered)} នាក់")

            # Display in table
            self.results_table.setRowCount(len(filtered))
            for i, row in enumerate(filtered):
                # Extract area category
                area_raw = str(row[db.PatientCol.AREA])
                area_cat = area_raw.split("::")[0] if "::" in area_raw else area_raw

                # Extract disease category
                disease_raw = str(row[db.PatientCol.DISEASE])
                disease_cat = disease_raw.split("::")[0] if "::" in disease_raw else disease_raw

                self.results_table.setItem(i, 0, QTableWidgetItem(str(row[db.PatientCol.SERIAL])))
                self.results_table.setItem(i, 1, QTableWidgetItem(str(row[db.PatientCol.NAME])))
                self.results_table.setItem(i, 2, QTableWidgetItem(str(row[db.PatientCol.SEX])))
                self.results_table.setItem(i, 3, QTableWidgetItem(str(row[db.PatientCol.AGE])))
                self.results_table.setItem(i, 4, QTableWidgetItem(area_cat))
                self.results_table.setItem(i, 5, QTableWidgetItem(disease_cat))
                self.results_table.setItem(i, 6, QTableWidgetItem(str(row[db.PatientCol.DIAGNOSIS])))
                self.results_table.setItem(i, 7, QTableWidgetItem(str(row[db.PatientCol.DATE])))
                self.results_table.setItem(i, 8, QTableWidgetItem(str(row[db.PatientCol.SERVICE])))

        except Exception as e:
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Error", f"មិនអាចស្វែងរកបានទេ: {str(e)}")

    def export_results(self):
        """Export results to Excel"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill
            from PyQt5.QtWidgets import QFileDialog, QMessageBox

            if self.results_table.rowCount() == 0:
                QMessageBox.warning(self, "Warning", "គ្មានទិន្នន័យសម្រាប់ Export!")
                return

            wb = Workbook()
            ws = wb.active
            ws.title = "Advanced Query Results"  # type: ignore

            # Title
            ws.merge_cells('A1:I1')  # type: ignore
            title_cell = ws['A1']  # type: ignore
            title_cell.value = "🔍 ស្វែងរកកម្រិតខ្ពស់ (Advanced Query Results)"
            title_cell.font = Font(bold=True, size=16)
            title_cell.alignment = Alignment(horizontal='center')  # type: ignore

            # Headers
            headers = ["លេខរៀង", "ឈ្មោះ", "ភេទ", "អាយុ", "តំបន់", "ជំងឺ", "រោគវិនិច្ឆ័យ", "កាលបរិច្ឆេទ", "សេវា"]
            header_fill = PatternFill(start_color='FF00CEC9', end_color='FF00CEC9', fill_type='solid')
            header_font = Font(bold=True, color='FFFFFFFF')
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=3, column=col, value=header)  # type: ignore
                cell.fill = header_fill  # type: ignore
                cell.font = header_font  # type: ignore
                cell.alignment = Alignment(horizontal='center')  # type: ignore

            # Data
            for row_idx in range(self.results_table.rowCount()):
                for col_idx in range(self.results_table.columnCount()):
                    item = self.results_table.item(row_idx, col_idx)
                    value = item.text() if item else ""
                    ws.cell(row=row_idx + 4, column=col_idx + 1, value=value).alignment = Alignment(horizontal='center')  # type: ignore

            # Save
            filename = f"AdvancedQuery_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
            file_path, _ = QFileDialog.getSaveFileName(self, "Save Excel", filename, "Excel Files (*.xlsx)")
            
            if file_path:
                wb.save(file_path)
                QMessageBox.information(self, "ជោគជ័យ", f"បាន Export ទិន្នន័យ {self.results_table.rowCount()} នាក់ ទៅ Excel!\n\n{os.path.basename(file_path)}")

        except Exception as e:
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.critical(self, "Error", f"Export បរាជ័យ: {str(e)}")


# ============================================================================
# Main Application Entry Point
# ============================================================================

if __name__ == "__main__":
    print("========================================")
    print("    CLINIC MANAGEMENT SYSTEM STARTING   ")
    print("========================================")
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)  # type: ignore
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)  # type: ignore
    app = QApplication(sys.argv)
    app.setStyleSheet(MESSAGE_BOX_STYLESHEET)
    font_db = QFontDatabase()
    for font_name in [
        "Noto Sans Khmer",
        "Khmer OS Battambang",
        "Khmer OS Siemreap",
        "Khmer OS",
        "Khmer UI",
        "Leelawadee UI",
        "Segoe UI"
    ]:
        if font_name in font_db.families():
            app.setFont(QFont(font_name, 12))
            break
    
    # --- New Application Loop for better Logout/Login flow ---
    # --- ត្រួតពិនិត្យ License ជាមុនសិន ---
    mid = get_machine_id()
    lic_info = db.get_license_info()
    online_valid, _online_msg = validate_saved_online_license(mid)

    needs_activation = False
    # Initialize local license metadata; app use still requires a valid license.
    if not lic_info or not lic_info[1]:
        # បើទើបដំឡើងដំបូង កត់ត្រាថ្ងៃ install
        db.save_license_info(datetime.now().strftime("%Y-%m-%d"), "", "", mid)
        lic_info = db.get_license_info()

    if not lic_info:
        QMessageBox.critical(None, "Error", "Failed to initialize license information.")  # type: ignore
        sys.exit(1)

    current_key = lic_info[2]
    current_email = lic_info[3]
    
    if online_valid:
        needs_activation = False
    elif not current_key:
        needs_activation = True
    elif is_online_license_key(current_key):
        needs_activation = True
    else:
        is_valid, _ = validate_license(current_email, mid, current_key)
        if not is_valid:
            needs_activation = True

    if needs_activation:
        lic_dlg = LicenseDialog(mid)
        if lic_dlg.exec_() != QDialog.Accepted:
            sys.exit()

    while True:
        login = LoginDialog()
        if login.exec_() == QDialog.Accepted:
            print("Login Successful! Opening Main Window...")
            username = login.user.text()
            win = App(username, login.user_context)
            win.show()
            app.exec_() # This blocks until the window is closed (by logout or 'X' button)
            print("Window closed. Returning to login screen...")
        else:
            # User cancelled the login dialog or closed it
            print("Login Cancelled. Exiting Application.")
            break # Exit the while loop and the application
    
    sys.exit()
